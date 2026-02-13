"""Tests for document_converter functions."""

import pytest
from docx import Document

from aegis.etls.key_themes.document_converter import (
    auto_bold_html_metrics,
    validate_document_content,
    HTMLToDocx,
    add_theme_header_with_background,
    get_standard_report_metadata,
)


# ---------------------------------------------------------------------------
# auto_bold_html_metrics
# ---------------------------------------------------------------------------
class TestAutoBoldHtmlMetrics:
    """Tests for auto_bold_html_metrics()."""

    def test_bolds_percentage(self):
        result = auto_bold_html_metrics("Revenue grew 5.2% this quarter.")
        assert "<b>5.2%</b>" in result

    def test_bolds_dollar_with_scale(self):
        result = auto_bold_html_metrics("NII was $5.2 BN.")
        assert "<b>$5.2 BN</b>" in result

    def test_bolds_basis_points(self):
        result = auto_bold_html_metrics("Spread widened by 15 bps.")
        assert "<b>15 bps</b>" in result

    def test_does_not_double_bold(self):
        result = auto_bold_html_metrics("Revenue grew <b>5%</b> this quarter.")
        assert result.count("<b>5%</b>") == 1
        assert "<b><b>" not in result

    def test_empty_string(self):
        assert auto_bold_html_metrics("") == ""

    def test_none_returns_none(self):
        assert auto_bold_html_metrics(None) is None

    def test_no_metrics_unchanged(self):
        text = "The bank reported strong results."
        assert auto_bold_html_metrics(text) == text

    def test_negative_dollar_with_scale(self):
        result = auto_bold_html_metrics("Loss was -$1.2 BN.")
        assert "<b>-$1.2 BN</b>" in result

    def test_negative_dollar_without_scale(self):
        result = auto_bold_html_metrics("Declined by -$500.")
        assert "<b>-$500</b>" in result

    def test_dollar_without_scale(self):
        result = auto_bold_html_metrics("Revenue was $1,200 million.")
        assert "<b>$1,200</b>" in result

    def test_inside_span_tags_not_double_bolded(self):
        text = '<span style="color: #1e4d8b;">NIM was <b>1.65%</b> for Q4</span>'
        result = auto_bold_html_metrics(text)
        # 1.65% already bolded, should not get double-wrapped
        assert result.count("<b>1.65%</b>") == 1

    def test_multiple_metrics_in_one_line(self):
        result = auto_bold_html_metrics("NII grew 5% to $5.2 BN with 15 bps expansion.")
        assert "<b>5%</b>" in result
        assert "<b>$5.2 BN</b>" in result
        assert "<b>15 bps</b>" in result


# ---------------------------------------------------------------------------
# validate_document_content
# ---------------------------------------------------------------------------
class TestValidateDocumentContent:
    """Tests for validate_document_content()."""

    def test_valid_document_passes(self):
        doc = Document()
        add_theme_header_with_background(doc, 1, "Net Interest Margin")
        para = doc.add_paragraph()
        para.add_run(
            "The CFO discussed NIM trends and expects expansion "
            "to 1.70% by mid next year as deposit costs normalize."
        )
        # Should not raise
        validate_document_content(doc)

    def test_empty_document_raises(self):
        doc = Document()
        with pytest.raises(ValueError, match="no paragraphs"):
            validate_document_content(doc)

    def test_missing_theme_header_raises(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run(
            "The CFO discussed NIM trends and expects expansion " "to 1.70% by mid next year."
        )
        with pytest.raises(ValueError, match="theme header"):
            validate_document_content(doc)

    def test_missing_body_text_raises(self):
        doc = Document()
        add_theme_header_with_background(doc, 1, "Credit Quality")
        with pytest.raises(ValueError, match="body text"):
            validate_document_content(doc)


# ---------------------------------------------------------------------------
# HTMLToDocx
# ---------------------------------------------------------------------------
class TestHTMLToDocx:
    """Tests for HTMLToDocx parser."""

    def test_plain_text(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed("Plain text here")
        parser.close()
        assert para.text == "Plain text here"

    def test_bold_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed("Revenue grew <b>5%</b> this quarter")
        parser.close()
        assert para.text == "Revenue grew 5% this quarter"
        bold_runs = [r for r in para.runs if r.font.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "5%"

    def test_italic_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed("<i>Personal Banking</i> grew strongly")
        parser.close()
        assert para.text == "Personal Banking grew strongly"
        italic_runs = [r for r in para.runs if r.font.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "Personal Banking"

    def test_underline_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed("We are <u>committed to reaching 1.80%</u> by year end")
        parser.close()
        assert "committed to reaching 1.80%" in para.text
        underline_runs = [r for r in para.runs if r.font.underline]
        assert len(underline_runs) == 1

    def test_highlight_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed('<mark style="background-color: #ffff99;">Critical insight</mark>')
        parser.close()
        assert "Critical insight" in para.text

    def test_span_color_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed(
            '<span style="color: #1e4d8b; font-size: 11pt; font-weight: bold;">'
            "Key question</span>"
        )
        parser.close()
        assert "Key question" in para.text
        # Should have color applied
        assert len(para.runs) >= 1

    def test_nested_tags(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed("Before <b>bold with <i>italic inside</i> still bold</b> after")
        parser.close()
        full_text = para.text
        assert "Before" in full_text
        assert "bold with" in full_text
        assert "italic inside" in full_text
        assert "after" in full_text

    def test_br_tag_inserts_newline(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed("Line one<br>Line two")
        parser.close()
        assert "Line one\nLine two" in para.text

    def test_self_closing_br_tag(self):
        doc = Document()
        para = doc.add_paragraph()
        parser = HTMLToDocx(para)
        parser.feed("Line one<br/>Line two")
        parser.close()
        assert "Line one\nLine two" in para.text


# ---------------------------------------------------------------------------
# get_standard_report_metadata
# ---------------------------------------------------------------------------
class TestGetStandardReportMetadata:
    """Tests for get_standard_report_metadata()."""

    def test_returns_required_keys(self):
        metadata = get_standard_report_metadata()
        assert "report_name" in metadata
        assert "report_description" in metadata
        assert "report_type" in metadata

    def test_report_type_is_key_themes(self):
        metadata = get_standard_report_metadata()
        assert metadata["report_type"] == "key_themes"
