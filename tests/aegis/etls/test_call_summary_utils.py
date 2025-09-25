"""
Test coverage for aegis.etls.call_summary.main utility functions.

These tests target the simple utility functions in the ETL script
to get initial coverage of the ETL modules.
"""

import pytest
from unittest.mock import patch, MagicMock, mock_open

from aegis.etls.call_summary.main import (
    get_model_for_stage,
    get_bank_type,
    load_research_plan_config,
    load_category_extraction_config,
    load_categories_from_xlsx,
    parse_and_format_text,
    add_page_numbers,
    setup_toc_styles
)


class TestGetModelForStage:
    """Test get_model_for_stage function."""

    @patch('aegis.etls.call_summary.main.config')
    @patch('aegis.etls.call_summary.main.logger')
    def test_get_model_for_stage_with_stage_override(self, mock_logger, mock_config):
        """Test stage-level model override."""
        mock_config.llm.large.model = "default-model"
        mock_overrides = {"RESEARCH_PLAN_MODEL": "research-override-model"}

        with patch('aegis.etls.call_summary.main.MODEL_OVERRIDES', mock_overrides):
            with patch('aegis.etls.call_summary.main.CATEGORY_MODEL_OVERRIDES', {}):
                result = get_model_for_stage("RESEARCH_PLAN_MODEL")
                assert result == "research-override-model"
                mock_logger.info.assert_called_once()

    @patch('aegis.etls.call_summary.main.config')
    @patch('aegis.etls.call_summary.main.logger')
    def test_get_model_for_stage_with_category_override(self, mock_logger, mock_config):
        """Test category-specific model override (highest priority)."""
        mock_config.llm.large.model = "default-model"
        mock_category_overrides = {"revenue": "revenue-specific-model"}

        with patch('aegis.etls.call_summary.main.CATEGORY_MODEL_OVERRIDES', mock_category_overrides):
            with patch('aegis.etls.call_summary.main.MODEL_OVERRIDES', {}):
                result = get_model_for_stage("CATEGORY_EXTRACTION_MODEL", "revenue")
                assert result == "revenue-specific-model"
                mock_logger.info.assert_called_once()

    @patch('aegis.etls.call_summary.main.config')
    @patch('aegis.etls.call_summary.main.logger')
    def test_get_model_for_stage_fallback_to_default(self, mock_logger, mock_config):
        """Test fallback to default model when no overrides."""
        mock_config.llm.large.model = "default-model"

        with patch('aegis.etls.call_summary.main.MODEL_OVERRIDES', {}):
            with patch('aegis.etls.call_summary.main.CATEGORY_MODEL_OVERRIDES', {}):
                result = get_model_for_stage("UNKNOWN_STAGE")
                assert result == "default-model"
                mock_logger.debug.assert_called_once()

    @patch('aegis.etls.call_summary.main.config')
    @patch('aegis.etls.call_summary.main.logger')
    def test_get_model_for_stage_category_no_override(self, mock_logger, mock_config):
        """Test category extraction without category override falls back to stage override."""
        mock_config.llm.large.model = "default-model"
        mock_stage_overrides = {"CATEGORY_EXTRACTION_MODEL": "stage-model"}

        with patch('aegis.etls.call_summary.main.MODEL_OVERRIDES', mock_stage_overrides):
            with patch('aegis.etls.call_summary.main.CATEGORY_MODEL_OVERRIDES', {}):
                result = get_model_for_stage("CATEGORY_EXTRACTION_MODEL", "unknown_category")
                assert result == "stage-model"

    @patch('aegis.etls.call_summary.main.config')
    @patch('aegis.etls.call_summary.main.logger')
    def test_get_model_for_stage_non_category_extraction_ignores_category(self, mock_logger, mock_config):
        """Test that non-CATEGORY_EXTRACTION_MODEL stages ignore category parameter."""
        mock_config.llm.large.model = "default-model"
        mock_category_overrides = {"revenue": "revenue-model"}

        with patch('aegis.etls.call_summary.main.CATEGORY_MODEL_OVERRIDES', mock_category_overrides):
            with patch('aegis.etls.call_summary.main.MODEL_OVERRIDES', {}):
                result = get_model_for_stage("RESEARCH_PLAN_MODEL", "revenue")
                assert result == "default-model"  # Category override ignored for non-category stages


class TestGetBankType:
    """Test get_bank_type function."""

    def test_get_bank_type_canadian_banks(self):
        """Test Canadian bank IDs (1-7)."""
        assert get_bank_type(1) == "Canadian_Banks"
        assert get_bank_type(3) == "Canadian_Banks"
        assert get_bank_type(7) == "Canadian_Banks"

    def test_get_bank_type_us_banks(self):
        """Test US bank IDs (8+)."""
        assert get_bank_type(8) == "US_Banks"
        assert get_bank_type(10) == "US_Banks"
        assert get_bank_type(15) == "US_Banks"

    def test_get_bank_type_edge_cases(self):
        """Test edge cases around the boundary."""
        assert get_bank_type(7) == "Canadian_Banks"  # Last Canadian
        assert get_bank_type(8) == "US_Banks"        # First US


class TestLoadResearchPlanConfig:
    """Test load_research_plan_config function."""

    @patch("aegis.etls.call_summary.main.yaml.safe_load")
    @patch("builtins.open", new_callable=mock_open)
    @patch("aegis.etls.call_summary.main.os.path.join")
    @patch("aegis.etls.call_summary.main.os.path.dirname")
    @patch("aegis.etls.call_summary.main.os.path.abspath")
    def test_load_research_plan_config_success(self, mock_abspath, mock_dirname, mock_join, mock_file, mock_yaml):
        """Test successful loading of research plan config."""
        # Setup mocks
        mock_abspath.return_value = "/path/to/main.py"
        mock_dirname.return_value = "/path/to"
        mock_join.return_value = "/path/to/research_plan_prompt.yaml"
        mock_yaml.return_value = {
            "system_template": "Research plan template",
            "tool": {"name": "research_tool"}
        }

        result = load_research_plan_config()

        assert result["system_template"] == "Research plan template"
        assert result["tool"] == {"name": "research_tool"}
        mock_file.assert_called_once_with("/path/to/research_plan_prompt.yaml", "r")


class TestLoadCategoryExtractionConfig:
    """Test load_category_extraction_config function."""

    @patch("aegis.etls.call_summary.main.yaml.safe_load")
    @patch("builtins.open", new_callable=mock_open)
    @patch("aegis.etls.call_summary.main.os.path.join")
    @patch("aegis.etls.call_summary.main.os.path.dirname")
    @patch("aegis.etls.call_summary.main.os.path.abspath")
    def test_load_category_extraction_config_success(self, mock_abspath, mock_dirname, mock_join, mock_file, mock_yaml):
        """Test successful loading of category extraction config."""
        # Setup mocks
        mock_abspath.return_value = "/path/to/main.py"
        mock_dirname.return_value = "/path/to"
        mock_join.return_value = "/path/to/category_extraction_prompt.yaml"
        mock_yaml.return_value = {
            "system_template": "Category extraction template",
            "tool": {"name": "extraction_tool"}
        }

        result = load_category_extraction_config()

        assert result["system_template"] == "Category extraction template"
        assert result["tool"] == {"name": "extraction_tool"}
        mock_file.assert_called_once_with("/path/to/category_extraction_prompt.yaml", "r")


class TestLoadCategoriesFromXlsx:
    """Test load_categories_from_xlsx function."""

    @patch("aegis.etls.call_summary.main.pd.read_excel")
    @patch("aegis.etls.call_summary.main.os.path.exists")
    @patch("aegis.etls.call_summary.main.os.path.join")
    @patch("aegis.etls.call_summary.main.os.path.dirname")
    @patch("aegis.etls.call_summary.main.os.path.abspath")
    def test_load_categories_canadian_banks(self, mock_abspath, mock_dirname, mock_join, mock_exists, mock_read_excel):
        """Test loading Canadian bank categories."""
        # Setup mocks
        mock_abspath.return_value = "/path/to/main.py"
        mock_dirname.return_value = "/path/to"
        mock_join.return_value = "/path/to/canadian_banks_categories.xlsx"
        mock_exists.return_value = True  # File exists

        # Mock DataFrame
        mock_df = MagicMock()
        mock_df.columns = ['transcripts_section', 'category_name', 'category_description']
        mock_df.to_dict.return_value = [
            {"transcripts_section": "MD", "category_name": "Revenue", "category_description": "Revenue metrics"}
        ]
        mock_read_excel.return_value = mock_df

        result = load_categories_from_xlsx("Canadian_Banks")

        assert len(result) == 1
        assert result[0]["category_name"] == "Revenue"
        mock_read_excel.assert_called_once_with("/path/to/canadian_banks_categories.xlsx", sheet_name=0)

    @patch("aegis.etls.call_summary.main.pd.read_excel")
    @patch("aegis.etls.call_summary.main.os.path.exists")
    @patch("aegis.etls.call_summary.main.os.path.join")
    @patch("aegis.etls.call_summary.main.os.path.dirname")
    @patch("aegis.etls.call_summary.main.os.path.abspath")
    def test_load_categories_us_banks(self, mock_abspath, mock_dirname, mock_join, mock_exists, mock_read_excel):
        """Test loading US bank categories."""
        # Setup mocks
        mock_abspath.return_value = "/path/to/main.py"
        mock_dirname.return_value = "/path/to"
        mock_join.return_value = "/path/to/us_banks_categories.xlsx"
        mock_exists.return_value = True  # File exists

        # Mock DataFrame
        mock_df = MagicMock()
        mock_df.columns = ['transcripts_section', 'category_name', 'category_description']
        mock_df.to_dict.return_value = [
            {"transcripts_section": "QA", "category_name": "Credit", "category_description": "Credit quality"}
        ]
        mock_read_excel.return_value = mock_df

        result = load_categories_from_xlsx("US_Banks")

        assert len(result) == 1
        assert result[0]["category_name"] == "Credit"
        mock_read_excel.assert_called_once_with("/path/to/us_banks_categories.xlsx", sheet_name=0)

    @patch("aegis.etls.call_summary.main.pd.read_excel")
    @patch("aegis.etls.call_summary.main.os.path.exists")
    @patch("aegis.etls.call_summary.main.os.path.join")
    @patch("aegis.etls.call_summary.main.os.path.dirname")
    @patch("aegis.etls.call_summary.main.os.path.abspath")
    def test_load_categories_unknown_bank_type(self, mock_abspath, mock_dirname, mock_join, mock_exists, mock_read_excel):
        """Test unknown bank type defaults to US banks."""
        # Setup mocks
        mock_abspath.return_value = "/path/to/main.py"
        mock_dirname.return_value = "/path/to"
        mock_join.return_value = "/path/to/us_banks_categories.xlsx"
        mock_exists.return_value = True  # File exists

        # Mock DataFrame
        mock_df = MagicMock()
        mock_df.columns = ['transcripts_section', 'category_name', 'category_description']
        mock_df.to_dict.return_value = [
            {"transcripts_section": "QA", "category_name": "Default", "category_description": "Default category"}
        ]
        mock_read_excel.return_value = mock_df

        result = load_categories_from_xlsx("Unknown_Banks")  # Should default to US

        assert len(result) == 1
        assert result[0]["category_name"] == "Default"
        # Should use US banks file for unknown types
        mock_read_excel.assert_called_once_with("/path/to/us_banks_categories.xlsx", sheet_name=0)


class TestParseAndFormatText:
    """Test parse_and_format_text function."""

    def test_parse_and_format_text_plain_text(self):
        """Test plain text without any markdown formatting."""
        # Create mock paragraph
        mock_paragraph = MagicMock()
        mock_run = MagicMock()
        mock_paragraph.add_run.return_value = mock_run

        # Test plain text
        parse_and_format_text(mock_paragraph, "Simple text without formatting")

        # Should add one run with the plain text
        mock_paragraph.add_run.assert_called_once_with("Simple text without formatting")
        assert mock_run.italic is False  # Default italic setting

    def test_parse_and_format_text_with_bold(self):
        """Test text with bold markdown formatting."""
        # Create mock paragraph
        mock_paragraph = MagicMock()
        mock_runs = [MagicMock(), MagicMock(), MagicMock()]  # Need 3 runs
        mock_paragraph.add_run.side_effect = mock_runs

        # Test text with bold
        parse_and_format_text(mock_paragraph, "This has **bold text** in it")

        # Should create three runs: before, bold, after
        assert mock_paragraph.add_run.call_count == 3
        calls = mock_paragraph.add_run.call_args_list
        assert calls[0][0][0] == "This has "
        assert calls[1][0][0] == "bold text"
        assert calls[2][0][0] == " in it"

        # Second run should be bold
        assert mock_runs[1].bold is True

    def test_parse_and_format_text_with_underline(self):
        """Test text with underline markdown formatting."""
        # Create mock paragraph
        mock_paragraph = MagicMock()
        mock_runs = [MagicMock(), MagicMock(), MagicMock()]  # Need 3 runs
        mock_paragraph.add_run.side_effect = mock_runs

        # Test text with underline
        parse_and_format_text(mock_paragraph, "This has __underlined text__ in it")

        # Should create three runs: before, underlined, after
        assert mock_paragraph.add_run.call_count == 3
        calls = mock_paragraph.add_run.call_args_list
        assert calls[0][0][0] == "This has "
        assert calls[1][0][0] == "underlined text"
        assert calls[2][0][0] == " in it"

        # Second run should be underlined
        assert mock_runs[1].underline is True

    def test_parse_and_format_text_with_base_formatting(self):
        """Test text with base font size and color applied."""
        # Create mock paragraph
        mock_paragraph = MagicMock()
        mock_run = MagicMock()
        mock_paragraph.add_run.return_value = mock_run

        from docx.shared import Pt
        from docx.shared import RGBColor

        # Test with base formatting options
        parse_and_format_text(
            mock_paragraph,
            "Simple text",
            base_font_size=Pt(12),
            base_color=RGBColor(255, 0, 0),
            base_italic=True
        )

        # Should apply base formatting
        assert mock_run.font.size == Pt(12)
        assert mock_run.font.color.rgb == RGBColor(255, 0, 0)
        assert mock_run.italic is True

    def test_parse_and_format_text_mixed_formatting(self):
        """Test text with both bold and underline formatting."""
        # Create mock paragraph
        mock_paragraph = MagicMock()
        mock_runs = [MagicMock(), MagicMock(), MagicMock(), MagicMock(), MagicMock()]  # Need 5 runs
        mock_paragraph.add_run.side_effect = mock_runs

        # Test text with both bold and underline
        parse_and_format_text(mock_paragraph, "Text with **bold** and __underline__ formatting")

        # Should create five runs: before, bold, middle, underline, after
        assert mock_paragraph.add_run.call_count == 5
        calls = mock_paragraph.add_run.call_args_list

        # Verify the text content
        assert calls[0][0][0] == "Text with "
        assert calls[1][0][0] == "bold"
        assert calls[2][0][0] == " and "
        assert calls[3][0][0] == "underline"
        assert calls[4][0][0] == " formatting"

        # Verify formatting
        assert mock_runs[1].bold is True
        assert mock_runs[3].underline is True


class TestAddPageNumbers:
    """Test add_page_numbers function."""

    def test_add_page_numbers_basic(self):
        """Test basic page number addition."""
        # Create mock document with sections
        mock_doc = MagicMock()
        mock_section = MagicMock()
        mock_doc.sections = [mock_section]

        # Setup footer mock
        mock_footer = MagicMock()
        mock_section.footer = mock_footer
        mock_paragraph = MagicMock()
        mock_footer.paragraphs = [mock_paragraph]

        # Setup run mock
        mock_run = MagicMock()
        mock_paragraph.add_run.return_value = mock_run
        mock_run._element = MagicMock()

        add_page_numbers(mock_doc)

        # Should clear existing footer content
        mock_paragraph.clear.assert_called_once()

        # Should set right alignment
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        assert mock_paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT

        # Should add a run
        mock_paragraph.add_run.assert_called_once()


class TestSetupTocStyles:
    """Test setup_toc_styles function."""

    def test_setup_toc_styles_existing_styles(self):
        """Test TOC style setup when styles already exist."""
        # Create mock document with existing styles
        mock_doc = MagicMock()
        mock_styles = MagicMock()
        mock_doc.styles = mock_styles

        # Mock existing TOC styles
        mock_toc1 = MagicMock()
        mock_toc2 = MagicMock()
        mock_styles.__getitem__.side_effect = lambda key: {
            'TOC 1': mock_toc1,
            'TOC 2': mock_toc2
        }[key]

        setup_toc_styles(mock_doc)

        # Should modify existing styles
        from docx.shared import Pt
        assert mock_toc1.font.size == Pt(8)
        assert mock_toc1.font.bold is True
        assert mock_toc2.font.size == Pt(7)
        assert mock_toc2.font.bold is False

    def test_setup_toc_styles_missing_styles(self):
        """Test TOC style setup when styles don't exist."""
        # Create mock document without existing styles
        mock_doc = MagicMock()
        mock_styles = MagicMock()
        mock_doc.styles = mock_styles

        # Mock KeyError for missing styles
        def mock_getitem(key):
            raise KeyError(f"Style '{key}' not found")

        mock_styles.__getitem__.side_effect = mock_getitem

        # Mock style creation
        mock_toc1 = MagicMock()
        mock_toc2 = MagicMock()
        mock_styles.add_style.side_effect = [mock_toc1, mock_toc2]

        setup_toc_styles(mock_doc)

        # Should create new styles
        assert mock_styles.add_style.call_count == 2

        # Should configure the created styles
        from docx.shared import Pt
        assert mock_toc1.font.size == Pt(8)
        assert mock_toc1.font.bold is True
        assert mock_toc2.font.size == Pt(7)
        assert mock_toc2.font.bold is False