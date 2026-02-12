"""Tests for document_converter functions including B6.1 and B6.2."""

import pytest
from docx import Document

from aegis.etls.call_summary.document_converter import (
    auto_bold_metrics,
    parse_and_format_text,
    add_structured_content_to_doc,
    validate_document_content,
    build_toc_fallback_text,
    add_table_of_contents,
)


# ---------------------------------------------------------------------------
# auto_bold_metrics
# ---------------------------------------------------------------------------
class TestAutoBoldMetrics:
    """Tests for auto_bold_metrics()."""

    def test_bolds_percentage(self):
        result = auto_bold_metrics("Revenue grew 5.2% this quarter.")
        assert "**5.2%**" in result

    def test_bolds_dollar_with_scale(self):
        result = auto_bold_metrics("NII was $5.2 BN.")
        assert "**$5.2 BN**" in result

    def test_bolds_basis_points(self):
        result = auto_bold_metrics("Spread widened by 15 bps.")
        assert "**15 bps**" in result

    def test_does_not_double_bold(self):
        result = auto_bold_metrics("Revenue grew **5%** this quarter.")
        assert result.count("**5%**") == 1
        assert "****" not in result

    def test_empty_string(self):
        assert auto_bold_metrics("") == ""

    def test_no_metrics_unchanged(self):
        text = "The bank reported strong results."
        assert auto_bold_metrics(text) == text

    def test_negative_dollar_with_scale(self):
        result = auto_bold_metrics("Loss was -$1.2 BN.")
        assert "**-$1.2 BN**" in result

    def test_negative_dollar_without_scale(self):
        result = auto_bold_metrics("Declined by -$500.")
        assert "**-$500**" in result


# ---------------------------------------------------------------------------
# parse_and_format_text
# ---------------------------------------------------------------------------
class TestParseAndFormatText:
    """Tests for parse_and_format_text()."""

    def test_plain_text(self):
        doc = Document()
        para = doc.add_paragraph()
        parse_and_format_text(para, "Plain text here")
        assert para.text == "Plain text here"
        assert len(para.runs) == 1

    def test_bold_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        parse_and_format_text(para, "Revenue grew **5%** this quarter")
        assert para.text == "Revenue grew 5% this quarter"
        bold_runs = [r for r in para.runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "5%"

    def test_underline_formatting(self):
        doc = Document()
        para = doc.add_paragraph()
        parse_and_format_text(para, "This is __important__ text")
        assert para.text == "This is important text"
        underline_runs = [r for r in para.runs if r.underline]
        assert len(underline_runs) == 1
        assert underline_runs[0].text == "important"

    def test_mixed_bold_and_underline(self):
        doc = Document()
        para = doc.add_paragraph()
        parse_and_format_text(para, "**Bold** and __underline__")
        assert para.text == "Bold and underline"
        bold_runs = [r for r in para.runs if r.bold]
        underline_runs = [r for r in para.runs if r.underline]
        assert len(bold_runs) == 1
        assert len(underline_runs) == 1


# ---------------------------------------------------------------------------
# add_structured_content_to_doc
# ---------------------------------------------------------------------------
class TestAddStructuredContent:
    """Tests for add_structured_content_to_doc()."""

    def test_adds_heading_and_statements(self, sample_extraction_result):
        doc = Document()
        add_structured_content_to_doc(doc, sample_extraction_result)
        headings = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        assert len(headings) == 1
        assert "Revenue & Income Breakdown" in headings[0].text

    def test_rejected_category_adds_nothing(self, sample_rejected_result):
        doc = Document()
        initial_count = len(doc.paragraphs)
        add_structured_content_to_doc(doc, sample_rejected_result)
        assert len(doc.paragraphs) == initial_count

    def test_evidence_added(self, sample_extraction_result):
        doc = Document()
        add_structured_content_to_doc(doc, sample_extraction_result)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "CFO" in all_text or "CEO" in all_text


# ---------------------------------------------------------------------------
# validate_document_content (B6.1)
# ---------------------------------------------------------------------------
class TestValidateDocumentContent:
    """Tests for validate_document_content()."""

    def test_valid_document_passes(self, sample_extraction_result):
        doc = Document()
        doc.add_heading("Section", level=1)
        add_structured_content_to_doc(doc, sample_extraction_result)
        # Should not raise
        validate_document_content(doc)

    def test_empty_document_raises(self):
        doc = Document()
        with pytest.raises(ValueError, match="no paragraphs"):
            validate_document_content(doc)

    def test_missing_heading1_raises(self):
        doc = Document()
        doc.add_heading("Category", level=2)
        doc.add_paragraph("Some body text here.")
        with pytest.raises(ValueError, match="section heading"):
            validate_document_content(doc)

    def test_missing_heading2_raises(self):
        doc = Document()
        doc.add_heading("Section", level=1)
        doc.add_paragraph("Some body text here.")
        with pytest.raises(ValueError, match="category heading"):
            validate_document_content(doc)

    def test_missing_body_text_raises(self):
        doc = Document()
        doc.add_heading("Section", level=1)
        doc.add_heading("Category", level=2)
        with pytest.raises(ValueError, match="body text"):
            validate_document_content(doc)


# ---------------------------------------------------------------------------
# build_toc_fallback_text (B6.2)
# ---------------------------------------------------------------------------
class TestBuildTocFallbackText:
    """Tests for build_toc_fallback_text()."""

    def test_level1_no_indent(self):
        result = build_toc_fallback_text([("Results Summary", 1)])
        assert result == "Results Summary"

    def test_level2_indented(self):
        result = build_toc_fallback_text([("Revenue", 2)])
        assert result == "    Revenue"

    def test_mixed_levels(self):
        entries = [
            ("Results Summary", 1),
            ("Revenue", 2),
            ("Credit Quality", 2),
            ("Strategic Outlook", 1),
            ("Forward Guidance", 2),
        ]
        result = build_toc_fallback_text(entries)
        lines = result.split("\n")
        assert lines[0] == "Results Summary"
        assert lines[1] == "    Revenue"
        assert lines[2] == "    Credit Quality"
        assert lines[3] == "Strategic Outlook"
        assert lines[4] == "    Forward Guidance"

    def test_empty_entries(self):
        assert build_toc_fallback_text([]) == ""


# ---------------------------------------------------------------------------
# add_table_of_contents with toc_entries (B6.2)
# ---------------------------------------------------------------------------
class TestAddTableOfContents:
    """Tests for add_table_of_contents() with optional toc_entries."""

    def test_default_placeholder(self):
        doc = Document()
        add_table_of_contents(doc)
        # Check that the placeholder text is in the document XML
        xml = doc.element.body.xml
        assert "Table of Contents will be generated here" in xml

    def test_toc_entries_in_fallback(self):
        doc = Document()
        entries = [("Results Summary", 1), ("Revenue", 2)]
        add_table_of_contents(doc, toc_entries=entries)
        xml = doc.element.body.xml
        assert "Results Summary" in xml
        assert "Revenue" in xml
        # Placeholder should NOT be present
        assert "Table of Contents will be generated here" not in xml
