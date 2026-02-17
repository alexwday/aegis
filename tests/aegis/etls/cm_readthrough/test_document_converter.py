"""Tests for CM readthrough document converter."""

import pytest
from docx import Document

from aegis.etls.cm_readthrough.document_converter import (
    auto_bold_metrics,
    create_combined_document,
    validate_document_content,
)


def _sample_results():
    return {
        "metadata": {
            "fiscal_year": 2024,
            "quarter": "Q3",
            "subtitle_section1": "Outlook: Pipelines remain resilient despite volatility",
            "subtitle_section2": "Conference calls: Analysts focus on volatility and regulation",
            "subtitle_section3": "Conference calls: Analysts question M&A conversion and pipelines",
        },
        "outlook": {
            "Royal Bank of Canada": {
                "bank_symbol": "RY-CA",
                "statements": [
                    {
                        "category": "Investment Banking Pipelines",
                        "statement": "Pipelines are up 12% year on year.",
                        "relevance_score": 8,
                    }
                ],
            }
        },
        "section2_questions": {
            "Royal Bank of Canada": {
                "bank_symbol": "RY-CA",
                "questions": [
                    {
                        "category": "Market Volatility",
                        "verbatim_question": "How are client hedging flows trending in this volatility regime?",
                    }
                ],
            }
        },
        "section3_questions": {
            "Royal Bank of Canada": {
                "bank_symbol": "RY-CA",
                "questions": [
                    {
                        "category": "M&A Activity",
                        "verbatim_question": "What are you seeing in conversion rates for announced M&A pipelines?",
                    }
                ],
            }
        },
    }


class TestAutoBoldMetrics:
    """Tests for metric emphasis helper."""

    def test_bolds_percentage(self):
        result = auto_bold_metrics("Pipelines rose 12% year on year.")
        assert "<strong><u>12%</u></strong>" in result

    def test_does_not_double_bold_existing_markup(self):
        source = "Revenue <strong><u>$5.2 BN</u></strong> this quarter."
        result = auto_bold_metrics(source)
        assert result.count("<strong><u>$5.2 BN</u></strong>") == 1


class TestValidateDocumentContent:
    """Tests for document validation before save."""

    def test_empty_document_raises(self):
        doc = Document()
        with pytest.raises(ValueError, match="no paragraphs"):
            validate_document_content(doc)

    def test_valid_document_passes(self):
        doc = Document()
        doc.add_paragraph("Read Through For Capital Markets: Q3/24 Select U.S. & European Banks")
        doc.add_paragraph("Outlook: Subtitle")
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "Body"
        validate_document_content(doc)


def test_create_combined_document_writes_docx(tmp_path):
    """End-to-end converter call should write a DOCX file."""
    output = tmp_path / "cm_readthrough.docx"
    create_combined_document(_sample_results(), str(output))
    assert output.exists()


def test_outlook_statements_sorted_by_relevance_score(tmp_path):
    """Statements within a category should appear sorted by relevance_score descending."""
    results = _sample_results()
    results["outlook"]["Royal Bank of Canada"]["statements"] = [
        {"category": "Pipelines", "statement": "Low score.", "relevance_score": 3},
        {"category": "Pipelines", "statement": "High score.", "relevance_score": 9},
        {"category": "Pipelines", "statement": "Mid score.", "relevance_score": 6},
    ]
    output = tmp_path / "sorted.docx"
    create_combined_document(results, str(output))

    doc = Document(str(output))
    # Find the outlook table (first table after titles)
    table = doc.tables[0]
    # The content cell is row 1, col 1
    content_cell = table.rows[1].cells[1]
    paragraphs = [p.text for p in content_cell.paragraphs]
    # After the category header, quotes should be ordered: 9, 6, 3
    quote_texts = [p for p in paragraphs if "score." in p]
    assert "High score." in quote_texts[0]
    assert "Mid score." in quote_texts[1]
    assert "Low score." in quote_texts[2]


def test_relevance_score_prefix_in_document(tmp_path):
    """Each statement should have a [score/10] prefix when score > 0."""
    results = _sample_results()
    output = tmp_path / "scored.docx"
    create_combined_document(results, str(output))

    doc = Document(str(output))
    table = doc.tables[0]
    content_cell = table.rows[1].cells[1]
    # Look for score prefix in runs
    all_text = "".join(run.text for p in content_cell.paragraphs for run in p.runs)
    assert "[8/10]" in all_text


def test_grouped_categories_render_group_headers(tmp_path):
    """When statements have category_group, group headers should appear in output."""
    results = _sample_results()
    results["outlook"]["Royal Bank of Canada"]["statements"] = [
        {
            "category": "M&A Activity",
            "category_group": "Investment Banking",
            "statement": "Strong M&A pipeline.",
            "relevance_score": 9,
        },
        {
            "category": "IPO Issuance",
            "category_group": "Investment Banking",
            "statement": "IPO pipeline building.",
            "relevance_score": 7,
        },
        {
            "category": "Trading",
            "category_group": "Markets & Trading",
            "statement": "Trading volumes elevated.",
            "relevance_score": 8,
        },
    ]
    output = tmp_path / "grouped.docx"
    create_combined_document(results, str(output))

    doc = Document(str(output))
    table = doc.tables[0]
    content_cell = table.rows[1].cells[1]
    all_text = "".join(p.text for p in content_cell.paragraphs)
    assert "Investment Banking" in all_text
    assert "Markets & Trading" in all_text


def test_no_group_headers_when_groups_empty(tmp_path):
    """When no category_group values exist, no group headers should appear."""
    results = _sample_results()
    results["outlook"]["Royal Bank of Canada"]["statements"] = [
        {"category": "Cat A", "statement": "Stmt A.", "relevance_score": 7},
        {"category": "Cat B", "statement": "Stmt B.", "relevance_score": 6},
    ]
    output = tmp_path / "no_groups.docx"
    create_combined_document(results, str(output))

    doc = Document(str(output))
    table = doc.tables[0]
    content_cell = table.rows[1].cells[1]
    paragraphs = [p.text for p in content_cell.paragraphs]
    # First paragraph should be a category header, not a group header
    assert paragraphs[0].endswith(":")
