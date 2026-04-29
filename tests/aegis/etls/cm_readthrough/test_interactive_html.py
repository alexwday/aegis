"""Tests for CM readthrough editor interactive HTML helpers."""

from aegis.etls.cm_readthrough.interactive_html import build_report_state, generate_html
from aegis.etls.cm_readthrough.docx_export import create_cm_readthrough_docx_from_state


def _docx_text(path) -> str:
    """Extract visible text from a generated DOCX."""
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(part for part in parts if part)


def test_docx_export_uses_cm_bank_grouped_tables() -> None:
    """DOCX export should mirror the visible multi-bank CM report tables."""
    banks_data = {
        "RY-CA": {
            "ticker": "RY",
            "company_name": "Royal Bank of Canada",
            "md_blocks": [],
            "qa_conversations": [],
        },
        "TD-CA": {
            "ticker": "TD",
            "company_name": "Toronto-Dominion Bank",
            "md_blocks": [],
            "qa_conversations": [],
        },
    }
    categories = [
        {
            "report_section": "Outlook",
            "category_name": "Advisory Pipeline",
            "category_description": "Capital markets advisory outlook.",
        },
        {
            "report_section": "Q&A",
            "category_name": "Trading Questions",
            "category_description": "Analyst questions about trading and markets.",
        },
    ]

    state = build_report_state(
        banks_data=banks_data,
        categories=categories,
        fiscal_year=2025,
        fiscal_quarter="Q2",
        min_importance=4.0,
    )
    html = generate_html(
        state=state,
        fiscal_year=2025,
        fiscal_quarter="Q2",
        min_importance=4.0,
    )
    save_docx_body = html.split("async function saveDocx()", 1)[1].split(
        "const doc = new Document",
        1,
    )[0]

    assert "function buildCmOutlookDocxTable(" in html
    assert "function buildCmQaDocxTable(" in html
    assert "children.push(buildCmOutlookDocxTable(docxLib, sectionBuckets));" in html
    assert "children.push(buildCmQaDocxTable(docxLib, sectionBuckets));" in html
    assert "docxCmHeaderCell(docxLib, 'Banks', 12)" in html
    assert "docxCmHeaderCell(docxLib, 'Relevant Questions', 68)" in html
    assert "const groups = getReportBucketGroups(bucket.id);" in html
    assert "const scopeLabel = bankIds.length > 1" in html
    assert "computeMDSubquotes(blk, bs, bankId)" in html
    assert "getSentenceReviewStatus(sent.sid, bankId)" in html
    assert "const quotes = getReportSubquotes(bankId, bucket.id);" not in save_docx_body


def test_server_docx_export_uses_cm_initial_report_state(tmp_path) -> None:
    """Server-side DOCX should be generated from selected editor-state findings."""
    banks_data = {
        "RY-CA": {
            "ticker": "RY",
            "company_name": "Royal Bank of Canada",
            "md_blocks": [
                {
                    "sentences": [
                        {
                            "sid": "md_1",
                            "text": "Advisory pipelines improved through the quarter.",
                            "condensed": "Advisory pipelines improved through the quarter.",
                            "status": "selected",
                            "selected_bucket_id": "bucket_0",
                        }
                    ]
                }
            ],
            "qa_conversations": [
                {
                    "question_sentences": [
                        {
                            "sid": "qa_1",
                            "text": "How are clients responding to market volatility?",
                            "verbatim_text": "How are clients responding to market volatility?",
                            "status": "selected",
                            "selected_bucket_id": "bucket_1",
                        }
                    ],
                    "answer_sentences": [],
                }
            ],
        }
    }
    categories = [
        {
            "report_section": "Outlook",
            "category_name": "Advisory Pipeline",
            "category_description": "Capital markets advisory outlook.",
        },
        {
            "report_section": "Q&A",
            "category_name": "Trading Questions",
            "category_description": "Analyst questions about trading and markets.",
        },
    ]
    state = build_report_state(
        banks_data=banks_data,
        categories=categories,
        fiscal_year=2025,
        fiscal_quarter="Q2",
        min_importance=4.0,
    )

    output_path = tmp_path / "CM_Readthrough_2025_Q2.docx"
    create_cm_readthrough_docx_from_state(report_state=state, output_path=str(output_path))

    assert output_path.exists()
    assert output_path.stat().st_size > 0


def test_server_docx_export_matches_cm_preview_grouping_and_verbatim_text(tmp_path) -> None:
    """Server-side DOCX should mirror the report preview, not raw sentence rows."""
    banks_data = {
        "RY-CA": {
            "ticker": "RY-CA",
            "symbol": "RY",
            "company_name": "Royal Bank of Canada",
            "md_blocks": [
                {
                    "id": "md_block_1",
                    "speaker": "Chief Financial Officer",
                    "speaker_title": "CFO",
                    "sentences": [
                        {
                            "sid": "md_1",
                            "text": "Capital markets pipelines improved in the quarter.",
                            "condensed": "Condensed pipeline summary.",
                            "status": "selected",
                            "selected_bucket_id": "bucket_0",
                            "importance_score": 8.0,
                            "scores": {"bucket_0": 8.0},
                        },
                        {
                            "sid": "md_2",
                            "text": "Management expects advisory activity to remain constructive.",
                            "condensed": "Condensed advisory summary.",
                            "status": "selected",
                            "selected_bucket_id": "bucket_0",
                            "importance_score": 7.0,
                            "scores": {"bucket_0": 7.0},
                        },
                    ],
                }
            ],
            "qa_conversations": [
                {
                    "id": "qa_conv_1",
                    "analyst_name": "Analyst",
                    "analyst_affiliation": "Example Securities",
                    "question_sentences": [
                        {
                            "sid": "qa_1",
                            "text": "Can you discuss trading momentum?",
                            "status": "selected",
                            "selected_bucket_id": "bucket_1",
                            "importance_score": 7.5,
                            "scores": {"bucket_1": 7.5},
                        },
                        {
                            "sid": "qa_2",
                            "text": "And how sustainable is the client activity?",
                            "status": "selected",
                            "selected_bucket_id": "bucket_1",
                            "importance_score": 7.0,
                            "scores": {"bucket_1": 7.0},
                        },
                    ],
                    "answer_sentences": [],
                }
            ],
        }
    }
    categories = [
        {
            "report_section": "Outlook",
            "category_name": "Advisory Pipeline",
            "category_description": "Capital markets advisory outlook.",
        },
        {
            "report_section": "Q&A",
            "category_name": "Trading Questions",
            "category_description": "Analyst questions about trading and markets.",
        },
    ]
    state = build_report_state(
        banks_data=banks_data,
        categories=categories,
        fiscal_year=2025,
        fiscal_quarter="Q2",
        min_importance=4.0,
    )

    output_path = tmp_path / "CM_Readthrough_2025_Q2.docx"
    create_cm_readthrough_docx_from_state(report_state=state, output_path=str(output_path))
    text = _docx_text(output_path)

    assert "RY-CA" not in text
    assert "RY" in text
    assert (
        "Capital markets pipelines improved in the quarter. "
        "Management expects advisory activity to remain constructive."
    ) in text
    assert "Condensed pipeline summary." not in text
    assert "Can you discuss trading momentum? And how sustainable is the client activity?" in text
