"""
Quarterly Newsletter ETL - Multi-Bank Summary Generator

This ETL processes all monitored financial institutions for a given quarter and generates
a combined newsletter document containing paragraph summaries for each bank's earnings call.

Key architecture:
- Processes multiple banks in sequence rather than individual bank selection
- Generates simplified paragraph summaries rather than detailed categorical analysis
- Combines all results into a single consolidated Word document
- Uses existing transcript subagent infrastructure for data retrieval

Usage:
    python -m aegis.etls.quarterly_newsletter.main --year 2024 --quarter Q3
"""

import argparse
import asyncio
import json
import sys
import uuid
import os
import yaml
import time
from datetime import datetime
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
from sqlalchemy import text
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import hashlib

# Import existing Aegis infrastructure
from aegis.model.subagents.transcripts.retrieval import retrieve_full_section
from aegis.model.subagents.transcripts.formatting import format_full_section_chunks
from aegis.utils.ssl import setup_ssl
from aegis.connections.oauth_connector import setup_authentication
from aegis.connections.llm_connector import complete
from aegis.connections.postgres_connector import get_connection
from aegis.utils.logging import setup_logging, get_logger
from aegis.utils.settings import config

from aegis.etls.quarterly_newsletter.config.config import MODELS, TEMPERATURE, MAX_TOKENS

# Import document converter functions
from aegis.etls.quarterly_newsletter.document_converter import (
    convert_docx_to_pdf,
    get_standard_report_metadata,
    bank_summaries_to_markdown
)

setup_logging()
logger = get_logger()


@dataclass
class BankSummary:
    """
    Data structure containing processing results for a single bank.

    Tracks both successful summaries and processing failures to enable
    comprehensive reporting and partial success handling.
    """
    bank_id: int
    bank_name: str
    bank_symbol: str
    fiscal_year: int
    quarter: str
    summary_paragraph: str
    bank_type: str  # "Canadian_Banks" or "US_Banks"
    processing_success: bool = True
    error_message: Optional[str] = None


def load_newsletter_prompt() -> str:
    """
    Load the newsletter summary generation prompt from YAML configuration.

    Returns the system prompt template for LLM summary generation.
    """
    prompt_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'prompts', 'newsletter_summary_prompt.yaml'
    )

    with open(prompt_path, 'r') as f:
        config = yaml.safe_load(f)

    return config['system_template']


def load_monitored_institutions() -> Dict[str, Dict]:
    """
    Load the complete list of monitored financial institutions from YAML configuration.

    Returns dictionary mapping bank symbols to institution metadata including
    ID, name, and bank type classification.
    """
    config_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'config', 'monitored_institutions.yaml'
    )

    with open(config_path, 'r') as f:
        institutions = yaml.safe_load(f)

    logger.info(f"Loaded {len(institutions)} monitored institutions for processing")
    return institutions


async def check_data_availability(bank_id: int, fiscal_year: int, quarter: str) -> bool:
    """
    Verify transcript data availability for specified bank and reporting period.

    Queries aegis_data_availability table to confirm 'transcripts' database
    is available before attempting processing.

    Args:
        bank_id: Unique bank identifier
        fiscal_year: Reporting year
        quarter: Reporting quarter (Q1-Q4)

    Returns:
        True if transcript data exists, False otherwise
    """
    async with get_connection() as conn:
        result = await conn.execute(text(
            """
            SELECT database_names
            FROM aegis_data_availability
            WHERE bank_id = :bank_id
              AND fiscal_year = :fiscal_year
              AND quarter = :quarter
            """
        ), {
            "bank_id": bank_id,
            "fiscal_year": fiscal_year,
            "quarter": quarter
        })

        row = result.fetchone()
        if row and row[0]:
            return 'transcripts' in row[0]

        return False


async def generate_bank_summary(
    bank_id: int,
    bank_name: str,
    bank_symbol: str,
    bank_type: str,
    fiscal_year: int,
    quarter: str,
    context: Dict[str, Any]
) -> str:
    """
    Generate a paragraph summary for a single bank's earnings call.

    Process flow:
    1. Create bank/period combination object for transcript retrieval
    2. Retrieve complete transcript using existing subagent functions
    3. Format transcript content for LLM processing
    4. Generate paragraph summary using configured prompt
    5. Return formatted summary text

    Args:
        bank_id: Unique bank identifier
        bank_name: Full legal bank name
        bank_symbol: Trading symbol
        bank_type: Bank classification (Canadian_Banks/US_Banks)
        fiscal_year: Reporting year
        quarter: Reporting quarter
        context: Execution context containing auth and SSL configuration

    Returns:
        Generated paragraph summary text

    Raises:
        ValueError: If no transcript data found
        Exception: If LLM processing fails
    """
    logger.info(
        "quarterly_newsletter.processing_bank",
        bank_id=bank_id,
        bank_name=bank_name
    )

    # Create combination object expected by transcript subagent functions
    combo = {
        "bank_id": bank_id,
        "bank_name": bank_name,
        "bank_symbol": bank_symbol,
        "fiscal_year": fiscal_year,
        "quarter": quarter,
        "query_intent": f"Generate newsletter summary for {bank_name} {quarter} {fiscal_year}"
    }

    # Retrieve complete transcript using existing infrastructure
    # This leverages all existing transcript processing logic
    chunks = await retrieve_full_section(
        combo=combo,
        sections="ALL",  # Include both management discussion and Q&A sections
        context=context
    )

    if not chunks:
        raise ValueError(f"No transcript data found for {bank_name} {quarter} {fiscal_year}")

    # Format transcript chunks into coherent text
    formatted_transcript = await format_full_section_chunks(
        chunks=chunks,
        combo=combo,
        context=context
    )

    if not formatted_transcript:
        raise ValueError(f"No formatted transcript content for {bank_name} {quarter} {fiscal_year}")

    logger.info(
        "quarterly_newsletter.transcript_retrieved",
        bank_id=bank_id,
        content_length=len(formatted_transcript)
    )

    # Load prompt template and format with bank-specific context
    prompt_template = load_newsletter_prompt()
    system_prompt = prompt_template.format(
        bank_name=bank_name,
        bank_symbol=bank_symbol,
        quarter=quarter,
        fiscal_year=fiscal_year
    )

    user_prompt = f"Create a newsletter summary paragraph for this earnings call transcript:\n\n{formatted_transcript}"

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    # Generate summary using configured LLM parameters with retry logic
    max_retries = 3
    response = None

    for attempt in range(max_retries):
        try:
            response = await complete(
                messages=messages,
                context=context,
                llm_params={
                    "model": MODELS["summary"],
                    "temperature": TEMPERATURE,
                    "max_tokens": MAX_TOKENS
                }
            )

            if response:
                break  # Success

        except Exception as e:
            if attempt < max_retries - 1:
                logger.warning(
                    f"Retry {attempt + 1}/{max_retries} for {bank_name} summary: {str(e)}"
                )
                await asyncio.sleep(2 ** attempt)  # Exponential backoff: 1s, 2s, 4s
                continue
            else:
                error_msg = f"Failed after {max_retries} attempts: {str(e)}"
                logger.error(
                    "quarterly_newsletter.summary_failed_all_retries",
                    bank_name=bank_name,
                    error=error_msg
                )
                raise ValueError(error_msg)

    if not response:
        raise ValueError(f"LLM returned no response for {bank_name} after {max_retries} attempts")

    # Handle different response formats
    if isinstance(response, dict) and 'choices' in response:
        summary = response['choices'][0]['message']['content'].strip()
    else:
        # Response might be a string directly
        summary = str(response).strip()

    logger.info(
        "quarterly_newsletter.bank_summary_complete",
        bank_id=bank_id,
        summary_length=len(summary)
    )

    return summary


async def process_all_banks(fiscal_year: int, quarter: str) -> List[BankSummary]:
    """
    Main processing orchestrator for multi-bank summary generation.

    Coordinates the complete workflow:
    1. Authentication and context setup
    2. Institution list loading
    3. Sequential processing of each bank
    4. Error collection and success tracking
    5. Results aggregation

    Uses sequential rather than parallel processing to avoid API rate limits
    and simplify error handling. Each bank failure is isolated and does not
    affect processing of remaining institutions.

    Args:
        fiscal_year: Target reporting year
        quarter: Target reporting quarter

    Returns:
        List of BankSummary objects containing results and errors
    """
    execution_id = str(uuid.uuid4())
    logger.info(
        "quarterly_newsletter.processing_started",
        execution_id=execution_id,
        fiscal_year=fiscal_year,
        quarter=quarter
    )

    # Setup authentication and SSL configuration for LLM calls
    ssl_config = setup_ssl()
    auth_config = await setup_authentication(execution_id, ssl_config)

    if not auth_config["success"]:
        raise RuntimeError(f"Authentication failed: {auth_config.get('error', 'Unknown error')}")

    context = {
        "execution_id": execution_id,
        "auth_config": auth_config,
        "ssl_config": ssl_config
    }

    # Load complete institution list from configuration
    institutions = load_monitored_institutions()
    summaries = []

    # Process each institution sequentially
    for symbol, bank_info in institutions.items():
        bank_id = bank_info['id']
        bank_name = bank_info['name']
        bank_type = bank_info['type']

        try:
            # Verify data availability before processing
            if not await check_data_availability(bank_id, fiscal_year, quarter):
                logger.info(
                    "quarterly_newsletter.bank_skipped",
                    bank_id=bank_id,
                    bank_name=bank_name,
                    reason="No transcript data available"
                )

                # Record unavailable data for reporting
                summaries.append(BankSummary(
                    bank_id=bank_id,
                    bank_name=bank_name,
                    bank_symbol=symbol,
                    fiscal_year=fiscal_year,
                    quarter=quarter,
                    summary_paragraph="",
                    bank_type=bank_type,
                    processing_success=False,
                    error_message=f"No transcript data available for {quarter} {fiscal_year}"
                ))
                continue

            # Generate summary for available data
            summary_text = await generate_bank_summary(
                bank_id=bank_id,
                bank_name=bank_name,
                bank_symbol=symbol,
                bank_type=bank_type,
                fiscal_year=fiscal_year,
                quarter=quarter,
                context=context
            )

            # Record successful processing
            summaries.append(BankSummary(
                bank_id=bank_id,
                bank_name=bank_name,
                bank_symbol=symbol,
                fiscal_year=fiscal_year,
                quarter=quarter,
                summary_paragraph=summary_text,
                bank_type=bank_type,
                processing_success=True
            ))

        except Exception as e:
            # Log error but continue processing remaining banks
            logger.error(
                "quarterly_newsletter.bank_failed",
                bank_id=bank_id,
                bank_name=bank_name,
                error=str(e)
            )

            # Record processing failure
            summaries.append(BankSummary(
                bank_id=bank_id,
                bank_name=bank_name,
                bank_symbol=symbol,
                fiscal_year=fiscal_year,
                quarter=quarter,
                summary_paragraph="",
                bank_type=bank_type,
                processing_success=False,
                error_message=str(e)
            ))

    # Log final processing statistics
    successful = [s for s in summaries if s.processing_success]
    failed = [s for s in summaries if not s.processing_success]

    logger.info(
        "quarterly_newsletter.processing_complete",
        execution_id=execution_id,
        total_banks=len(summaries),
        successful_count=len(successful),
        failed_count=len(failed)
    )

    return summaries


def add_page_numbers(doc):
    """Add page numbers to the footer of the document."""
    for section in doc.sections:
        footer = section.footer

        # Clear existing footer content
        footer.paragraphs[0].clear()

        # Create a paragraph for the page number
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add page number field
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run._element.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)


async def create_newsletter_document(summaries: List[BankSummary], fiscal_year: int, quarter: str) -> Dict[str, Any]:
    """
    Generate Word document containing all bank summaries in newsletter format.

    Document structure:
    - Title page with quarter/year and generation date
    - Canadian Banks section with individual bank summaries
    - US Banks section with individual bank summaries
    - Processing notes section (if any failures occurred)

    Each bank summary appears as a level-2 heading followed by the paragraph content.
    Failed processing attempts are documented in the processing notes section.

    Args:
        summaries: List of BankSummary objects from processing
        fiscal_year: Reporting year for document title
        quarter: Reporting quarter for document title

    Returns:
        Dictionary containing document paths and metadata
    """
    logger.info(
        "quarterly_newsletter.creating_document",
        total_summaries=len(summaries)
    )

    doc = Document()

    # Configure document margins (standardize with call_summary narrow margins)
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.5)
        section.gutter = Inches(0)

    # Add page numbers to footer
    add_page_numbers(doc)

    # Check for banner image in config folder
    etl_dir = os.path.dirname(os.path.abspath(__file__))
    config_dir = os.path.join(etl_dir, 'config')
    banner_path = None

    for ext in ['jpg', 'jpeg', 'png']:
        potential_banner = os.path.join(config_dir, f'banner.{ext}')
        if os.path.exists(potential_banner):
            banner_path = potential_banner
            break

    # If not found, check in call_summary config directory (fallback)
    if not banner_path:
        call_summary_config_dir = os.path.join(os.path.dirname(os.path.dirname(etl_dir)), 'call_summary', 'config')
        for ext in ['jpg', 'jpeg', 'png']:
            potential_banner = os.path.join(call_summary_config_dir, f'banner.{ext}')
            if os.path.exists(potential_banner):
                banner_path = potential_banner
                break

    # Add banner image if found
    if banner_path:
        try:
            # Add the banner image at the top, adjusted for narrow margins
            doc.add_picture(banner_path, width=Inches(7.4))  # Full width with narrow margins
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(3)

            logger.info(f"Banner added from: {banner_path}")
        except Exception as e:
            logger.warning(f"Could not add banner: {str(e)}")

    # Document title and metadata
    title_text = f"Quarterly Banking Newsletter - {quarter} {fiscal_year}"
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].italic = True

    doc.add_page_break()

    # Separate successful and failed summaries for document organization
    successful_summaries = [s for s in summaries if s.processing_success]
    failed_summaries = [s for s in summaries if not s.processing_success]

    # Group successful summaries by bank type for sectioned presentation
    canadian_banks = [s for s in successful_summaries if s.bank_type == "Canadian_Banks"]
    us_banks = [s for s in successful_summaries if s.bank_type == "US_Banks"]

    # Canadian Banks section
    if canadian_banks:
        doc.add_heading("Canadian Banks", level=1)

        for summary in sorted(canadian_banks, key=lambda x: x.bank_name):
            bank_heading = doc.add_heading(f"{summary.bank_name} ({summary.bank_symbol})", level=2)
            para = doc.add_paragraph(summary.summary_paragraph)
            para.paragraph_format.space_after = Pt(12)

    # US Banks section
    if us_banks:
        doc.add_heading("US Banks", level=1)

        for summary in sorted(us_banks, key=lambda x: x.bank_name):
            bank_heading = doc.add_heading(f"{summary.bank_name} ({summary.bank_symbol})", level=2)
            para = doc.add_paragraph(summary.summary_paragraph)
            para.paragraph_format.space_after = Pt(12)

    # Processing notes section for failed attempts
    if failed_summaries:
        doc.add_page_break()
        doc.add_heading("Processing Notes", level=1)

        intro_para = doc.add_paragraph(
            f"This newsletter includes summaries for {len(successful_summaries)} banks. "
            f"The following {len(failed_summaries)} banks could not be processed:"
        )

        for failed in failed_summaries:
            failure_para = doc.add_paragraph(
                f"• {failed.bank_name}: {failed.error_message}",
                style='List Bullet'
            )

    # Save document with timestamp-based filename
    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')
    os.makedirs(output_dir, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"Quarterly_Newsletter_{quarter}_{fiscal_year}_{timestamp}"
    docx_filename = f"{base_filename}.docx"
    filepath = os.path.join(output_dir, docx_filename)

    doc.save(filepath)

    logger.info(
        "quarterly_newsletter.document_saved",
        filepath=filepath,
        successful_banks=len(successful_summaries),
        failed_banks=len(failed_summaries)
    )

    # Generate PDF
    pdf_path = None
    pdf_filename = None
    pdf_path_target = os.path.join(output_dir, f"{base_filename}.pdf")
    logger.info(f"Generating PDF: {pdf_path_target}")
    pdf_result = convert_docx_to_pdf(filepath, pdf_path_target)
    if pdf_result:
        pdf_path = pdf_result
        pdf_filename = f"{base_filename}.pdf"
        logger.info(f"PDF generated successfully: {pdf_path}")
    else:
        logger.warning("PDF generation failed")

    # Generate Markdown content for database storage
    logger.info("Generating markdown for database storage")
    markdown_content = bank_summaries_to_markdown(summaries, quarter, fiscal_year)

    # Save to database
    logger.info("Saving report to aegis_reports table")
    report_metadata = get_standard_report_metadata()
    generation_timestamp = datetime.now()
    execution_id = str(uuid.uuid4())

    database_saved = False
    report_id = None

    try:
        async with get_connection() as conn:
            # Since this is a multi-bank report, we use a special bank_id of 0
            # to indicate it's a consolidated report
            bank_id = 0  # Special ID for multi-bank reports (0 = all banks)
            bank_name = "All Monitored Banks"
            bank_symbol = "ALL"

            # Delete any existing quarterly newsletter for this period
            delete_result = await conn.execute(text(
                """
                DELETE FROM aegis_reports
                WHERE bank_id = :bank_id
                  AND fiscal_year = :fiscal_year
                  AND quarter = :quarter
                  AND report_type = :report_type
                RETURNING id
                """
            ), {
                "bank_id": bank_id,
                "fiscal_year": fiscal_year,
                "quarter": quarter,
                "report_type": report_metadata["report_type"]
            })
            deleted = delete_result.fetchall()

            if deleted:
                logger.info(f"Deleted {len(deleted)} existing quarterly newsletter(s)")

            # Insert new report
            result = await conn.execute(text(
                """
                INSERT INTO aegis_reports (
                report_name,
                report_description,
                report_type,
                bank_id,
                bank_name,
                bank_symbol,
                fiscal_year,
                quarter,
                local_filepath,
                s3_document_name,
                s3_pdf_name,
                markdown_content,
                generation_date,
                generated_by,
                execution_id,
                metadata
            ) VALUES (
                :report_name,
                :report_description,
                :report_type,
                :bank_id,
                :bank_name,
                :bank_symbol,
                :fiscal_year,
                :quarter,
                :local_filepath,
                :s3_document_name,
                :s3_pdf_name,
                :markdown_content,
                :generation_date,
                :generated_by,
                :execution_id,
                :metadata
                )
                RETURNING id
                """
            ), {
                "report_name": report_metadata["report_name"],
                "report_description": report_metadata["report_description"],
                "report_type": report_metadata["report_type"],
                "bank_id": bank_id,
                "bank_name": bank_name,
                "bank_symbol": bank_symbol,
                "fiscal_year": fiscal_year,
                "quarter": quarter,
                "local_filepath": filepath,
                "s3_document_name": docx_filename,
                "s3_pdf_name": pdf_filename,
                "markdown_content": markdown_content,
                "generation_date": generation_timestamp,
                "generated_by": "quarterly_newsletter_etl",
                "execution_id": execution_id,
                "metadata": json.dumps({
                    "total_banks": len(summaries),
                    "successful_banks": len(successful_summaries),
                    "failed_banks": len(failed_summaries),
                    "canadian_banks": len(canadian_banks),
                    "us_banks": len(us_banks)
                })
            })
            report_row = result.fetchone()
            await conn.commit()
            report_id = report_row.id
            database_saved = True
            logger.info(f"Report saved to database with ID: {report_id}")

    except Exception as e:
        logger.error(f"Database error: {str(e)}")
        # Continue even if database save fails

    # Return comprehensive metadata
    return {
        "filepath": filepath,
        "pdf_path": pdf_path,
        "docx_filename": docx_filename,
        "pdf_filename": pdf_filename,
        "total_banks": len(summaries),
        "successful_banks": len(successful_summaries),
        "failed_banks": len(failed_summaries),
        "execution_id": execution_id,
        "database_saved": database_saved,
        "report_id": report_id
    }


async def generate_quarterly_newsletter(fiscal_year: int, quarter: str) -> str:
    """
    Main entry point orchestrating complete newsletter generation workflow.

    Coordinates bank processing and document generation, returning formatted
    status report with processing results and output file location.

    Args:
        fiscal_year: Target reporting year
        quarter: Target reporting quarter

    Returns:
        Formatted status report string
    """
    try:
        logger.info(
            "quarterly_newsletter.started",
            fiscal_year=fiscal_year,
            quarter=quarter
        )

        # Execute multi-bank processing
        summaries = await process_all_banks(fiscal_year, quarter)

        # Generate consolidated document with PDF and database storage
        document_result = await create_newsletter_document(summaries, fiscal_year, quarter)
        document_path = document_result["filepath"]
        pdf_path = document_result.get("pdf_path")

        # Compile processing statistics
        successful = [s for s in summaries if s.processing_success]
        failed = [s for s in summaries if not s.processing_success]

        report = f"""
================================================================================
QUARTERLY NEWSLETTER GENERATION COMPLETE
================================================================================
Period: {quarter} {fiscal_year}
Generated: {datetime.now().isoformat()}
Execution ID: {document_result['execution_id']}
================================================================================

DOCUMENT OUTPUTS:
- Word Document: {document_path}
- PDF Document: {pdf_path if pdf_path else 'PDF generation failed'}
- Database Entry: {'Saved to aegis_reports table' if document_result['database_saved'] else 'Not saved'}

PROCESSING RESULTS:
- Total Banks: {len(summaries)}
- Successfully Processed: {len(successful)}
- Failed to Process: {len(failed)}
"""

        if successful:
            report += "\nSUCCESSFUL BANKS:\n"
            for s in successful:
                report += f"  ✓ {s.bank_name} ({s.bank_symbol})\n"

        if failed:
            report += "\nFAILED BANKS:\n"
            for s in failed:
                report += f"  ✗ {s.bank_name} ({s.bank_symbol}): {s.error_message}\n"

        report += f"""
================================================================================
Newsletter saved to: {document_path}
================================================================================
"""

        logger.info(
            "quarterly_newsletter.completed",
            fiscal_year=fiscal_year,
            quarter=quarter,
            document_path=document_path
        )

        return report

    except Exception as e:
        error_msg = f"Error generating quarterly newsletter: {str(e)}"
        logger.error(
            "quarterly_newsletter.error",
            fiscal_year=fiscal_year,
            quarter=quarter,
            error=error_msg,
            exc_info=True
        )
        return f"❌ {error_msg}"


def main():
    """Command-line interface for quarterly newsletter generation."""
    parser = argparse.ArgumentParser(
        description="Generate quarterly newsletter with summaries from all monitored banks",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python -m aegis.etls.quarterly_newsletter.main --year 2024 --quarter Q3
  python -m aegis.etls.quarterly_newsletter.main --year 2025 --quarter Q1
        """
    )

    parser.add_argument(
        "--year",
        type=int,
        required=True,
        help="Fiscal year (e.g., 2024)"
    )

    parser.add_argument(
        "--quarter",
        required=True,
        choices=["Q1", "Q2", "Q3", "Q4"],
        help="Quarter (Q1, Q2, Q3, Q4)"
    )

    args = parser.parse_args()

    print(f"\n📰 Generating quarterly newsletter for {args.quarter} {args.year}...")

    result = asyncio.run(generate_quarterly_newsletter(
        fiscal_year=args.year,
        quarter=args.quarter
    ))

    print(result)


if __name__ == "__main__":
    main()