"""
Document converter utilities for WM Readthrough ETL - MINIMAL VERSION.
Quick and simple document generation.
"""

import os
import subprocess
import platform
from typing import Dict, Any
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from aegis.utils.logging import get_logger

logger = get_logger()


def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """Convert DOCX to PDF using available tools."""
    system = platform.system()

    try:
        if system == "Darwin":  # macOS
            # Try LibreOffice
            try:
                result = subprocess.run(
                    ["soffice", "--headless", "--convert-to", "pdf", "--outdir",
                     os.path.dirname(pdf_path), docx_path],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0:
                    generated_pdf = os.path.join(
                        os.path.dirname(pdf_path),
                        Path(docx_path).stem + ".pdf"
                    )
                    if generated_pdf != pdf_path:
                        os.rename(generated_pdf, pdf_path)
                    logger.info(f"PDF created: {pdf_path}")
                    return True
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass

    except Exception as e:
        logger.error(f"Error converting to PDF: {e}")

    logger.warning(f"PDF conversion failed for {docx_path}")
    return False


def create_combined_document(results: Dict[str, Any], output_path: str) -> None:
    """Create a simple Word document with all WM readthrough results."""
    doc = Document()
    metadata = results["metadata"]

    # Title
    title = doc.add_heading("WM Readthrough Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"{metadata['fiscal_year']} {metadata['quarter']}")
    run.font.size = Pt(16)

    doc.add_paragraph()

    # Summary
    summary = doc.add_paragraph()
    summary.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}\n")
    summary.add_run(f"Banks with Page 1: {metadata['banks_with_page1']}\n")
    summary.add_run(f"Banks with Page 2: {metadata['banks_with_page2']}\n")
    summary.add_run(f"Banks with Page 3: {metadata['banks_with_page3']}\n")
    summary.add_run(f"Banks with Page 4: {metadata['banks_with_page4']}\n")
    summary.add_run(f"Banks with Page 5: {metadata['banks_with_page5']}\n")

    # PAGE 1: WM Narratives
    if results.get("page1_results"):
        doc.add_page_break()
        doc.add_heading("Page 1: Wealth Management Narratives", 1)

        for bank_name, data in results["page1_results"].items():
            doc.add_heading(f"{bank_name} ({data.get('bank_symbol', '')})", 2)

            # Key theme
            theme = data.get("key_theme", "")
            if theme:
                p = doc.add_paragraph()
                p.add_run("Key Theme: ").bold = True
                p.add_run(theme)

            # Narrative
            narrative = data.get("narrative_summary", "")
            if narrative:
                doc.add_paragraph(narrative)

            # Supporting quotes
            quotes = data.get("supporting_quotes", [])
            if quotes:
                doc.add_paragraph().add_run("Supporting Quotes:").bold = True
                for quote in quotes:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'"{quote.get("quote", "")}" - {quote.get("speaker", "")}')

    # PAGE 2: Three-Theme Q&A
    if results.get("page2_results"):
        doc.add_page_break()
        doc.add_heading("Page 2: Three-Theme Q&A", 1)

        for bank_name, data in results["page2_results"].items():
            doc.add_heading(f"{bank_name} ({data.get('bank_symbol', '')})", 2)

            questions = data.get("questions", [])
            for q in questions:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"[{q.get('theme', '')}] ").bold = True
                p.add_run(q.get('verbatim_question', ''))

    # PAGE 3: Canadian AM
    if results.get("page3_results"):
        doc.add_page_break()
        doc.add_heading("Page 3: Canadian Asset Management", 1)

        for bank_name, data in results["page3_results"].items():
            doc.add_heading(f"{bank_name} ({data.get('bank_symbol', '')})", 2)

            # AUM/Net Flows
            aum = data.get("aum_netflows", {})
            if aum:
                doc.add_paragraph().add_run("AUM & Net Flows:").bold = True
                doc.add_paragraph(f"Total AUM: {aum.get('total_aum', 'N/A')}")
                doc.add_paragraph(f"AUM Breakdown: {aum.get('aum_breakdown', 'N/A')}")
                doc.add_paragraph(f"Net Flows: {aum.get('net_flows', 'N/A')}")
                doc.add_paragraph(f"Notable Metrics: {aum.get('notable_metrics', 'N/A')}")

            # Focus areas
            focus_areas = data.get("focus_areas", [])
            if focus_areas:
                doc.add_paragraph().add_run("Focus Areas:").bold = True
                for area in focus_areas:
                    doc.add_heading(area.get("theme_title", ""), 3)
                    for q in area.get("questions", []):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(q.get('verbatim_question', ''))

    # PAGE 4: Three-Column Table
    if results.get("page4_results"):
        doc.add_page_break()
        doc.add_heading("Page 4: Banking Metrics (NII/NIM, Credit/PCL, Tariffs)", 1)

        for bank_name, data in results["page4_results"].items():
            doc.add_heading(f"{bank_name} ({data.get('bank_symbol', '')})", 2)
            doc.add_paragraph(f"NII/NIM: {data.get('nii_nim_summary', 'N/A')}")
            doc.add_paragraph(f"Credit/PCL: {data.get('credit_pcl_summary', 'N/A')}")
            doc.add_paragraph(f"Tariffs/Uncertainty: {data.get('tariff_uncertainty_summary', 'N/A')}")

    # PAGE 5: Six-Theme Q&A
    if results.get("page5_results"):
        doc.add_page_break()
        doc.add_heading("Page 5: Six-Theme Q&A", 1)

        for bank_name, data in results["page5_results"].items():
            doc.add_heading(f"{bank_name} ({data.get('bank_symbol', '')})", 2)

            questions = data.get("questions", [])
            for q in questions:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"[{q.get('theme', '')}] ").bold = True
                p.add_run(q.get('verbatim_question', ''))

    # Save
    doc.save(output_path)
    logger.info(f"Document saved to {output_path}")


def structured_data_to_markdown(results: Dict[str, Any]) -> str:
    """Convert results to markdown for database storage."""
    lines = [
        "# WM Readthrough Report",
        f"## {results['metadata']['fiscal_year']} {results['metadata']['quarter']}",
        "",
        "## Summary",
        f"- Banks with Page 1: {results['metadata']['banks_with_page1']}",
        f"- Banks with Page 2: {results['metadata']['banks_with_page2']}",
        f"- Banks with Page 3: {results['metadata']['banks_with_page3']}",
        f"- Banks with Page 4: {results['metadata']['banks_with_page4']}",
        f"- Banks with Page 5: {results['metadata']['banks_with_page5']}",
        ""
    ]

    # Add page summaries
    for page_num in range(1, 6):
        page_key = f"page{page_num}_results"
        if results.get(page_key):
            lines.append(f"## Page {page_num}")
            lines.append(f"Banks: {len(results[page_key])}")
            lines.append("")

    return "\n".join(lines)
