"""
Document converter utilities for key themes ETL.

This module provides helper functions for creating Word documents from themed Q&A data
and utilities for markdown processing, document validation, and HTML metric auto-bolding.
"""

import os
import re
from typing import Dict
from html.parser import HTMLParser
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from aegis.utils.logging import get_logger

logger = get_logger()

# Patterns for financial metrics that should be bolded (most specific first)
_METRIC_PATTERNS = [
    r"-?\$[\d,]+(?:\.\d+)?\s*(?:MM|BN|TN|K|M|B)\b",  # Dollar with scale: $1.2 BN
    r"-?\$[\d,]+(?:\.\d+)?(?!\s*(?:MM|BN|TN|K|M|B))\b",  # Dollar without scale: $1,200
    r"\d+(?:\.\d+)?\s*bps\b",  # Basis points: 15 bps
    r"\d+(?:\.\d+)?%",  # Percentages: 12.3%
]


def validate_document_content(doc) -> None:
    """
    Validate generated document has meaningful content before saving.

    Checks that:
    1. Document has at least one theme header
    2. Every theme header has at least one body text paragraph following it
    3. Document is not empty

    Args:
        doc: Word Document object to validate

    Raises:
        ValueError: If document fails any validation check
    """
    if not doc.paragraphs:
        raise ValueError("Document has no paragraphs")

    theme_count = 0
    empty_themes = []
    current_theme = None
    current_theme_has_body = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check for theme headers (created by add_theme_header_with_background)
        is_theme_header = False
        if text.startswith("Theme ") and para.runs:
            for run in para.runs:
                if run.font.bold:
                    is_theme_header = True
                    break

        if is_theme_header:
            # Finalize previous theme
            if current_theme is not None and not current_theme_has_body:
                empty_themes.append(current_theme)

            current_theme = text
            current_theme_has_body = False
            theme_count += 1

        # Check for body text (non-header paragraphs with content)
        elif current_theme is not None and len(text) > 20 and not text.startswith("_"):
            current_theme_has_body = True

    # Finalize last theme
    if current_theme is not None and not current_theme_has_body:
        empty_themes.append(current_theme)

    missing = []
    if theme_count == 0:
        missing.append("theme header")
    if theme_count > 0 and theme_count == len(empty_themes):
        missing.append("body text")
    if missing:
        raise ValueError(f"Document missing required content: {', '.join(missing)}")

    if empty_themes:
        logger.warning(
            "document.empty_themes",
            empty_themes=empty_themes,
            total_themes=theme_count,
        )


def auto_bold_html_metrics(text: str) -> str:
    """
    Auto-bold financial metrics not already wrapped in <b> tags.

    HTML equivalent of call_summary's auto_bold_metrics(). Catches dollar amounts,
    percentages, and basis point figures that the LLM failed to bold. Does not modify
    text already inside <b>...</b> tags.

    Args:
        text: HTML-formatted text potentially containing unbolded metrics

    Returns:
        Text with financial metrics wrapped in <b>...</b> tags
    """
    if not text:
        return text

    for pattern in _METRIC_PATTERNS:
        text = _bold_html_unbolded_matches(text, pattern)

    return text


def _bold_html_unbolded_matches(text: str, pattern: str) -> str:
    """Bold regex matches that are not already inside <b>...</b> tags."""

    def replacer(match):
        start = match.start()
        prefix = text[:start]

        # Count unclosed <b> tags — if inside a <b> block, skip
        open_count = len(re.findall(r"<b\b[^>]*>", prefix, re.IGNORECASE))
        close_count = len(re.findall(r"</b>", prefix, re.IGNORECASE))
        if open_count > close_count:
            return match.group(0)

        # Check if immediately preceded by <b> opening tag
        stripped = prefix.rstrip()
        if stripped.endswith(">"):
            tag_match = re.search(r"<b\b[^>]*>\s*$", prefix, re.IGNORECASE)
            if tag_match:
                return match.group(0)

        return f"<b>{match.group(0)}</b>"

    return re.sub(pattern, replacer, text)


class HTMLToDocx(HTMLParser):
    """
    Parser to convert HTML-formatted text to Word document formatting.
    Supports: <b>, <strong>, <i>, <em>, <u>, <mark>, <span style="..."> tags and their nesting.
    """

    def __init__(self, paragraph, font_size=Pt(9)):
        super().__init__()
        self.paragraph = paragraph
        self.font_size = font_size
        self.text_buffer = ""
        self.format_stack = []
        self.style_stack = []

    def handle_starttag(self, tag, attrs):
        self._flush_text()

        if tag in ["b", "strong"]:
            self.format_stack.append("bold")
        elif tag in ["i", "em"]:
            self.format_stack.append("italic")
        elif tag == "u":
            self.format_stack.append("underline")
        elif tag == "mark":
            style_dict = self._parse_style(attrs)
            if style_dict and "background-color" in style_dict:
                self.format_stack.append("highlight_yellow")
            else:
                self.format_stack.append("highlight")
        elif tag == "span":
            style_dict = self._parse_style(attrs)
            self.style_stack.append(style_dict if style_dict else {})
        elif tag == "br":
            # Line break — flush buffer and add a newline run
            self.text_buffer += "\n"

    def handle_endtag(self, tag):
        self._flush_text()

        if tag in ["b", "strong"] and "bold" in self.format_stack:
            self.format_stack.remove("bold")
        elif tag in ["i", "em"] and "italic" in self.format_stack:
            self.format_stack.remove("italic")
        elif tag == "u" and "underline" in self.format_stack:
            self.format_stack.remove("underline")
        elif tag == "mark":
            if "highlight_yellow" in self.format_stack:
                self.format_stack.remove("highlight_yellow")
            elif "highlight" in self.format_stack:
                self.format_stack.remove("highlight")
        elif tag == "span" and self.style_stack:
            self.style_stack.pop()

    def handle_data(self, data):
        self.text_buffer += data

    def _parse_style(self, attrs):
        """Parse style attribute from HTML tags."""
        style_dict = {}
        for attr_name, attr_value in attrs:
            if attr_name == "style":

                styles = attr_value.split(";")
                for style in styles:
                    if ":" in style:
                        prop, value = style.split(":", 1)
                        style_dict[prop.strip().lower()] = value.strip()
        return style_dict

    def _flush_text(self):
        """Apply formatting and add text to document."""
        if not self.text_buffer:
            return

        run = self.paragraph.add_run(self.text_buffer)

        if self.style_stack:
            current_style = self.style_stack[-1]

            if "color" in current_style:
                color_hex = current_style["color"].lstrip("#")
                if color_hex.startswith("1e4d8b"):
                    run.font.color.rgb = RGBColor(0x1E, 0x4D, 0x8B)
                elif color_hex.startswith("4d94ff"):
                    run.font.color.rgb = RGBColor(0x4D, 0x94, 0xFF)

            if "font-size" in current_style:
                size_str = current_style["font-size"]
                if "pt" in size_str:
                    size = float(size_str.replace("pt", "").strip())
                    run.font.size = Pt(size)
                else:
                    run.font.size = self.font_size
            else:
                run.font.size = self.font_size

            if "font-weight" in current_style and current_style["font-weight"] == "bold":
                run.font.bold = True
        else:
            run.font.size = self.font_size

        if "bold" in self.format_stack:
            run.font.bold = True
        if "italic" in self.format_stack:
            run.font.italic = True
        if "underline" in self.format_stack:
            run.font.underline = True
        if "highlight" in self.format_stack or "highlight_yellow" in self.format_stack:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        self.text_buffer = ""

    def close(self):
        """Ensure any remaining text is flushed."""
        self._flush_text()
        super().close()


def add_page_numbers_with_footer(doc, bank_symbol, quarter, fiscal_year):
    """Add custom footer with bank info, document title, and page numbers."""
    for section in doc.sections:
        footer = section.footer

        footer.paragraphs[0].clear()

        tbl = footer.add_table(
            1,
            2,
            width=doc.sections[0].page_width
            - doc.sections[0].left_margin
            - doc.sections[0].right_margin,
        )
        tbl.autofit = False
        tbl.allow_autofit = False

        tbl.style = "Table Grid"
        for row in tbl.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    '<w:top w:val="nil"/><w:left w:val="nil"/>'
                    '<w:bottom w:val="nil"/><w:right w:val="nil"/></w:tcBorders>'
                )
                tcPr.append(tcBorders)

        left_cell = tbl.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_text = (
            f"{bank_symbol} | {quarter}/{str(fiscal_year)[-2:]} | Investor Call - Key Themes"
        )
        run_left = left_para.add_run(left_text)
        run_left.font.size = Pt(9)
        run_left.font.color.rgb = RGBColor(68, 68, 68)

        right_cell = tbl.rows[0].cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run_label = right_para.add_run("Page ")
        run_label.font.size = Pt(9)
        run_label.font.color.rgb = RGBColor(68, 68, 68)

        run = right_para.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(68, 68, 68)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._element.append(fldChar1)

        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        run._element.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._element.append(fldChar2)


def add_theme_header_with_background(doc, theme_number, theme_title):
    """Add a theme header with dark blue text on light blue background."""
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(8)
    heading.paragraph_format.keep_with_next = True

    full_text = f"Theme {theme_number}: {theme_title}"
    run = heading.add_run(full_text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:val"), "clear")
    shading_elm.set(qn("w:color"), "auto")
    shading_elm.set(qn("w:fill"), "D4E1F5")
    heading._element.get_or_add_pPr().append(shading_elm)

    return heading


def add_banner_image(doc, config_dir: str) -> None:
    """Add banner image to document if found in config directory."""
    banner_path = None
    for ext in ["jpg", "jpeg", "png"]:
        potential_banner = os.path.join(config_dir, f"banner.{ext}")
        if os.path.exists(potential_banner):
            banner_path = potential_banner
            break

    if banner_path:
        try:
            doc.add_picture(banner_path, width=Inches(7.4))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(6)
        except (FileNotFoundError, OSError, ValueError):
            pass


def get_standard_report_metadata() -> Dict[str, str]:
    """
    Get standard metadata for key themes reports.

    Returns:
        Dictionary with report_name, report_description, and report_type
    """
    return {
        "report_name": "Key Themes Analysis",
        "report_description": (
            "AI-generated thematic analysis of earnings call Q&A sessions, "
            "identifying and grouping key discussion topics between analysts "
            "and executives. Provides consolidated insights into major themes "
            "with supporting conversation excerpts."
        ),
        "report_type": "key_themes",
    }
