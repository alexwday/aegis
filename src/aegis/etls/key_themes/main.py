"""
Key Themes ETL - Extract and Group Earnings Call Themes

This ETL processes earnings call Q&A sessions to:
1. Load all Q&A blocks into an index
2. Process each independently to extract themes (parallelizable)
3. Make ONE comprehensive grouping decision with full visibility
4. Apply grouping programmatically
5. Generate formatted document

Usage:
    python -m aegis.etls.key_themes.main --bank "Royal Bank of Canada" --year 2024 --quarter Q3
"""

import argparse
import asyncio
import json
import random
import time
import uuid
import os
import sys
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime
from functools import lru_cache
from typing import Dict, Any, List, Optional
from pydantic import BaseModel, ValidationError
import pandas as pd
from sqlalchemy import text
from sqlalchemy.exc import SQLAlchemyError
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from aegis.etls.key_themes.document_converter import (
    HTMLToDocx,
    add_banner_image,
    add_page_numbers_with_footer,
    add_theme_header_with_background,
    get_standard_report_metadata,
    validate_document_content,
    auto_bold_html_metrics,
)
from aegis.etls.key_themes.transcript_utils import (
    retrieve_full_section,
    SECTIONS_KEY_MD,
    SECTIONS_KEY_QA,
    SECTIONS_KEY_ALL,
    VALID_SECTION_KEYS,
)
from aegis.utils.ssl import setup_ssl
from aegis.connections.oauth_connector import setup_authentication
from aegis.connections.llm_connector import complete, complete_with_tools
from aegis.connections.postgres_connector import get_connection
from aegis.utils.logging import setup_logging, get_logger
from aegis.utils.prompt_loader import load_prompt_from_db
from aegis.utils.sql_prompt import postgresql_prompts
from aegis.utils.settings import config

import yaml

setup_logging()
logger = get_logger()

# --- Module-level defaults (used as fallbacks if YAML config keys are missing) ---
MAX_RETRIES = 3
RETRY_BASE_DELAY = 1.0
RETRY_MAX_DELAY = 10.0

# --- Display / Logging Configuration ---
LOG_SNIPPET_LENGTH = 80

# --- Concurrency Configuration ---
MAX_CONCURRENT_FORMATTING = 5


# --- ETL Exception Hierarchy ---


class KeyThemesError(Exception):
    """Base exception for key themes ETL errors."""


class KeyThemesSystemError(KeyThemesError):
    """Unexpected system/infrastructure error."""


class KeyThemesUserError(KeyThemesError):
    """Expected user-facing error (bad input, no data, etc.)."""


@dataclass
class KeyThemesResult:
    """Successful key themes generation result."""

    filepath: str
    theme_groups: int
    valid_qa: int
    invalid_qa_filtered: int
    total_cost: float = 0.0
    total_tokens: int = 0


# --- Pydantic Models for LLM Response Validation ---


class ThemeExtractionResponse(BaseModel):
    """Validation model for theme extraction LLM response."""

    is_valid: bool
    completion_status: str = ""
    reasoning: str = ""
    category_name: str = ""
    summary: str = ""
    rejection_reason: str = ""


class ThemeGroupItem(BaseModel):
    """A single theme group from the grouping LLM response."""

    group_title: str
    qa_ids: List[str]
    rationale: str = ""


class ThemeGroupingResponse(BaseModel):
    """Top-level theme grouping response from the LLM."""

    theme_groups: List[ThemeGroupItem]


class ETLConfig:
    """
    ETL configuration loader that reads YAML configs and resolves model references.

    This class loads ETL-specific configuration from YAML files and provides
    easy access to configuration values with automatic model tier resolution.
    """

    def __init__(self, config_path: str):
        """Initialize the ETL configuration loader."""
        self.config_path = config_path
        self._config = self._load_config()

    def _load_config(self) -> Dict[str, Any]:
        """Load configuration from YAML file."""
        if not os.path.exists(self.config_path):
            raise FileNotFoundError(f"Configuration file not found: {self.config_path}")

        with open(self.config_path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f)

    def get_model(self, model_key: str) -> str:
        """
        Get the actual model name for a given model key.

        Resolves model tier references (small/medium/large) to actual model names
        from the global settings configuration.
        """
        if "models" not in self._config or model_key not in self._config["models"]:
            raise KeyError(f"Model key '{model_key}' not found in configuration")

        tier = self._config["models"][model_key].get("tier")
        if not tier:
            raise ValueError(f"No tier specified for model '{model_key}'")

        # Resolve tier to actual model from global config
        tier_map = {
            "small": config.llm.small.model,
            "medium": config.llm.medium.model,
            "large": config.llm.large.model,
        }

        if tier not in tier_map:
            raise ValueError(
                f"Invalid tier '{tier}' for model '{model_key}'. "
                f"Valid tiers: {list(tier_map.keys())}"
            )

        return tier_map[tier]

    @property
    def temperature(self) -> float:
        """Get the LLM temperature parameter."""
        return self._config.get("llm", {}).get("temperature", 0.1)

    def get_max_tokens(self, task_key: str) -> int:
        """
        Get max_tokens for a specific task, falling back to default.

        Supports both legacy (flat int) and per-task (dict) config formats.

        Args:
            task_key: Task identifier (e.g., "theme_extraction", "html_formatting")

        Returns:
            max_tokens value for the task
        """
        max_tokens_config = self._config.get("llm", {}).get("max_tokens", {})
        if isinstance(max_tokens_config, int):
            return max_tokens_config
        return max_tokens_config.get(task_key, max_tokens_config.get("default", 32768))

    @property
    def max_concurrent_formatting(self) -> int:
        """Get the maximum number of concurrent HTML formatting tasks."""
        return self._config.get("concurrency", {}).get(
            "max_concurrent_formatting", MAX_CONCURRENT_FORMATTING
        )

    @property
    def max_retries(self) -> int:
        """Get the maximum number of LLM call retries."""
        return self._config.get("retry", {}).get("max_retries", MAX_RETRIES)

    @property
    def retry_base_delay(self) -> float:
        """Get the base delay in seconds for retry backoff."""
        return self._config.get("retry", {}).get("base_delay", RETRY_BASE_DELAY)

    @property
    def retry_max_delay(self) -> float:
        """Get the maximum delay in seconds for retry backoff."""
        return self._config.get("retry", {}).get("max_delay", RETRY_MAX_DELAY)


etl_config = ETLConfig(os.path.join(os.path.dirname(__file__), "config", "config.yaml"))


@lru_cache(maxsize=1)
def _load_monitored_institutions() -> Dict[int, Dict[str, Any]]:
    """
    Load and cache monitored institutions configuration.

    Uses @lru_cache for thread-safe, testable caching (clear via .cache_clear()).

    Returns:
        Dictionary mapping bank_id to institution details (id, name, symbol, type, path_safe_name)
    """
    config_path = os.path.join(os.path.dirname(__file__), "config", "monitored_institutions.yaml")
    with open(config_path, "r", encoding="utf-8") as f:
        yaml_data = yaml.safe_load(f)

    institutions = {}
    for key, value in yaml_data.items():
        symbol = key.split("-")[0]  # Extract symbol from "RY-CA" -> "RY"
        institutions[value["id"]] = {**value, "symbol": symbol}
    return institutions


def _sanitize_for_prompt(text_val: str) -> str:
    """
    Escape curly braces in text for safe use in .format() templates.

    Prevents KeyError/IndexError when XLSX-sourced content contains { or }.

    Args:
        text_val: Raw text string

    Returns:
        Text with { and } escaped as {{ and }}
    """
    return text_val.replace("{", "{{").replace("}", "}}")


def _timing_summary(marks: list) -> dict:
    """
    Convert timing marks to elapsed-seconds summary.

    Args:
        marks: List of (stage_name, timestamp) tuples

    Returns:
        Dict mapping "{stage}_s" to elapsed seconds, plus "total_s"
    """
    if len(marks) < 2:
        return {}
    summary = {}
    for i in range(1, len(marks)):
        summary[f"{marks[i][0]}_s"] = round(marks[i][1] - marks[i - 1][1], 2)
    summary["total_s"] = round(marks[-1][1] - marks[0][1], 2)
    return summary


def _accumulate_llm_cost(context: Dict[str, Any], metrics: Dict[str, Any]) -> None:
    """
    Accumulate LLM cost and token usage from a response's metrics dict.

    Appends to context["_llm_costs"] list for later aggregation.
    Safe for asyncio concurrency (single-threaded event loop).

    Args:
        context: Execution context with _llm_costs list
        metrics: Response metrics dict with prompt_tokens, completion_tokens, total_cost
    """
    if "_llm_costs" not in context:
        context["_llm_costs"] = []
    context["_llm_costs"].append(
        {
            "prompt_tokens": metrics.get("prompt_tokens", 0),
            "completion_tokens": metrics.get("completion_tokens", 0),
            "total_cost": metrics.get("total_cost", 0),
        }
    )


def _get_total_llm_cost(context: Dict[str, Any]) -> Dict[str, Any]:
    """
    Aggregate all accumulated LLM costs from context.

    Returns:
        Dict with total_cost, total_prompt_tokens, total_completion_tokens, total_tokens
    """
    costs = context.get("_llm_costs", [])
    total_cost = sum(c.get("total_cost", 0) for c in costs)
    prompt_tokens = sum(c.get("prompt_tokens", 0) for c in costs)
    completion_tokens = sum(c.get("completion_tokens", 0) for c in costs)
    return {
        "total_cost": round(total_cost, 4),
        "total_prompt_tokens": prompt_tokens,
        "total_completion_tokens": completion_tokens,
        "total_tokens": prompt_tokens + completion_tokens,
        "llm_calls": len(costs),
    }


def get_bank_info_from_config(bank_identifier: str) -> Dict[str, Any]:
    """
    Look up bank from monitored institutions configuration file.

    Args:
        bank_identifier: Bank ID (as string/int), symbol (e.g., "RY"), or name

    Returns:
        Dictionary with bank_id, bank_name, bank_symbol, bank_type

    Raises:
        ValueError: If bank not found in monitored institutions
    """
    institutions = _load_monitored_institutions()

    # Try lookup by ID
    if bank_identifier.isdigit():
        bank_id = int(bank_identifier)
        if bank_id in institutions:
            inst = institutions[bank_id]
            return {
                "bank_id": inst["id"],
                "bank_name": inst["name"],
                "bank_symbol": inst["symbol"],
                "bank_type": inst["type"],
            }

    # Try lookup by symbol or name
    bank_identifier_upper = bank_identifier.upper()
    bank_identifier_lower = bank_identifier.lower()

    for inst in institutions.values():
        # Match by symbol (case-insensitive)
        if inst["symbol"].upper() == bank_identifier_upper:
            return {
                "bank_id": inst["id"],
                "bank_name": inst["name"],
                "bank_symbol": inst["symbol"],
                "bank_type": inst["type"],
            }

        # Match by name (case-insensitive, partial match)
        if bank_identifier_lower in inst["name"].lower():
            return {
                "bank_id": inst["id"],
                "bank_name": inst["name"],
                "bank_symbol": inst["symbol"],
                "bank_type": inst["type"],
            }

    # Build helpful error message with available banks
    available = [f"{inst['symbol']} ({inst['name']})" for inst in institutions.values()]
    raise ValueError(
        f"Bank '{bank_identifier}' not found in monitored institutions.\n"
        f"Available banks: {', '.join(sorted(available))}"
    )


async def verify_and_get_availability(
    bank_id: int, bank_name: str, fiscal_year: int, quarter: str
) -> None:
    """
    Verify transcript data is available for the specified bank and period.

    Raises ValueError with available periods if data not found.

    Args:
        bank_id: Bank ID
        bank_name: Bank name (for error messages)
        fiscal_year: Year (e.g., 2024)
        quarter: Quarter (e.g., "Q3")

    Raises:
        ValueError: If transcript data not available (includes available periods)
    """
    async with get_connection() as conn:
        result = await conn.execute(
            text(
                """
                SELECT database_names
                FROM aegis_data_availability
                WHERE bank_id = :bank_id
                  AND fiscal_year = :fiscal_year
                  AND quarter = :quarter
                """
            ),
            {"bank_id": bank_id, "fiscal_year": fiscal_year, "quarter": quarter},
        )
        row = result.fetchone()

        if row and row[0] and "transcripts" in row[0]:
            return

        # Query available transcript periods for this bank to provide helpful context
        available_result = await conn.execute(
            text(
                """
                SELECT fiscal_year, quarter
                FROM aegis_data_availability
                WHERE bank_id = :bank_id
                  AND 'transcripts' = ANY(database_names)
                ORDER BY fiscal_year DESC, quarter DESC
                """
            ),
            {"bank_id": bank_id},
        )
        available_periods = [f"{r[1]} {r[0]}" for r in available_result.fetchall()]

        if available_periods:
            raise ValueError(
                f"No transcript data available for {bank_name} {quarter} {fiscal_year}. "
                f"Available periods: {', '.join(available_periods)}"
            )
        raise ValueError(
            f"No transcript data available for {bank_name} {quarter} {fiscal_year}. "
            f"No transcript data found for this bank in any period."
        )


def format_categories_for_prompt(categories: List[Dict[str, Any]]) -> str:
    """
    Format category dictionaries into standardized XML format for prompt injection.

    This is the standardized formatting function used across all ETLs (Call Summary,
    Key Themes, CM Readthrough) to ensure consistent category presentation to LLMs.

    Args:
        categories: List of category dicts with standardized 6-column format

    Returns:
        Formatted XML string with category information
    """
    formatted_sections = []

    for cat in categories:
        # Map transcript_sections to human-readable description
        section_desc = {
            SECTIONS_KEY_MD: "Management Discussion section only",
            SECTIONS_KEY_QA: "Q&A section only",
            SECTIONS_KEY_ALL: "Both Management Discussion and Q&A sections",
        }.get(cat.get("transcript_sections", SECTIONS_KEY_ALL), "ALL sections")

        section = "<category>\n"
        section += f"<name>{_sanitize_for_prompt(cat['category_name'])}</name>\n"
        section += f"<section>{section_desc}</section>\n"
        section += (
            f"<description>{_sanitize_for_prompt(cat['category_description'])}</description>\n"
        )

        # Collect non-empty examples
        examples = []
        for i in range(1, 4):
            example_key = f"example_{i}"
            if (
                cat.get(example_key)
                and isinstance(cat[example_key], str)
                and cat[example_key].strip()
            ):
                examples.append(cat[example_key])

        if examples:
            section += "<examples>\n"
            for example in examples:
                section += f"  <example>{_sanitize_for_prompt(example)}</example>\n"
            section += "</examples>\n"

        section += "</category>"
        formatted_sections.append(section)

    return "\n\n".join(formatted_sections)


def load_categories_from_xlsx(execution_id: str) -> List[Dict[str, Any]]:
    """
    Load categories from the key themes categories XLSX file.

    Args:
        execution_id: Execution ID for logging

    Returns:
        List of dictionaries with transcript_sections, category_name, category_description,
        example_1, example_2, example_3
    """
    file_name = "key_themes_categories.xlsx"

    current_dir = os.path.dirname(os.path.abspath(__file__))
    xlsx_path = os.path.join(current_dir, "config", "categories", file_name)

    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(f"Categories file not found: {xlsx_path}")

    try:
        df = pd.read_excel(xlsx_path, sheet_name=0)

        # Required columns for standard format
        required_columns = ["transcript_sections", "category_name", "category_description"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"Missing required columns in {file_name}: {missing_columns}")

        # Optional example columns
        optional_columns = ["example_1", "example_2", "example_3"]
        for col in optional_columns:
            if col not in df.columns:
                df[col] = ""  # Add empty column if not present

        # Convert to list of dicts, ensuring all required fields are non-empty
        categories = []
        for idx, row in df.iterrows():
            for field in required_columns:
                if pd.isna(row[field]) or str(row[field]).strip() == "":
                    raise ValueError(f"Missing value for '{field}' in {file_name} (row {idx + 2})")

            transcript_sections = str(row["transcript_sections"]).strip()
            if transcript_sections not in VALID_SECTION_KEYS:
                raise ValueError(
                    f"Invalid transcript_sections '{transcript_sections}' "
                    f"in {file_name} (row {idx + 2}). Must be one of: {VALID_SECTION_KEYS}"
                )

            category = {
                "transcript_sections": transcript_sections,
                "category_name": str(row["category_name"]).strip(),
                "category_description": str(row["category_description"]).strip(),
                "example_1": str(row["example_1"]).strip() if pd.notna(row["example_1"]) else "",
                "example_2": str(row["example_2"]).strip() if pd.notna(row["example_2"]) else "",
                "example_3": str(row["example_3"]).strip() if pd.notna(row["example_3"]) else "",
            }
            categories.append(category)

        if not categories:
            raise ValueError(f"No categories in {file_name}")

        logger.info(
            "etl.key_themes.categories_loaded",
            execution_id=execution_id,
            file_name=file_name,
            num_categories=len(categories),
        )
        return categories

    except Exception as e:
        logger.error(
            "etl.key_themes.categories_load_error",
            execution_id=execution_id,
            xlsx_path=xlsx_path,
            error=str(e),
        )
        raise RuntimeError(f"Failed to load categories from {xlsx_path}: {str(e)}") from e


@dataclass
class QABlock:
    """Represents a single Q&A block with its extracted information."""

    qa_id: str
    position: int
    original_content: str
    category_name: Optional[str] = None
    summary: Optional[str] = None
    formatted_content: Optional[str] = None
    assigned_group: Optional[str] = None
    is_valid: bool = True
    completion_status: str = "complete"


@dataclass
class ThemeGroup:
    """Represents a group of related Q&A blocks under a unified theme."""

    group_title: str
    qa_ids: List[str]
    rationale: str = ""
    qa_blocks: list = field(default_factory=list)


async def load_qa_blocks(
    bank_name: str, fiscal_year: int, quarter: str, context: Dict[str, Any]
) -> Dict[str, QABlock]:
    """
    Step 1: Load all Q&A blocks and create an index.
    Uses retrieve_full_section from transcript_utils for consistency.

    Returns:
        Dictionary indexed by qa_id containing QABlock objects
    """
    bank_info = get_bank_info_from_config(bank_name)

    combo = {
        "bank_name": bank_info["bank_name"],
        "bank_id": bank_info["bank_id"],
        "bank_symbol": bank_info["bank_symbol"],
        "fiscal_year": fiscal_year,
        "quarter": quarter,
    }

    # Use retrieve_full_section to get Q&A chunks
    chunks = await retrieve_full_section(combo=combo, sections=SECTIONS_KEY_QA, context=context)

    logger.debug(
        "load_qa_index.retrieved",
        bank_name=bank_name,
        quarter=quarter,
        fiscal_year=fiscal_year,
        num_chunks=len(chunks),
    )

    if not chunks:
        return {}

    # Group chunks by qa_group_id
    qa_groups = {}
    for chunk in chunks:
        qa_group_id = chunk.get("qa_group_id")
        if qa_group_id is not None:
            if qa_group_id not in qa_groups:
                qa_groups[qa_group_id] = []
            qa_groups[qa_group_id].append(chunk)

    # Build QABlock index
    qa_index = {}
    for qa_group_id, group_chunks in qa_groups.items():
        # Concatenate content from all chunks in this Q&A group
        qa_content = "\n".join(
            [chunk.get("content", "") for chunk in group_chunks if chunk.get("content")]
        )

        if qa_content:
            qa_id = f"qa_{qa_group_id}"
            qa_index[qa_id] = QABlock(qa_id, qa_group_id, qa_content)

    return qa_index


async def classify_qa_block(
    qa_block: QABlock,
    categories: List[Dict[str, str]],
    previous_classifications: List[Dict[str, str]],
    extraction_prompts: Dict[str, Any],
    context: Dict[str, Any],
):
    """
    Step 1A: Validate and classify a single Q&A block into predefined category.
    Uses cumulative context from previous classifications for consistency.

    Args:
        qa_block: Q&A block to classify
        categories: List of predefined categories from xlsx
        previous_classifications: List of prior classifications for context
        extraction_prompts: Prompt data loaded from DB (theme_extraction)
        context: Execution context
    """
    execution_id = context.get("execution_id")

    logger.info(
        "category_classification.processing_qa",
        execution_id=execution_id,
        qa_id=qa_block.qa_id,
        content_length=len(qa_block.original_content),
        num_previous_classifications=len(previous_classifications),
    )

    prompt_data = extraction_prompts

    # Format categories using standardized XML format
    categories_str = format_categories_for_prompt(categories)

    # Format previous classifications
    if previous_classifications:
        prev_class_str = "\n".join(
            [
                f"{pc['qa_id']}: {pc['category_name']} - {pc['summary'][:100]}..."
                for pc in previous_classifications
            ]
        )
    else:
        prev_class_str = "No previous classifications yet (this is the first Q&A)."

    system_prompt = prompt_data["system_prompt"].format(
        bank_name=context.get("bank_name", "Bank"),
        quarter=context.get("quarter", "Q"),
        fiscal_year=context.get("fiscal_year", "Year"),
        categories_list=categories_str,
        num_categories=len(categories),
        previous_classifications=prev_class_str,
    )

    user_prompt = prompt_data["user_prompt"].format(
        qa_content=_sanitize_for_prompt(qa_block.original_content)
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    result = None

    for attempt in range(etl_config.max_retries):
        try:
            response = await complete_with_tools(
                messages=messages,
                tools=[prompt_data["tool_definition"]],
                context=context,
                llm_params={
                    "model": etl_config.get_model("theme_extraction"),
                    "temperature": etl_config.temperature,
                    "max_tokens": etl_config.get_max_tokens("theme_extraction"),
                },
            )

            if response:
                tool_calls = (
                    response.get("choices", [{}])[0].get("message", {}).get("tool_calls", [])
                )
                if tool_calls:
                    raw_data = json.loads(tool_calls[0]["function"]["arguments"])
                    result = ThemeExtractionResponse.model_validate(raw_data)

                    # Log and accumulate LLM usage
                    metrics = response.get("metrics", {})
                    _accumulate_llm_cost(context, metrics)
                    logger.info(
                        "etl.key_themes.llm_usage",
                        execution_id=execution_id,
                        stage=f"classification:{qa_block.qa_id}",
                        prompt_tokens=metrics.get("prompt_tokens", 0),
                        completion_tokens=metrics.get("completion_tokens", 0),
                        total_cost=metrics.get("total_cost", 0),
                        response_time=metrics.get("response_time", 0),
                    )
                    break

        except (KeyError, IndexError, json.JSONDecodeError, TypeError, ValidationError) as e:
            logger.warning(
                "category_classification.parse_error",
                execution_id=execution_id,
                qa_id=qa_block.qa_id,
                attempt=attempt + 1,
                error=str(e),
            )
            if attempt < etl_config.max_retries - 1:
                continue
            raise RuntimeError(
                f"Failed to parse classification for {qa_block.qa_id} "
                f"after {etl_config.max_retries} retries"
            ) from e
        except Exception as e:
            logger.error(
                "etl.key_themes.classification_error",
                execution_id=execution_id,
                qa_id=qa_block.qa_id,
                error=str(e),
                attempt=attempt + 1,
            )
            if attempt < etl_config.max_retries - 1:
                delay = min(etl_config.retry_base_delay * (2**attempt), etl_config.retry_max_delay)
                delay += random.uniform(0, 0.5 * delay)
                await asyncio.sleep(delay)
                continue
            raise RuntimeError(
                f"Classification failed for {qa_block.qa_id} "
                f"after {etl_config.max_retries} retries: {e}"
            ) from e

    if not result:
        raise RuntimeError(
            f"Classification failed for {qa_block.qa_id}: "
            f"LLM returned no tool calls after {etl_config.max_retries} retries"
        )

    if not result.is_valid:
        logger.warning(
            "category_classification.rejected",
            execution_id=execution_id,
            qa_id=qa_block.qa_id,
            rejection_reason=result.rejection_reason or "No reason provided",
        )
        qa_block.is_valid = False
        qa_block.category_name = None
        qa_block.summary = None
        qa_block.completion_status = ""
    else:
        logger.info(
            "category_classification.accepted",
            execution_id=execution_id,
            qa_id=qa_block.qa_id,
            category_name=result.category_name,
            completion_status=result.completion_status,
            summary_preview=result.summary[:100],
        )

        qa_block.is_valid = True
        qa_block.category_name = result.category_name
        qa_block.summary = result.summary
        qa_block.completion_status = result.completion_status or "complete"


async def format_qa_html(
    qa_block: QABlock, formatting_prompts: Dict[str, Any], context: Dict[str, Any]
):
    """
    Step 2B: Format Q&A block with HTML tags for emphasis.
    Only formats valid Q&A blocks that passed validation.

    Args:
        qa_block: Q&A block to format
        formatting_prompts: Prompt data loaded from DB (html_formatting)
        context: Execution context
    """

    if not qa_block.is_valid:
        qa_block.formatted_content = None
        return

    execution_id = context.get("execution_id")
    prompt_data = formatting_prompts

    system_prompt = prompt_data["system_prompt"].format(
        bank_name=context.get("bank_name", "Bank"),
        quarter=context.get("quarter", "Q"),
        fiscal_year=context.get("fiscal_year", "Year"),
    )

    # Build completion note for incomplete Q&As
    completion_note = ""
    if qa_block.completion_status == "question_only":
        completion_note = (
            "[NOTE: This block contains only the analyst question. "
            "Executive response may be in a separate block.]\n\n"
        )
    elif qa_block.completion_status == "answer_only":
        completion_note = (
            "[NOTE: This block contains only the executive response. "
            "Analyst question may be in a separate block.]\n\n"
        )

    user_prompt = prompt_data["user_prompt"].format(
        completion_note=completion_note,
        qa_content=_sanitize_for_prompt(qa_block.original_content),
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    for attempt in range(etl_config.max_retries):
        try:
            response = await complete(
                messages,
                context,
                {
                    "model": etl_config.get_model("html_formatting"),
                    "temperature": etl_config.temperature,
                    "max_tokens": etl_config.get_max_tokens("html_formatting"),
                },
            )

            if isinstance(response, dict):
                qa_block.formatted_content = (
                    response.get("choices", [{}])[0].get("message", {}).get("content", "")
                )
            else:
                qa_block.formatted_content = str(response)

            # Log and accumulate LLM usage
            if isinstance(response, dict):
                metrics = response.get("metrics", {})
                _accumulate_llm_cost(context, metrics)
                logger.info(
                    "etl.key_themes.llm_usage",
                    execution_id=execution_id,
                    stage=f"formatting:{qa_block.qa_id}",
                    prompt_tokens=metrics.get("prompt_tokens", 0),
                    completion_tokens=metrics.get("completion_tokens", 0),
                    total_cost=metrics.get("total_cost", 0),
                    response_time=metrics.get("response_time", 0),
                )
            break

        except Exception as e:
            if attempt < etl_config.max_retries - 1:
                delay = min(etl_config.retry_base_delay * (2**attempt), etl_config.retry_max_delay)
                delay += random.uniform(0, 0.5 * delay)
                await asyncio.sleep(delay)
            else:
                raise RuntimeError(
                    f"HTML formatting failed for {qa_block.qa_id} "
                    f"after {etl_config.max_retries} retries: {e}"
                ) from e


async def classify_all_qa_blocks_sequential(
    qa_index: Dict[str, QABlock],
    categories: List[Dict[str, str]],
    extraction_prompts: Dict[str, Any],
    context: Dict[str, Any],
):
    """
    Step 1: Classify all Q&A blocks sequentially with cumulative context.

    Args:
        qa_index: Dictionary of QA blocks indexed by qa_id
        categories: List of predefined categories from xlsx
        extraction_prompts: Prompt data loaded from DB (theme_extraction)
        context: Execution context
    """
    previous_classifications = []

    # Process in order by position
    sorted_qa_blocks = sorted(qa_index.values(), key=lambda x: x.position)

    for qa_block in sorted_qa_blocks:
        await classify_qa_block(
            qa_block, categories, previous_classifications, extraction_prompts, context
        )

        # Add to cumulative context if valid
        if qa_block.is_valid:
            previous_classifications.append(
                {
                    "qa_id": qa_block.qa_id,
                    "category_name": qa_block.category_name,
                    "summary": qa_block.summary,
                }
            )

    logger.info(
        "classification.completed",
        execution_id=context.get("execution_id"),
        total_classified=len(previous_classifications),
        total_invalid=len([qa for qa in qa_index.values() if not qa.is_valid]),
    )


async def format_all_qa_blocks_parallel(
    qa_index: Dict[str, QABlock], formatting_prompts: Dict[str, Any], context: Dict[str, Any]
):
    """
    Step 2: Format all valid Q&A blocks in parallel with HTML tags.

    Uses semaphore for concurrency control.

    Args:
        qa_index: Dictionary of QA blocks indexed by qa_id
        formatting_prompts: Prompt data loaded from DB (html_formatting)
        context: Execution context
    """
    # Only format valid Q&As
    valid_qa_blocks = [qa for qa in qa_index.values() if qa.is_valid]

    if not valid_qa_blocks:
        return

    semaphore = asyncio.Semaphore(etl_config.max_concurrent_formatting)

    async def _format_with_semaphore(qa_block: QABlock):
        async with semaphore:
            await format_qa_html(qa_block, formatting_prompts, context)

    # Parallel formatting with concurrency limit
    tasks = [_format_with_semaphore(qa_block) for qa_block in valid_qa_blocks]
    await asyncio.gather(*tasks)

    logger.info(
        "formatting.completed",
        execution_id=context.get("execution_id"),
        total_formatted=len(valid_qa_blocks),
    )


async def determine_comprehensive_grouping(
    qa_index: Dict[str, QABlock],
    categories: List[Dict[str, str]],
    grouping_prompts: Dict[str, Any],
    context: Dict[str, Any],
) -> tuple:
    """
    Step 3: Make ONE comprehensive grouping decision for all themes.
    Only processes valid Q&A blocks.

    Args:
        qa_index: Dictionary of QA blocks indexed by qa_id
        categories: List of category definitions from xlsx
        grouping_prompts: Prompt data loaded from DB (theme_grouping)
        context: Execution context with auth, ssl, bank info

    Returns:
        List of ThemeGroup objects

    Raises:
        RuntimeError: If grouping fails after all retries
    """

    valid_qa_blocks = {qa_id: qa_block for qa_id, qa_block in qa_index.items() if qa_block.is_valid}

    if not valid_qa_blocks:
        return []

    execution_id = context.get("execution_id")
    prompt_data = grouping_prompts

    qa_blocks_info = []
    for qa_id, qa_block in sorted(valid_qa_blocks.items(), key=lambda x: x[1].position):
        qa_blocks_info.append(
            f"ID: {qa_id}\n"
            f"Category: {qa_block.category_name}\n"
            f"Summary: {qa_block.summary}\n"
        )

    qa_blocks_str = _sanitize_for_prompt("\n\n".join(qa_blocks_info))

    # Format categories using standardized XML format
    categories_str = format_categories_for_prompt(categories)

    system_prompt = prompt_data["system_prompt"].format(
        bank_name=context.get("bank_name", "Bank"),
        bank_symbol=context.get("bank_symbol", "BANK"),
        quarter=context.get("quarter", "Q"),
        fiscal_year=context.get("fiscal_year", "Year"),
        total_qa_blocks=len(valid_qa_blocks),
        qa_blocks_info=qa_blocks_str,
        categories_list=categories_str,
        num_categories=len(categories),
    )

    user_prompt = prompt_data["user_prompt"]

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    logger.info(
        "regrouping.llm_request",
        execution_id=execution_id,
        num_qa_blocks=len(valid_qa_blocks),
        system_prompt_length=len(system_prompt),
    )

    result = None

    for attempt in range(etl_config.max_retries):
        try:
            response = await complete_with_tools(
                messages=messages,
                tools=[prompt_data["tool_definition"]],
                context=context,
                llm_params={
                    "model": etl_config.get_model("theme_grouping"),
                    "temperature": etl_config.temperature,
                    "max_tokens": etl_config.get_max_tokens("theme_grouping"),
                },
            )

            if response:
                tool_calls = (
                    response.get("choices", [{}])[0].get("message", {}).get("tool_calls", [])
                )

                if tool_calls:
                    raw_args = tool_calls[0]["function"]["arguments"]
                    if isinstance(raw_args, dict):
                        raw_data = raw_args
                    else:
                        raw_data = json.loads(str(raw_args).strip())

                    result = ThemeGroupingResponse.model_validate(raw_data)

                    # Log and accumulate LLM usage
                    metrics = response.get("metrics", {})
                    _accumulate_llm_cost(context, metrics)
                    logger.info(
                        "etl.key_themes.llm_usage",
                        execution_id=execution_id,
                        stage="grouping",
                        prompt_tokens=metrics.get("prompt_tokens", 0),
                        completion_tokens=metrics.get("completion_tokens", 0),
                        total_cost=metrics.get("total_cost", 0),
                        response_time=metrics.get("response_time", 0),
                    )

                    logger.info(
                        "regrouping.parsed_result",
                        execution_id=execution_id,
                        num_groups=len(result.theme_groups),
                    )
                    break

        except (KeyError, IndexError, json.JSONDecodeError, TypeError, ValidationError) as e:
            logger.warning(
                "regrouping.parse_error",
                execution_id=execution_id,
                error=str(e),
                attempt=attempt + 1,
            )
            if attempt < etl_config.max_retries - 1:
                continue
            raise RuntimeError("Failed to parse theme regrouping after multiple retries.") from e
        except Exception as e:
            logger.error(
                "regrouping.unexpected_error",
                execution_id=execution_id,
                error=str(e),
                error_type=type(e).__name__,
                attempt=attempt + 1,
            )
            if attempt < etl_config.max_retries - 1:
                delay = min(etl_config.retry_base_delay * (2**attempt), etl_config.retry_max_delay)
                delay += random.uniform(0, 0.5 * delay)
                await asyncio.sleep(delay)
                continue
            raise RuntimeError("Failed to complete theme regrouping after multiple retries.") from e

    if result:
        theme_groups = []
        for group_item in result.theme_groups:
            group = ThemeGroup(
                group_title=group_item.group_title or "Theme Group",
                qa_ids=group_item.qa_ids,
                rationale=group_item.rationale or "Regrouped by category",
            )
            theme_groups.append(group)

        return theme_groups

    raise RuntimeError(
        "Theme grouping failed: LLM returned no tool calls after all retries"
    )


def validate_grouping_assignments(
    qa_index: Dict[str, QABlock], theme_groups: List[ThemeGroup], execution_id: str
):
    """
    Ensure regrouping output covers every valid QA exactly once.
    """
    valid_ids = {qa_id for qa_id, qa_block in qa_index.items() if qa_block.is_valid}
    provided_ids = []

    for group in theme_groups:
        provided_ids.extend(group.qa_ids)

    provided_set = set(provided_ids)
    duplicates = [qa_id for qa_id, count in Counter(provided_ids).items() if count > 1]
    missing_ids = sorted(valid_ids - provided_set)
    unknown_ids = sorted(provided_set - valid_ids)

    if duplicates or missing_ids or unknown_ids:
        logger.error(
            "regrouping.validation_failed",
            execution_id=execution_id,
            duplicates=duplicates or None,
            missing_ids=missing_ids or None,
            unknown_ids=unknown_ids or None,
            expected_total=len(valid_ids),
            provided_total=len(provided_ids),
        )
        raise ValueError(
            "theme_grouping output failed validation; check logs for duplicates or missing IDs."
        )


def apply_grouping_to_index(qa_index: Dict[str, QABlock], theme_groups: List[ThemeGroup]):
    """
    Step 4: Apply grouping decisions to the Q&A index.
    """

    for qa_block in qa_index.values():
        qa_block.assigned_group = None

    for group in theme_groups:
        for qa_id in group.qa_ids:
            if qa_id in qa_index:
                qa_block = qa_index[qa_id]
                qa_block.assigned_group = group
                group.qa_blocks.append(qa_block)


def create_document(
    theme_groups: List[ThemeGroup],
    bank_name: str,
    bank_symbol: str,
    fiscal_year: int,
    quarter: str,
    output_path: str,
):
    """
    Step 5: Create Word document with grouped themes matching call summary style.
    """
    doc = Document()

    resolved_symbol = bank_symbol or (bank_name.split()[0] if bank_name else "BANK")

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.5)
        section.gutter = Inches(0)

    add_page_numbers_with_footer(doc, resolved_symbol, quarter, fiscal_year)

    config_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config")
    add_banner_image(doc, config_dir)

    for i, group in enumerate(theme_groups, 1):

        add_theme_header_with_background(doc, i, group.group_title)

        sorted_blocks = sorted(group.qa_blocks, key=lambda x: x.position)

        for j, qa_block in enumerate(sorted_blocks, 1):

            conv_para = doc.add_paragraph()
            conv_para.paragraph_format.space_before = Pt(6)
            conv_para.paragraph_format.space_after = Pt(4)

            conv_text = f"Conversation {j}:"
            conv_run = conv_para.add_run(conv_text)
            conv_run.font.underline = True
            conv_run.font.size = Pt(11)
            conv_run.font.color.rgb = RGBColor(0, 0, 0)

            # Apply auto_bold_html_metrics to formatted content
            content = qa_block.formatted_content or qa_block.original_content
            content = auto_bold_html_metrics(content)

            for line in content.split("\n"):
                if line.strip():
                    if line.strip() in ["---", "***", "___", "<hr>", "<hr/>", "<hr />"]:
                        continue

                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.15

                    parser = HTMLToDocx(p, font_size=Pt(9))
                    parser.feed(line.strip())
                    parser.close()

                    if not p.runs:
                        run = p.add_run(line.strip())
                        run.font.size = Pt(9)

            if j < len(sorted_blocks):
                separator = doc.add_paragraph()
                separator.paragraph_format.left_indent = Inches(0.3)
                separator.paragraph_format.space_before = Pt(6)
                separator.paragraph_format.space_after = Pt(6)

                separator_run = separator.add_run("_" * 50)
                separator_run.font.size = Pt(8)
                separator_run.font.color.rgb = RGBColor(200, 200, 200)

        if i < len(theme_groups):
            spacing = doc.add_paragraph()
            spacing.paragraph_format.space_before = Pt(12)
            spacing.paragraph_format.space_after = Pt(12)

    validate_document_content(doc)
    doc.save(output_path)


async def _save_to_database(
    theme_groups: List[ThemeGroup],
    qa_index: Dict[str, QABlock],
    filepath: str,
    docx_filename: str,
    etl_context: dict,
) -> None:
    """
    Save report metadata to database.

    Args:
        theme_groups: Final grouped themes
        qa_index: All Q&A blocks (for invalid count)
        filepath: Local file path
        docx_filename: Document filename
        etl_context: Dict with keys: bank_info, quarter, fiscal_year, execution_id,
    """
    bank_info = etl_context["bank_info"]
    quarter = etl_context["quarter"]
    fiscal_year = etl_context["fiscal_year"]
    execution_id = etl_context["execution_id"]

    report_metadata = get_standard_report_metadata()
    generation_timestamp = datetime.now()

    stage = "connecting"
    try:
        async with get_connection() as conn:
            stage = "deleting existing report"
            deleted = await conn.execute(
                text(
                    """
                DELETE FROM aegis_reports
                WHERE bank_id = :bank_id
                  AND fiscal_year = :fiscal_year
                  AND quarter = :quarter
                  AND report_type = :report_type
                RETURNING id
                """
                ),
                {
                    "bank_id": bank_info["bank_id"],
                    "fiscal_year": fiscal_year,
                    "quarter": quarter,
                    "report_type": report_metadata["report_type"],
                },
            )
            deleted.fetchall()

            stage = "inserting new report"
            result = await conn.execute(
                text(
                    """
                INSERT INTO aegis_reports (
                    report_name,
                    report_description,
                    report_type,
                    bank_id,
                    bank_name,
                    bank_symbol,
                    fiscal_year,
                    quarter,
                    local_filepath,
                    s3_document_name,
                    s3_pdf_name,
                    generation_date,
                    generated_by,
                    execution_id,
                    metadata
                ) VALUES (
                    :report_name,
                    :report_description,
                    :report_type,
                    :bank_id,
                    :bank_name,
                    :bank_symbol,
                    :fiscal_year,
                    :quarter,
                    :local_filepath,
                    :s3_document_name,
                    :s3_pdf_name,
                    :generation_date,
                    :generated_by,
                    :execution_id,
                    :metadata
                )
                RETURNING id
                """
                ),
                {
                    "report_name": report_metadata["report_name"],
                    "report_description": report_metadata["report_description"],
                    "report_type": report_metadata["report_type"],
                    "bank_id": bank_info["bank_id"],
                    "bank_name": bank_info["bank_name"],
                    "bank_symbol": bank_info["bank_symbol"],
                    "fiscal_year": fiscal_year,
                    "quarter": quarter,
                    "local_filepath": filepath,
                    "s3_document_name": docx_filename,
                    "s3_pdf_name": None,
                    "generation_date": generation_timestamp,
                    "generated_by": "key_themes_etl",
                    "execution_id": execution_id,
                    "metadata": json.dumps(
                        {
                            "theme_groups": len(theme_groups),
                            "total_qa_blocks": sum(
                                len(group.qa_blocks) for group in theme_groups
                            ),
                            "invalid_qa_filtered": sum(
                                1 for qa in qa_index.values() if not qa.is_valid
                            ),
                        }
                    ),
                },
            )
            result.fetchone()

            await conn.commit()

    except SQLAlchemyError as e:
        logger.error(
            "etl.key_themes.database_error",
            execution_id=execution_id,
            stage=stage,
            filepath=filepath,
            error=str(e),
        )
        raise


async def generate_key_themes(bank_name: str, fiscal_year: int, quarter: str) -> KeyThemesResult:
    """
    Generate key themes report from earnings call Q&A.

    Args:
        bank_name: ID, name, or symbol of the bank
        fiscal_year: Year (e.g., 2024)
        quarter: Quarter (e.g., "Q3")

    Returns:
        KeyThemesResult with filepath and counts

    Raises:
        KeyThemesUserError: For expected errors (bad input, no data)
        KeyThemesSystemError: For unexpected system/infrastructure errors
    """
    marks = [("start", time.monotonic())]
    execution_id = str(uuid.uuid4())
    logger.info(
        "etl.key_themes.started",
        execution_id=execution_id,
        bank_name=bank_name,
        fiscal_year=fiscal_year,
        quarter=quarter,
    )

    try:
        bank_info = get_bank_info_from_config(bank_name)

        await verify_and_get_availability(
            bank_info["bank_id"], bank_info["bank_name"], fiscal_year, quarter
        )

        ssl_config = setup_ssl()
        auth_config = await setup_authentication(execution_id, ssl_config)

        if not auth_config["success"]:
            error_msg = f"Authentication failed: {auth_config.get('error', 'Unknown error')}"
            logger.error("etl.key_themes.auth_failed", execution_id=execution_id, error=error_msg)
            raise KeyThemesSystemError(error_msg)

        context = {
            "execution_id": execution_id,
            "ssl_config": ssl_config,
            "auth_config": auth_config,
            "bank_name": bank_info["bank_name"],
            "bank_symbol": bank_info["bank_symbol"],
            "quarter": quarter,
            "fiscal_year": fiscal_year,
        }

        marks.append(("setup", time.monotonic()))

        # Load categories from xlsx
        categories = load_categories_from_xlsx(execution_id)

        qa_index = await load_qa_blocks(bank_info["bank_name"], fiscal_year, quarter, context)

        if not qa_index:
            raise KeyThemesUserError(
                f"No Q&A data found for {bank_info['bank_name']} {quarter} {fiscal_year}"
            )

        marks.append(("retrieval", time.monotonic()))

        # Load all prompts once (matches Call Summary pattern)
        extraction_prompts = load_prompt_from_db(
            layer="key_themes_etl",
            name="theme_extraction",
            compose_with_globals=False,
            available_databases=None,
            execution_id=execution_id,
        )
        formatting_prompts = load_prompt_from_db(
            layer="key_themes_etl",
            name="html_formatting",
            compose_with_globals=False,
            available_databases=None,
            execution_id=execution_id,
        )
        grouping_prompts = load_prompt_from_db(
            layer="key_themes_etl",
            name="theme_grouping",
            compose_with_globals=False,
            available_databases=None,
            execution_id=execution_id,
        )

        # Stage 1: Sequential classification with cumulative context
        await classify_all_qa_blocks_sequential(
            qa_index, categories, extraction_prompts, context
        )

        marks.append(("classification", time.monotonic()))

        # Stage 2: Parallel HTML formatting
        await format_all_qa_blocks_parallel(qa_index, formatting_prompts, context)

        marks.append(("formatting", time.monotonic()))

        # Stage 3: Review classifications, regroup if needed, generate final titles
        theme_groups = await determine_comprehensive_grouping(
            qa_index, categories, grouping_prompts, context
        )

        validate_grouping_assignments(qa_index, theme_groups, execution_id)

        logger.info(
            "etl.key_themes.grouping_completed",
            execution_id=execution_id,
            num_groups=len(theme_groups),
        )

        marks.append(("grouping", time.monotonic()))

        apply_grouping_to_index(qa_index, theme_groups)

        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
        os.makedirs(output_dir, exist_ok=True)

        filename_base = f"{bank_info['bank_symbol']}_{fiscal_year}_{quarter}"
        docx_filename = f"{filename_base}.docx"
        filepath = os.path.join(output_dir, docx_filename)

        create_document(
            theme_groups,
            bank_info["bank_name"],
            bank_info["bank_symbol"],
            fiscal_year,
            quarter,
            filepath,
        )

        marks.append(("document", time.monotonic()))

        logger.info("etl.key_themes.document_saved", execution_id=execution_id, filepath=filepath)

        await _save_to_database(
            theme_groups=theme_groups,
            qa_index=qa_index,
            filepath=filepath,
            docx_filename=docx_filename,
            etl_context={
                "bank_info": bank_info,
                "quarter": quarter,
                "fiscal_year": fiscal_year,
                "execution_id": execution_id,
            },
        )

        marks.append(("save", time.monotonic()))

        total_qa = sum(len(group.qa_blocks) for group in theme_groups)
        invalid_qa = sum(1 for qa in qa_index.values() if not qa.is_valid)
        cost_summary = _get_total_llm_cost(context)
        logger.info(
            "etl.key_themes.completed",
            execution_id=execution_id,
            theme_groups=len(theme_groups),
            valid_qa=total_qa,
            invalid_qa_filtered=invalid_qa,
            filepath=filepath,
            llm_calls=cost_summary["llm_calls"],
            total_tokens=cost_summary["total_tokens"],
            total_cost=cost_summary["total_cost"],
            **_timing_summary(marks),
        )

        return KeyThemesResult(
            filepath=filepath,
            theme_groups=len(theme_groups),
            valid_qa=total_qa,
            invalid_qa_filtered=invalid_qa,
            total_cost=cost_summary["total_cost"],
            total_tokens=cost_summary["total_tokens"],
        )

    except KeyThemesError:
        raise
    except (ValueError, RuntimeError) as e:
        logger.error("etl.key_themes.error", execution_id=execution_id, error=str(e))
        raise KeyThemesUserError(str(e)) from e
    except Exception as e:
        error_msg = f"Error generating key themes report: {str(e)}"
        logger.error(
            "etl.key_themes.error", execution_id=execution_id, error=error_msg, exc_info=True
        )
        raise KeyThemesSystemError(error_msg) from e


def main():
    """Main entry point for command-line execution."""
    parser = argparse.ArgumentParser(
        description="Generate key themes report from earnings call Q&A",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument("--bank", required=True, help="Bank ID, name, or symbol")
    parser.add_argument("--year", type=int, required=True, help="Fiscal year")
    parser.add_argument(
        "--quarter", required=True, choices=["Q1", "Q2", "Q3", "Q4"], help="Quarter"
    )

    args = parser.parse_args()

    postgresql_prompts()

    print(f"\n🔄 Generating key themes report for {args.bank} {args.quarter} {args.year}...\n")

    try:
        result = asyncio.run(
            generate_key_themes(bank_name=args.bank, fiscal_year=args.year, quarter=args.quarter)
        )
        print(
            f"✅ Complete: {result.filepath}\n"
            f"   Theme groups: {result.theme_groups}, "
            f"Valid Q&A: {result.valid_qa}, Filtered: {result.invalid_qa_filtered}\n"
            f"   LLM cost: ${result.total_cost:.4f}, Tokens: {result.total_tokens:,}"
        )
    except KeyThemesUserError as e:
        print(f"⚠️ {e}", file=sys.stderr)
        sys.exit(1)
    except KeyThemesError as e:
        print(f"❌ {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
