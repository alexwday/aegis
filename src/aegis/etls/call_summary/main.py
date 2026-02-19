"""
Call Summary ETL Script.

Usage:
    python -m aegis.etls.call_summary.main --bank "Royal Bank of Canada" --year 2024 --quarter Q3
    python -m aegis.etls.call_summary.main --bank RY --year 2024 --quarter Q3
    python -m aegis.etls.call_summary.main --bank 1 --year 2024 --quarter Q3
"""

import argparse
import asyncio
import json
import random
import time
import uuid
import os
import sys
from dataclasses import dataclass
from datetime import datetime
from functools import lru_cache
from itertools import groupby
from typing import Dict, Any, List, Optional
from pydantic import BaseModel, Field, ValidationError
from sqlalchemy import text
from sqlalchemy.exc import SQLAlchemyError
import pandas as pd
import yaml
from docx import Document
from docx.shared import Pt

from aegis.etls.call_summary.document_converter import (
    get_standard_report_metadata,
    setup_document_formatting,
    add_banner_image,
    add_document_title,
    add_section_heading,
    add_table_of_contents,
    mark_document_for_update,
    add_structured_content_to_doc,
    validate_document_content,
)
from aegis.etls.call_summary.transcript_utils import (
    retrieve_full_section,
    format_full_section_chunks,
    SECTION_MD,
    SECTION_QA,
    SECTIONS_KEY_MD,
    SECTIONS_KEY_QA,
    SECTIONS_KEY_ALL,
    VALID_SECTION_KEYS,
)
from aegis.utils.ssl import setup_ssl
from aegis.connections.oauth_connector import setup_authentication
from aegis.connections.llm_connector import complete_with_tools
from aegis.connections.postgres_connector import get_connection
from aegis.utils.logging import setup_logging, get_logger
from aegis.utils.prompt_loader import load_prompt_from_db
from aegis.utils.sql_prompt import postgresql_prompts
from aegis.utils.settings import config

setup_logging()
logger = get_logger()

# --- Retry Configuration ---
MAX_RETRIES = 3
RETRY_BASE_DELAY = 1.0
RETRY_MAX_DELAY = 10.0

# --- Display / Logging Configuration ---
LOG_SNIPPET_LENGTH = 80

# --- Concurrency Configuration ---
MAX_CONCURRENT_EXTRACTIONS = 5


# --- ETL Exception Hierarchy ---


class CallSummaryError(Exception):
    """Base exception for call summary ETL errors."""


class CallSummarySystemError(CallSummaryError):
    """Unexpected system/infrastructure error."""


class CallSummaryUserError(CallSummaryError):
    """Expected user-facing error (bad input, no data, etc.)."""


@dataclass
class CallSummaryResult:
    """Successful call summary generation result."""

    filepath: str
    total_categories: int
    included_categories: int
    total_cost: float = 0.0
    total_tokens: int = 0


class ETLConfig:
    """
    ETL configuration loader that reads YAML configs and resolves model references.

    This class loads ETL-specific configuration from YAML files and provides
    easy access to configuration values with automatic model tier resolution.
    """

    def __init__(self, config_path: str):
        """Initialize the ETL configuration loader."""
        self.config_path = config_path
        self._config = self._load_config()

    def _load_config(self) -> Dict[str, Any]:
        """Load configuration from YAML file."""
        if not os.path.exists(self.config_path):
            raise FileNotFoundError(f"Configuration file not found: {self.config_path}")

        with open(self.config_path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f)

    def get_model(self, model_key: str) -> str:
        """
        Get the actual model name for a given model key.

        Resolves model tier references (small/medium/large) to actual model names
        from the global settings configuration.
        """
        if "models" not in self._config or model_key not in self._config["models"]:
            raise KeyError(f"Model key '{model_key}' not found in configuration")

        tier = self._config["models"][model_key].get("tier")
        if not tier:
            raise ValueError(f"No tier specified for model '{model_key}'")

        # Resolve tier to actual model from global config
        tier_map = {
            "small": config.llm.small.model,
            "medium": config.llm.medium.model,
            "large": config.llm.large.model,
        }

        if tier not in tier_map:
            raise ValueError(
                f"Invalid tier '{tier}' for model '{model_key}'. "
                f"Valid tiers: {list(tier_map.keys())}"
            )

        return tier_map[tier]

    @property
    def temperature(self) -> float:
        """Get the LLM temperature parameter."""
        return self._config.get("llm", {}).get("temperature", 0.1)

    def get_max_tokens(self, task_key: str) -> int:
        """
        Get max_tokens for a specific task, falling back to default.

        Supports both legacy (flat int) and per-task (dict) config formats.

        Args:
            task_key: Task identifier (e.g., "research_plan", "category_extraction")

        Returns:
            max_tokens value for the task
        """
        max_tokens_config = self._config.get("llm", {}).get("max_tokens", {})
        if isinstance(max_tokens_config, int):
            return max_tokens_config
        return max_tokens_config.get(task_key, max_tokens_config.get("default", 32768))

    @property
    def max_concurrent_extractions(self) -> int:
        """Get the maximum number of concurrent category extractions."""
        return self._config.get("concurrency", {}).get(
            "max_concurrent_extractions", MAX_CONCURRENT_EXTRACTIONS
        )

    @property
    def max_retries(self) -> int:
        """Get the maximum number of LLM call retries."""
        return self._config.get("retry", {}).get("max_retries", MAX_RETRIES)

    @property
    def retry_base_delay(self) -> float:
        """Get the base delay in seconds for retry backoff."""
        return self._config.get("retry", {}).get("base_delay", RETRY_BASE_DELAY)

    @property
    def retry_max_delay(self) -> float:
        """Get the maximum delay in seconds for retry backoff."""
        return self._config.get("retry", {}).get("max_delay", RETRY_MAX_DELAY)


etl_config = ETLConfig(os.path.join(os.path.dirname(__file__), "config", "config.yaml"))


# --- Pydantic Models for LLM Response Validation ---


class CategoryPlan(BaseModel):
    """A single category's research plan from the LLM."""

    index: int
    name: str
    extraction_strategy: str
    cross_category_notes: str = ""
    relevant_qa_groups: List[int] = Field(default_factory=list)


class ResearchPlanResponse(BaseModel):
    """Top-level research plan response from the LLM."""

    category_plans: List[CategoryPlan]


class Evidence(BaseModel):
    """A single piece of supporting evidence for a statement."""

    type: str = "quote"
    content: str = ""
    speaker: str = ""


class SummaryStatement(BaseModel):
    """A single summary statement with supporting evidence."""

    statement: str
    evidence: List[Evidence] = Field(default_factory=list)


class CategoryExtractionResponse(BaseModel):
    """Category extraction response from the LLM."""

    reasoning: Optional[str] = None
    rejected: bool
    rejection_reason: Optional[str] = None
    title: Optional[str] = None
    summary_statements: List[SummaryStatement] = Field(default_factory=list)


class DuplicateStatement(BaseModel):
    """A statement removal instruction from the LLM dedup pass."""

    category_index: int
    statement_index: int
    duplicate_of_category_index: int
    duplicate_of_statement_index: int
    reasoning: str = ""


class DuplicateEvidence(BaseModel):
    """An evidence removal instruction from the LLM dedup pass."""

    category_index: int
    statement_index: int
    evidence_index: int
    duplicate_of_category_index: int
    duplicate_of_statement_index: int
    duplicate_of_evidence_index: int
    reasoning: str = ""


class DeduplicationResponse(BaseModel):
    """Deduplication response from the LLM."""

    analysis_notes: str = ""
    duplicate_statements: List[DuplicateStatement] = Field(default_factory=list)
    duplicate_evidence: List[DuplicateEvidence] = Field(default_factory=list)


@lru_cache(maxsize=1)
def _load_monitored_institutions() -> Dict[int, Dict[str, Any]]:
    """
    Load and cache monitored institutions configuration.

    Uses @lru_cache for thread-safe, testable caching (clear via .cache_clear()).

    Returns:
        Dictionary mapping bank_id to institution details (id, name, symbol, type, path_safe_name)
    """
    config_path = os.path.join(os.path.dirname(__file__), "config", "monitored_institutions.yaml")
    with open(config_path, "r", encoding="utf-8") as f:
        yaml_data = yaml.safe_load(f)

    institutions = {}
    for key, value in yaml_data.items():
        symbol = key.split("-")[0]  # Extract symbol from "RY-CA" -> "RY"
        institutions[value["id"]] = {**value, "symbol": symbol}
    return institutions


def _sanitize_for_prompt(text: str) -> str:
    """
    Escape curly braces in text for safe use in .format() templates.

    Prevents KeyError/IndexError when XLSX-sourced content contains { or }.

    Args:
        text: Raw text string

    Returns:
        Text with { and } escaped as {{ and }}
    """
    return text.replace("{", "{{").replace("}", "}}")


def _accumulate_llm_cost(context: Dict[str, Any], metrics: Dict[str, Any]) -> None:
    """
    Accumulate LLM cost and token usage from a response's metrics dict.

    Appends to context["_llm_costs"] list for later aggregation.
    Safe for asyncio concurrency (single-threaded event loop).

    Args:
        context: Execution context with _llm_costs list
        metrics: Response metrics dict with prompt_tokens, completion_tokens, total_cost
    """
    if "_llm_costs" not in context:
        context["_llm_costs"] = []
    context["_llm_costs"].append(
        {
            "prompt_tokens": metrics.get("prompt_tokens", 0),
            "completion_tokens": metrics.get("completion_tokens", 0),
            "total_cost": metrics.get("total_cost", 0),
        }
    )


def _get_total_llm_cost(context: Dict[str, Any]) -> Dict[str, Any]:
    """
    Aggregate all accumulated LLM costs from context.

    Returns:
        Dict with total_cost, total_prompt_tokens, total_completion_tokens, total_tokens
    """
    costs = context.get("_llm_costs", [])
    total_cost = sum(c.get("total_cost", 0) for c in costs)
    prompt_tokens = sum(c.get("prompt_tokens", 0) for c in costs)
    completion_tokens = sum(c.get("completion_tokens", 0) for c in costs)
    return {
        "total_cost": round(total_cost, 4),
        "total_prompt_tokens": prompt_tokens,
        "total_completion_tokens": completion_tokens,
        "total_tokens": prompt_tokens + completion_tokens,
        "llm_calls": len(costs),
    }


def format_categories_for_prompt(categories: List[Dict[str, Any]]) -> str:
    """
    Format category dictionaries into standardized XML format for prompt injection.

    This is the standardized formatting function used across all ETLs (Call Summary,
    Key Themes, CM Readthrough) to ensure consistent category presentation to LLMs.

    Args:
        categories: List of category dicts with standardized 6-column format

    Returns:
        Formatted XML string with category information
    """
    formatted_sections = []

    for cat in categories:
        # Map transcript_sections to human-readable description
        section_desc = {
            SECTIONS_KEY_MD: "Management Discussion section only",
            SECTIONS_KEY_QA: "Q&A section only",
            SECTIONS_KEY_ALL: "Both Management Discussion and Q&A sections",
        }.get(cat.get("transcript_sections", SECTIONS_KEY_ALL), "ALL sections")

        section = "<category>\n"
        section += f"<name>{_sanitize_for_prompt(cat['category_name'])}</name>\n"
        section += f"<section>{section_desc}</section>\n"
        section += (
            f"<description>{_sanitize_for_prompt(cat['category_description'])}</description>\n"
        )

        # Collect non-empty examples
        examples = []
        for i in range(1, 4):
            example_key = f"example_{i}"
            if (
                cat.get(example_key)
                and isinstance(cat[example_key], str)
                and cat[example_key].strip()
            ):
                examples.append(cat[example_key])

        if examples:
            section += "<examples>\n"
            for example in examples:
                section += f"  <example>{_sanitize_for_prompt(example)}</example>\n"
            section += "</examples>\n"

        section += "</category>"
        formatted_sections.append(section)

    return "\n\n".join(formatted_sections)


def load_categories_from_xlsx(bank_type: str, execution_id: str) -> List[Dict[str, Any]]:
    """
    Load categories from the appropriate XLSX file based on bank type.

    Args:
        bank_type: Either "Canadian_Banks" or "US_Banks"
        execution_id: Execution ID for logging

    Returns:
        List of dictionaries with transcript_sections, category_name, category_description,
        example_1, example_2, example_3
    """
    file_name = (
        "canadian_banks_categories.xlsx"
        if bank_type == "Canadian_Banks"
        else "us_banks_categories.xlsx"
    )

    current_dir = os.path.dirname(os.path.abspath(__file__))
    xlsx_path = os.path.join(current_dir, "config", "categories", file_name)

    if not os.path.exists(xlsx_path):
        raise FileNotFoundError(f"Categories file not found: {xlsx_path}")

    try:
        df = pd.read_excel(xlsx_path, sheet_name=0)

        # Required columns for standard format
        required_columns = ["transcript_sections", "category_name", "category_description"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"Missing required columns in {file_name}: {missing_columns}")

        # Optional columns with defaults for backward compatibility
        optional_columns = ["example_1", "example_2", "example_3"]
        for col in optional_columns:
            if col not in df.columns:
                df[col] = ""  # Add empty column if not present

        # Ensure report_section column exists even for legacy sheets
        if "report_section" not in df.columns:
            df["report_section"] = "Results Summary"

        # Convert to list of dicts, ensuring all required fields are non-empty
        categories = []
        for idx, row in df.iterrows():
            for field in required_columns:
                if pd.isna(row[field]) or str(row[field]).strip() == "":
                    raise ValueError(f"Missing value for '{field}' in {file_name} (row {idx + 2})")

            transcript_sections = str(row["transcript_sections"]).strip()
            valid_sections = VALID_SECTION_KEYS
            if transcript_sections not in valid_sections:
                raise ValueError(
                    f"Invalid transcript_sections '{transcript_sections}' "
                    f"in {file_name} (row {idx + 2}). Must be one of: {valid_sections}"
                )

            category = {
                "transcript_sections": transcript_sections,
                "report_section": (
                    str(row["report_section"]).strip()
                    if pd.notna(row["report_section"])
                    else "Results Summary"
                ),
                "category_name": str(row["category_name"]).strip(),
                "category_description": str(row["category_description"]).strip(),
                "example_1": str(row["example_1"]).strip() if pd.notna(row["example_1"]) else "",
                "example_2": str(row["example_2"]).strip() if pd.notna(row["example_2"]) else "",
                "example_3": str(row["example_3"]).strip() if pd.notna(row["example_3"]) else "",
            }
            categories.append(category)

        if not categories:
            raise ValueError(f"No categories in {file_name}")

        logger.info(
            "etl.call_summary.categories_loaded",
            execution_id=execution_id,
            file_name=file_name,
            num_categories=len(categories),
        )
        return categories

    except Exception as e:
        error_msg = f"Failed to load categories from {xlsx_path}: {str(e)}"
        logger.error(
            "etl.call_summary.categories_load_error",
            execution_id=execution_id,
            xlsx_path=xlsx_path,
            error=str(e),
        )
        raise RuntimeError(error_msg) from e


def get_bank_info_from_config(bank_identifier: str) -> Dict[str, Any]:
    """
    Look up bank from monitored institutions configuration file.

    Args:
        bank_identifier: Bank ID (as string/int), symbol (e.g., "RY"), or name

    Returns:
        Dictionary with bank_id, bank_name, bank_symbol, bank_type

    Raises:
        ValueError: If bank not found in monitored institutions
    """
    institutions = _load_monitored_institutions()
    bank_identifier = bank_identifier.strip()

    def _to_bank_info(inst: Dict[str, Any]) -> Dict[str, Any]:
        return {
            "bank_id": inst["id"],
            "bank_name": inst["name"],
            "bank_symbol": inst["symbol"],
            "bank_type": inst["type"],
        }

    # Try lookup by ID
    if bank_identifier.isdigit():
        bank_id = int(bank_identifier)
        if bank_id in institutions:
            return _to_bank_info(institutions[bank_id])

    # Try lookup by symbol first (supports full ticker input, e.g. "C-US")
    bank_identifier_upper = bank_identifier.upper()
    bank_identifier_lower = bank_identifier.lower()
    symbol_candidate = bank_identifier_upper.split("-")[0]

    for inst in institutions.values():
        if inst["symbol"].upper() == symbol_candidate:
            return _to_bank_info(inst)

    # Fallback: match by name (case-insensitive, partial match)
    for inst in institutions.values():
        if bank_identifier_lower in inst["name"].lower():
            return _to_bank_info(inst)

    # Build helpful error message with available banks
    available = [f"{inst['symbol']} ({inst['name']})" for inst in institutions.values()]
    raise ValueError(
        f"Bank '{bank_identifier}' not found in monitored institutions.\n"
        f"Available banks: {', '.join(sorted(available))}"
    )


async def verify_and_get_availability(
    bank_id: int, bank_name: str, fiscal_year: int, quarter: str
) -> None:
    """
    Verify transcript data is available for the specified bank and period.

    Args:
        bank_id: Bank ID
        bank_name: Bank name (for error messages)
        fiscal_year: Year (e.g., 2024)
        quarter: Quarter (e.g., "Q3")

    Raises:
        ValueError: If transcript data not available (includes available periods)
    """
    async with get_connection() as conn:
        result = await conn.execute(
            text(
                """
                SELECT database_names
                FROM aegis_data_availability
                WHERE bank_id = :bank_id
                  AND fiscal_year = :fiscal_year
                  AND quarter = :quarter
                """
            ),
            {"bank_id": bank_id, "fiscal_year": fiscal_year, "quarter": quarter},
        )
        row = result.fetchone()

        if row and row[0] and "transcripts" in row[0]:
            return

        # Query available transcript periods for this bank to provide helpful context
        available_result = await conn.execute(
            text(
                """
                SELECT fiscal_year, quarter
                FROM aegis_data_availability
                WHERE bank_id = :bank_id
                  AND 'transcripts' = ANY(database_names)
                ORDER BY fiscal_year DESC, quarter DESC
                """
            ),
            {"bank_id": bank_id},
        )
        available_periods = [f"{r[1]} {r[0]}" for r in available_result.fetchall()]

        if available_periods:
            raise ValueError(
                f"No transcript data available for {bank_name} {quarter} {fiscal_year}. "
                f"Available periods: {', '.join(available_periods)}"
            )
        raise ValueError(
            f"No transcript data available for {bank_name} {quarter} {fiscal_year}. "
            f"No transcript data found for this bank in any period."
        )


async def _generate_research_plan(
    context: dict, research_prompts: dict, transcript_text: str, execution_id: str
) -> dict:
    """Generate research plan using LLM."""
    system_prompt = research_prompts["system_prompt"]
    user_prompt = research_prompts["user_prompt_template"].format(
        transcript_text=_sanitize_for_prompt(transcript_text)
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    max_retries = etl_config.max_retries
    for attempt in range(max_retries):
        try:
            response = await complete_with_tools(
                messages=messages,
                tools=[research_prompts["tool_definition"]],
                context=context,
                llm_params={
                    "model": etl_config.get_model("research_plan"),
                    "temperature": etl_config.temperature,
                    "max_tokens": etl_config.get_max_tokens("research_plan"),
                },
            )

            tool_call = response["choices"][0]["message"]["tool_calls"][0]
            raw_data = json.loads(tool_call["function"]["arguments"])
            validated = ResearchPlanResponse.model_validate(raw_data)
            research_plan_data = validated.model_dump()

            metrics = response.get("metrics", {})
            _accumulate_llm_cost(context, metrics)
            logger.info(
                "etl.call_summary.llm_usage",
                execution_id=execution_id,
                stage="research_plan",
                prompt_tokens=metrics.get("prompt_tokens", 0),
                completion_tokens=metrics.get("completion_tokens", 0),
                total_cost=metrics.get("total_cost", 0),
                response_time=metrics.get("response_time", 0),
            )

            logger.info(
                "etl.call_summary.research_plan_generated",
                execution_id=execution_id,
                num_plans=len(research_plan_data["category_plans"]),
            )
            return research_plan_data

        except (KeyError, IndexError, json.JSONDecodeError, TypeError, ValidationError) as e:
            logger.warning(
                "etl.call_summary.research_plan_parse_error",
                execution_id=execution_id,
                error=str(e),
                attempt=attempt + 1,
            )
            if attempt < max_retries - 1:
                continue
            raise RuntimeError(
                f"Error generating research plan after {max_retries} attempts: {str(e)}"
            ) from e
        except Exception as e:  # Catch transport/LLM failures (httpx/OpenAI/etc.)
            logger.error(
                "etl.call_summary.research_plan_error",
                execution_id=execution_id,
                error=str(e),
                attempt=attempt + 1,
            )
            if attempt < max_retries - 1:
                delay = min(etl_config.retry_base_delay * (2**attempt), etl_config.retry_max_delay)
                delay += random.uniform(0, 0.5 * delay)
                await asyncio.sleep(delay)
                continue
            raise RuntimeError(
                f"Error generating research plan after {max_retries} attempts: {str(e)}"
            ) from e

    raise RuntimeError(f"Failed to generate research plan after {max_retries} attempts")


def _build_rejection_result(index: int, category: dict, reason: str) -> dict:
    """
    Build a standardized rejection result for a category.

    Args:
        index: 1-based category index
        category: Category configuration dict
        reason: Human-readable rejection reason

    Returns:
        Rejection result dict with index, name, report_section, rejected, rejection_reason
    """
    return {
        "index": index,
        "name": category["category_name"],
        "report_section": category.get("report_section", "Results Summary"),
        "rejected": True,
        "rejection_reason": reason,
    }


def _format_categories_for_dedup(category_results: list) -> str:
    """
    Format non-rejected category results as XML with explicit indices for dedup prompt.

    Args:
        category_results: List of category result dicts from extraction

    Returns:
        XML string with category/statement/evidence indices for LLM consumption
    """
    parts = []
    for cat_idx, result in enumerate(category_results):
        if result.get("rejected", False) or "summary_statements" not in result:
            continue

        cat_lines = [f'<category index="{cat_idx}" name="{_sanitize_for_prompt(result["name"])}">']

        for stmt_idx, stmt in enumerate(result.get("summary_statements", [])):
            stmt_text = _sanitize_for_prompt(stmt.get("statement", ""))
            cat_lines.append(f'  <statement index="{stmt_idx}">{stmt_text}</statement>')

            for ev_idx, ev in enumerate(stmt.get("evidence", [])):
                ev_content = _sanitize_for_prompt(ev.get("content", ""))
                cat_lines.append(f'    <evidence index="{ev_idx}">{ev_content}</evidence>')

        cat_lines.append("</category>")
        parts.append("\n".join(cat_lines))

    return "\n\n".join(parts)


def _apply_dedup_removals(
    category_results: list, dedup_response: DeduplicationResponse, execution_id: str
) -> tuple:
    """
    Apply validated removal instructions from the LLM dedup response.

    Processes removals in reverse index order to prevent index shift corruption.
    Validates all indices before removal and skips invalid ones.

    Args:
        category_results: List of category result dicts (modified in-place)
        dedup_response: Validated deduplication response from LLM
        execution_id: Execution ID for logging

    Returns:
        Tuple of (statements_removed, evidence_removed) counts
    """
    statements_removed = 0
    evidence_removed = 0

    # --- Remove duplicate evidence (reverse order to avoid index shift) ---
    sorted_evidence = sorted(
        dedup_response.duplicate_evidence,
        key=lambda e: (e.category_index, e.statement_index, e.evidence_index),
        reverse=True,
    )
    for ev in sorted_evidence:
        # Validate indices
        if ev.category_index < 0 or ev.category_index >= len(category_results):
            logger.debug(
                "etl.call_summary.dedup.invalid_category_index",
                execution_id=execution_id,
                category_index=ev.category_index,
            )
            continue

        result = category_results[ev.category_index]
        if result.get("rejected", False) or "summary_statements" not in result:
            continue

        stmts = result["summary_statements"]
        if ev.statement_index < 0 or ev.statement_index >= len(stmts):
            logger.debug(
                "etl.call_summary.dedup.invalid_statement_index",
                execution_id=execution_id,
                statement_index=ev.statement_index,
            )
            continue

        evidence_list = stmts[ev.statement_index].get("evidence", [])
        if ev.evidence_index < 0 or ev.evidence_index >= len(evidence_list):
            logger.debug(
                "etl.call_summary.dedup.invalid_evidence_index",
                execution_id=execution_id,
                evidence_index=ev.evidence_index,
            )
            continue

        removed_ev = evidence_list.pop(ev.evidence_index)
        evidence_removed += 1
        logger.debug(
            "etl.call_summary.dedup.evidence_removed",
            execution_id=execution_id,
            category=result["name"],
            snippet=removed_ev.get("content", "")[:LOG_SNIPPET_LENGTH],
            reasoning=ev.reasoning,
        )

    # --- Remove duplicate statements (reverse order to avoid index shift) ---
    sorted_statements = sorted(
        dedup_response.duplicate_statements,
        key=lambda s: (s.category_index, s.statement_index),
        reverse=True,
    )
    for stmt in sorted_statements:
        if stmt.category_index < 0 or stmt.category_index >= len(category_results):
            logger.debug(
                "etl.call_summary.dedup.invalid_category_index",
                execution_id=execution_id,
                category_index=stmt.category_index,
            )
            continue

        result = category_results[stmt.category_index]
        if result.get("rejected", False) or "summary_statements" not in result:
            continue

        stmts = result["summary_statements"]
        if stmt.statement_index < 0 or stmt.statement_index >= len(stmts):
            logger.debug(
                "etl.call_summary.dedup.invalid_statement_index",
                execution_id=execution_id,
                statement_index=stmt.statement_index,
            )
            continue

        removed_stmt = stmts.pop(stmt.statement_index)
        statements_removed += 1
        logger.debug(
            "etl.call_summary.dedup.statement_removed",
            execution_id=execution_id,
            category=result["name"],
            statement=removed_stmt.get("statement", "")[:LOG_SNIPPET_LENGTH],
            reasoning=stmt.reasoning,
        )

    return statements_removed, evidence_removed


async def _deduplicate_results_llm(
    category_results: list, dedup_prompts: dict, context: dict, execution_id: str
) -> list:
    """
    LLM-based post-extraction deduplication of statements and evidence.

    Sends all non-rejected category results to an LLM to identify semantic duplicates
    (same insight expressed differently) across categories. On failure, returns results
    unchanged with a warning log.

    Args:
        category_results: List of category result dicts from extraction
        dedup_prompts: Loaded prompt dict with system_prompt, user_prompt, tool_definition
        context: Auth/SSL context dict
        execution_id: Execution ID for logging

    Returns:
        Modified category_results with LLM-identified duplicates removed
    """
    # Skip if fewer than 2 non-rejected categories
    non_rejected = [r for r in category_results if not r.get("rejected", False)]
    if len(non_rejected) < 2:
        logger.info(
            "etl.call_summary.dedup.skipped",
            execution_id=execution_id,
            reason="fewer than 2 non-rejected categories",
        )
        return category_results

    formatted_categories = _format_categories_for_dedup(category_results)

    system_prompt = dedup_prompts["system_prompt"]
    user_prompt = dedup_prompts["user_prompt"].format(
        categories_xml=_sanitize_for_prompt(formatted_categories)
    )

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    max_retries = etl_config.max_retries
    for attempt in range(max_retries):
        try:
            response = await complete_with_tools(
                messages=messages,
                tools=[dedup_prompts["tool_definition"]],
                context=context,
                llm_params={
                    "model": etl_config.get_model("deduplication"),
                    "temperature": etl_config.temperature,
                    "max_tokens": etl_config.get_max_tokens("deduplication"),
                },
            )

            tool_call = response["choices"][0]["message"]["tool_calls"][0]
            raw_data = json.loads(tool_call["function"]["arguments"])
            dedup_response = DeduplicationResponse.model_validate(raw_data)

            metrics = response.get("metrics", {})
            _accumulate_llm_cost(context, metrics)
            logger.info(
                "etl.call_summary.llm_usage",
                execution_id=execution_id,
                stage="deduplication",
                prompt_tokens=metrics.get("prompt_tokens", 0),
                completion_tokens=metrics.get("completion_tokens", 0),
                total_cost=metrics.get("total_cost", 0),
                response_time=metrics.get("response_time", 0),
            )

            statements_removed, evidence_removed = _apply_dedup_removals(
                category_results, dedup_response, execution_id
            )

            if statements_removed > 0 or evidence_removed > 0:
                logger.info(
                    "etl.call_summary.dedup.summary",
                    execution_id=execution_id,
                    statements_removed=statements_removed,
                    evidence_removed=evidence_removed,
                )

            return category_results

        except (KeyError, IndexError, json.JSONDecodeError, TypeError, ValidationError) as e:
            logger.warning(
                "etl.call_summary.dedup_parse_error",
                execution_id=execution_id,
                error=str(e),
                attempt=attempt + 1,
            )
            if attempt < max_retries - 1:
                continue
            raise RuntimeError(
                f"Deduplication failed after {max_retries} attempts: {str(e)}"
            ) from e
        except Exception as e:  # Transport/LLM failures
            logger.error(
                "etl.call_summary.dedup_error",
                execution_id=execution_id,
                error=str(e),
                attempt=attempt + 1,
            )
            if attempt < max_retries - 1:
                delay = min(etl_config.retry_base_delay * (2**attempt), etl_config.retry_max_delay)
                delay += random.uniform(0, 0.5 * delay)
                await asyncio.sleep(delay)
                continue
            raise RuntimeError(
                f"Deduplication failed after {max_retries} attempts: {str(e)}"
            ) from e

    raise RuntimeError(f"Deduplication failed after {max_retries} attempts")


def _filter_chunks_for_category(
    section_cache: dict,
    transcript_sections: str,
    relevant_qa_groups: list,
) -> list:
    """
    Filter cached transcript chunks for a specific category.

    MD chunks are always included in full (no visible block IDs for filtering).
    QA chunks are filtered to only those whose qa_group_id is in relevant_qa_groups.
    If relevant_qa_groups is empty/None, all QA chunks are included as a safety fallback.

    Args:
        section_cache: Dict with keys "MD", "QA", "ALL" mapping to chunk lists
        transcript_sections: "MD", "QA", or "ALL" — which sections this category uses
        relevant_qa_groups: List of Q&A group IDs relevant to this category

    Returns:
        Filtered list of chunks for the category
    """
    md_chunks = section_cache.get(SECTIONS_KEY_MD, [])
    qa_chunks = section_cache.get(SECTIONS_KEY_QA, [])

    # Filter QA chunks if we have a non-empty relevance list
    if relevant_qa_groups:
        qa_group_set = set(relevant_qa_groups)
        filtered_qa = [c for c in qa_chunks if c.get("qa_group_id") in qa_group_set]
    else:
        # Safety fallback: no filtering when list is empty/None
        filtered_qa = qa_chunks

    # Combine based on transcript_sections
    if transcript_sections == SECTIONS_KEY_MD:
        return list(md_chunks)
    if transcript_sections == SECTIONS_KEY_QA:
        return list(filtered_qa)
    # SECTIONS_KEY_ALL or any other value
    return list(md_chunks) + list(filtered_qa)


async def _extract_single_category(  # noqa: E501  pylint: disable=too-many-arguments,too-many-positional-arguments,too-many-locals
    index: int,
    category: dict,
    research_plan_data: dict,
    extraction_prompts: dict,
    etl_context: dict,
    semaphore: asyncio.Semaphore,
    total_categories: int,
) -> dict:
    """
    Extract data for a single category (parallel-safe coroutine).

    Acquires the semaphore to limit concurrency, then performs chunk filtering,
    transcript formatting, and LLM extraction with retry logic. Never raises;
    returns a rejection result on unrecoverable errors.

    Args:
        index: 1-based category index
        category: Category configuration dict
        research_plan_data: Research plan from LLM
        extraction_prompts: Prompts for extraction
        etl_context: Shared read-only context dict
        semaphore: Concurrency-limiting semaphore
        total_categories: Total number of categories being processed

    Returns:
        Extraction result dict (accepted or rejected)
    """
    async with semaphore:
        retrieval_params = etl_context["retrieval_params"]
        bank_info = etl_context["bank_info"]
        quarter = etl_context["quarter"]
        fiscal_year = etl_context["fiscal_year"]
        context = etl_context["context"]
        execution_id = etl_context["execution_id"]

        category_plan = next(
            (p for p in research_plan_data["category_plans"] if p.get("index") == index),
            None,
        )

        if not category_plan:
            logger.info(
                "etl.call_summary.category_not_in_plan",
                execution_id=execution_id,
                category_name=category["category_name"],
                category_index=index,
                reason="Not in research plan — proceeding with full unfiltered transcript",
            )

        # Extract Q&A group filter from research plan (empty = no filtering)
        relevant_qa_groups = category_plan.get("relevant_qa_groups", []) if category_plan else []

        # Filter chunks to relevant content; no plan = full unfiltered transcript
        cache = etl_context.get("section_cache", {})
        chunks = _filter_chunks_for_category(
            section_cache=cache,
            transcript_sections=category["transcript_sections"],
            relevant_qa_groups=relevant_qa_groups,
        )

        # Fallback if section_cache missed this section type entirely
        if not chunks and not cache.get(category["transcript_sections"]):
            chunks = await retrieve_full_section(
                combo=retrieval_params,
                sections=category["transcript_sections"],
                context=context,
            )

        # Log chunk filtering results
        total_section_chunks = len(cache.get(category["transcript_sections"], []))
        logger.debug(
            "etl.call_summary.chunk_filtering",
            execution_id=execution_id,
            category_name=category["category_name"],
            transcript_sections=category["transcript_sections"],
            relevant_qa_groups=relevant_qa_groups,
            total_chunks=total_section_chunks,
            filtered_chunks=len(chunks),
        )

        if not chunks:
            return _build_rejection_result(
                index,
                category,
                f"No {category['transcript_sections']} section data available",
            )

        formatted_section = format_full_section_chunks(
            chunks=chunks, combo=retrieval_params, context=context
        )

        extraction_strategy = (
            category_plan["extraction_strategy"]
            if category_plan
            else "No research plan available — extract all relevant content from the transcript."
        )
        cross_category_notes = (
            category_plan.get("cross_category_notes", "") if category_plan else ""
        )

        system_prompt = extraction_prompts["system_prompt"].format(
            category_index=index,
            total_categories=total_categories,
            bank_name=bank_info["bank_name"],
            bank_symbol=bank_info["bank_symbol"],
            quarter=quarter,
            fiscal_year=fiscal_year,
            category_name=_sanitize_for_prompt(category["category_name"]),
            category_description=_sanitize_for_prompt(category["category_description"]),
            transcripts_section=category["transcript_sections"],
            research_plan=_sanitize_for_prompt(extraction_strategy),
            cross_category_notes=_sanitize_for_prompt(cross_category_notes),
        )

        user_prompt = extraction_prompts["user_prompt"].format(
            formatted_section=_sanitize_for_prompt(formatted_section)
        )

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]

        max_retries = etl_config.max_retries
        for attempt in range(max_retries):
            try:
                response = await complete_with_tools(
                    messages=messages,
                    tools=[extraction_prompts["tool_definition"]],
                    context=context,
                    llm_params={
                        "model": etl_config.get_model("category_extraction"),
                        "temperature": etl_config.temperature,
                        "max_tokens": etl_config.get_max_tokens("category_extraction"),
                    },
                )

                tool_call = response["choices"][0]["message"]["tool_calls"][0]
                raw_data = json.loads(tool_call["function"]["arguments"])
                validated = CategoryExtractionResponse.model_validate(raw_data)
                extracted_data = validated.model_dump()

                extracted_data["index"] = index
                extracted_data["name"] = category["category_name"]
                extracted_data["report_section"] = category.get("report_section", "Results Summary")
                if not extracted_data.get("title"):
                    extracted_data["title"] = category["category_name"]

                metrics = response.get("metrics", {})
                _accumulate_llm_cost(context, metrics)
                logger.info(
                    "etl.call_summary.llm_usage",
                    execution_id=execution_id,
                    stage=f"extraction:{category['category_name']}",
                    prompt_tokens=metrics.get("prompt_tokens", 0),
                    completion_tokens=metrics.get("completion_tokens", 0),
                    total_cost=metrics.get("total_cost", 0),
                    response_time=metrics.get("response_time", 0),
                )

                logger.info(
                    "etl.call_summary.category_completed",
                    execution_id=execution_id,
                    category_name=category["category_name"],
                    rejected=extracted_data.get("rejected", False),
                )

                return extracted_data

            except (KeyError, IndexError, json.JSONDecodeError, TypeError, ValidationError) as e:
                logger.warning(
                    "etl.call_summary.category_extraction_parse_error",
                    execution_id=execution_id,
                    category_name=category["category_name"],
                    error=str(e),
                    attempt=attempt + 1,
                )

                if attempt < max_retries - 1:
                    continue
                raise RuntimeError(
                    f"Category extraction failed for '{category['category_name']}' "
                    f"after {max_retries} attempts: {str(e)}"
                ) from e
            except Exception as e:
                logger.error(
                    "etl.call_summary.category_extraction_error",
                    execution_id=execution_id,
                    category_name=category["category_name"],
                    error=str(e),
                    attempt=attempt + 1,
                )

                if attempt < max_retries - 1:
                    delay = min(
                        etl_config.retry_base_delay * (2**attempt),
                        etl_config.retry_max_delay,
                    )
                    delay += random.uniform(0, 0.5 * delay)
                    await asyncio.sleep(delay)
                    continue
                raise RuntimeError(
                    f"Category extraction failed for '{category['category_name']}' "
                    f"after {max_retries} attempts: {str(e)}"
                ) from e

        raise RuntimeError(
            f"Category extraction failed for '{category['category_name']}' "
            f"after {max_retries} attempts"
        )


async def _process_categories(
    categories: list, research_plan_data: dict, extraction_prompts: dict, etl_context: dict
) -> list:
    """
    Process all categories concurrently and extract data from transcripts.

    Uses asyncio.gather with a semaphore to limit concurrency. Each category
    executes independently with no shared mutable state. Results are sorted
    by index for deterministic ordering regardless of completion order.

    Args:
        categories: List of category configurations
        research_plan_data: Research plan from LLM
        extraction_prompts: Prompts for extraction
        etl_context: Dict with keys: retrieval_params, bank_info, quarter,
            fiscal_year, context, execution_id

    Returns:
        List of category results (both accepted and rejected), sorted by index
    """
    semaphore = asyncio.Semaphore(etl_config.max_concurrent_extractions)
    total_categories = len(categories)

    tasks = [
        _extract_single_category(
            index=i,
            category=cat,
            research_plan_data=research_plan_data,
            extraction_prompts=extraction_prompts,
            etl_context=etl_context,
            semaphore=semaphore,
            total_categories=total_categories,
        )
        for i, cat in enumerate(categories, 1)
    ]

    results = await asyncio.gather(*tasks)
    return sorted(results, key=lambda r: r.get("index", 0))


def _insert_toc_at_position(doc, toc_entries: list) -> bool:
    """
    Insert TOC paragraphs before the first Heading 1 in the document body.

    Uses lxml addprevious() to position the TOC after banner/title but before
    body content. Falls back to appending at the end if position detection fails.

    Args:
        doc: Word Document object
        toc_entries: List of (title, heading_level) tuples

    Returns:
        True if TOC was inserted at the correct position, False if fallback needed
    """
    try:
        body = doc.element.body
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        # Find the first Heading 1 paragraph element
        first_h1 = None
        for para_elem in body.iterchildren(f"{ns}p"):
            for ppr in para_elem.iterchildren(f"{ns}pPr"):
                for pstyle in ppr.iterchildren(f"{ns}pStyle"):
                    if pstyle.get(f"{ns}val") == "Heading1":
                        first_h1 = para_elem
                        break
                if first_h1:
                    break
            if first_h1:
                break

        if first_h1 is None:
            return False

        # Count paragraphs before adding TOC
        paras_before = len(list(body.iterchildren(f"{ns}p")))

        # Build TOC paragraphs and insert before first Heading 1
        add_table_of_contents(doc, toc_entries)

        # The TOC was appended at the end; move its paragraphs before first_h1
        # Determine which paragraphs were added by comparing counts
        all_paras = list(body.iterchildren(f"{ns}p"))
        toc_paras = all_paras[paras_before:]

        for toc_para in toc_paras:
            body.remove(toc_para)
            first_h1.addprevious(toc_para)

        return True

    except Exception as e:  # pylint: disable=broad-except
        logger.warning(
            "etl.call_summary.toc_insertion_failed",
            error=str(e),
        )
        return False


def _generate_document(valid_categories: list, etl_context: dict) -> tuple:
    """
    Generate Word document from category results.

    Builds body content first (collecting TOC entries), then inserts the TOC
    before the first section heading with real heading titles as fallback text.

    Args:
        valid_categories: List of accepted category results
        etl_context: Dict with keys: bank_info, quarter, fiscal_year, execution_id

    Returns:
        Tuple of (filepath, docx_filename)
    """
    bank_info = etl_context["bank_info"]
    quarter = etl_context["quarter"]
    fiscal_year = etl_context["fiscal_year"]
    execution_id = etl_context["execution_id"]
    doc = Document()
    setup_document_formatting(doc)

    etl_dir = os.path.dirname(os.path.abspath(__file__))
    config_dir = os.path.join(etl_dir, "config")
    add_banner_image(doc, config_dir)

    add_document_title(doc, quarter, fiscal_year, bank_info["bank_symbol"])

    # Build body content, collecting TOC entries as we go
    toc_entries = []

    sorted_categories = sorted(
        valid_categories,
        key=lambda x: (
            0 if x.get("report_section", "Results Summary") == "Results Summary" else 1,
            x.get("report_section", "Results Summary"),
            x.get("index", 0),
        ),
    )

    for idx, (section_name, section_categories) in enumerate(
        groupby(sorted_categories, key=lambda x: x.get("report_section", "Results Summary"))
    ):
        section_categories = list(section_categories)
        add_section_heading(doc, section_name, is_first_section=idx == 0)
        toc_entries.append((section_name, 1))

        for i, category_data in enumerate(section_categories, 1):
            title = category_data.get("title") or category_data.get("name") or "Untitled"
            if not category_data.get("rejected", False):
                toc_entries.append((title, 2))

            add_structured_content_to_doc(doc, category_data, heading_level=2)

            if i < len(section_categories):
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(6)
                spacer.add_run()

    # Insert TOC with real heading titles before the first section heading
    if not _insert_toc_at_position(doc, toc_entries):
        add_table_of_contents(doc, toc_entries)

    try:
        mark_document_for_update(doc)
    except (AttributeError, ValueError):
        pass

    validate_document_content(doc)

    output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
    os.makedirs(output_dir, exist_ok=True)

    filename_base = f"{bank_info['bank_symbol']}_{fiscal_year}_{quarter}_call_summary"
    docx_filename = f"{filename_base}.docx"
    filepath = os.path.join(output_dir, docx_filename)
    doc.save(filepath)

    logger.info("etl.call_summary.document_saved", execution_id=execution_id, filepath=filepath)

    return filepath, docx_filename


async def _save_to_database(
    category_results: list,
    valid_categories: list,
    filepath: str,
    docx_filename: str,
    etl_context: dict,
) -> None:
    """
    Save report metadata to database.

    Args:
        category_results: All category results
        valid_categories: Accepted category results
        filepath: Local file path
        docx_filename: Document filename
        etl_context: Dict with keys: bank_info, quarter, fiscal_year, bank_type, execution_id
    """
    bank_info = etl_context["bank_info"]
    quarter = etl_context["quarter"]
    fiscal_year = etl_context["fiscal_year"]
    bank_type = etl_context["bank_type"]
    execution_id = etl_context["execution_id"]

    report_metadata = get_standard_report_metadata()
    generation_timestamp = datetime.now()

    stage = "connecting"
    try:
        async with get_connection() as conn:
            stage = "deleting existing report"
            delete_result = await conn.execute(
                text(
                    """
                DELETE FROM aegis_reports
                WHERE bank_id = :bank_id
                  AND fiscal_year = :fiscal_year
                  AND quarter = :quarter
                  AND report_type = :report_type
                RETURNING id
                """
                ),
                {
                    "bank_id": bank_info["bank_id"],
                    "fiscal_year": fiscal_year,
                    "quarter": quarter,
                    "report_type": report_metadata["report_type"],
                },
            )
            delete_result.fetchall()

            stage = "inserting new report"
            result = await conn.execute(
                text(
                    """
                INSERT INTO aegis_reports (
                    report_name,
                    report_description,
                    report_type,
                    bank_id,
                    bank_name,
                    bank_symbol,
                    fiscal_year,
                    quarter,
                    local_filepath,
                    s3_document_name,
                    s3_pdf_name,
                    generation_date,
                    generated_by,
                    execution_id,
                    metadata
                ) VALUES (
                    :report_name,
                    :report_description,
                    :report_type,
                    :bank_id,
                    :bank_name,
                    :bank_symbol,
                    :fiscal_year,
                    :quarter,
                    :local_filepath,
                    :s3_document_name,
                    :s3_pdf_name,
                    :generation_date,
                    :generated_by,
                    :execution_id,
                    :metadata
                )
                RETURNING id
                """
                ),
                {
                    "report_name": report_metadata["report_name"],
                    "report_description": report_metadata["report_description"],
                    "report_type": report_metadata["report_type"],
                    "bank_id": bank_info["bank_id"],
                    "bank_name": bank_info["bank_name"],
                    "bank_symbol": bank_info["bank_symbol"],
                    "fiscal_year": fiscal_year,
                    "quarter": quarter,
                    "local_filepath": filepath,
                    "s3_document_name": docx_filename,
                    "s3_pdf_name": None,
                    "generation_date": generation_timestamp,
                    "generated_by": "call_summary_etl",
                    "execution_id": execution_id,
                    "metadata": json.dumps(
                        {
                            "bank_type": bank_type,
                            "categories_processed": len(category_results),
                            "categories_included": len(valid_categories),
                            "categories_rejected": len(category_results) - len(valid_categories),
                        }
                    ),
                },
            )
            result.fetchone()

            await conn.commit()

    except SQLAlchemyError as e:
        logger.error(
            "etl.call_summary.database_error",
            execution_id=execution_id,
            stage=stage,
            filepath=filepath,
            error=str(e),
        )
        raise


def _timing_summary(marks: list) -> dict:
    """
    Convert timing marks to elapsed-seconds summary.

    Args:
        marks: List of (stage_name, timestamp) tuples

    Returns:
        Dict mapping "{stage}_s" to elapsed seconds, plus "total_s"
    """
    if len(marks) < 2:
        return {}
    summary = {}
    for i in range(1, len(marks)):
        summary[f"{marks[i][0]}_s"] = round(marks[i][1] - marks[i - 1][1], 2)
    summary["total_s"] = round(marks[-1][1] - marks[0][1], 2)
    return summary


async def generate_call_summary(  # pylint: disable=too-many-statements
    bank_name: str, fiscal_year: int, quarter: str
) -> CallSummaryResult:
    """
    Generate a call summary by directly calling transcript functions.

    Args:
        bank_name: ID, name, or symbol of the bank
        fiscal_year: Year (e.g., 2024)
        quarter: Quarter (e.g., "Q3")

    Returns:
        CallSummaryResult with filepath and category counts

    Raises:
        CallSummaryUserError: For expected errors (bad input, no data)
        CallSummarySystemError: For unexpected system/infrastructure errors
    """
    marks = [("start", time.monotonic())]
    execution_id = str(uuid.uuid4())
    logger.info(
        "etl.call_summary.started",
        execution_id=execution_id,
        bank_name=bank_name,
        fiscal_year=fiscal_year,
        quarter=quarter,
    )

    try:
        # Get bank info from monitored institutions config
        bank_info = get_bank_info_from_config(bank_name)

        # Verify data availability (single database check)
        await verify_and_get_availability(
            bank_info["bank_id"], bank_info["bank_name"], fiscal_year, quarter
        )

        ssl_config = setup_ssl()
        auth_config = await setup_authentication(execution_id, ssl_config)

        if not auth_config["success"]:
            error_msg = f"Authentication failed: {auth_config.get('error', 'Unknown error')}"
            logger.error("etl.call_summary.auth_failed", execution_id=execution_id, error=error_msg)
            raise CallSummarySystemError(error_msg)

        context = {
            "execution_id": execution_id,
            "auth_config": auth_config,
            "ssl_config": ssl_config,
        }

        marks.append(("setup", time.monotonic()))

        categories = load_categories_from_xlsx(bank_info["bank_type"], execution_id)

        retrieval_params = {
            "bank_id": bank_info["bank_id"],
            "bank_name": bank_info["bank_name"],
            "bank_symbol": bank_info["bank_symbol"],
            "fiscal_year": fiscal_year,
            "quarter": quarter,
            "query_intent": "Generate comprehensive research plan for earnings call summary",
        }

        chunks = await retrieve_full_section(
            combo=retrieval_params, sections=SECTIONS_KEY_ALL, context=context
        )

        if not chunks:
            raise CallSummaryUserError(
                f"No transcript chunks found for {bank_info['bank_name']} "
                f"{quarter} {fiscal_year}"
            )

        formatted_transcript = format_full_section_chunks(
            chunks=chunks, combo=retrieval_params, context=context
        )

        marks.append(("retrieval", time.monotonic()))

        research_prompts = load_prompt_from_db(
            layer="call_summary_etl",
            name="research_plan",
            compose_with_globals=False,
            available_databases=None,
            execution_id=execution_id,
        )

        # Format categories using standardized XML format
        categories_text = format_categories_for_prompt(categories)

        research_prompts["system_prompt"] = research_prompts["system_prompt"].format(
            categories_list=categories_text,
            bank_name=bank_info["bank_name"],
            bank_symbol=bank_info["bank_symbol"],
            quarter=quarter,
            fiscal_year=fiscal_year,
        )

        research_prompts["user_prompt_template"] = research_prompts["user_prompt"]

        research_plan_data = await _generate_research_plan(
            context, research_prompts, formatted_transcript, execution_id
        )

        # Check for coverage gaps in research plan
        planned_indices = {p.get("index") for p in research_plan_data.get("category_plans", [])}
        all_indices = set(range(1, len(categories) + 1))
        skipped_indices = all_indices - planned_indices
        if skipped_indices:
            skipped_names = [
                categories[idx - 1]["category_name"] for idx in sorted(skipped_indices)
            ]
            logger.info(
                "etl.call_summary.research_plan_skipped_categories",
                execution_id=execution_id,
                skipped_count=len(skipped_indices),
                skipped_categories=skipped_names,
            )

        # Check Q&A group coverage across research plan
        qa_chunks = [c for c in chunks if c.get("section_name") == SECTION_QA]
        all_qa_ids = {c["qa_group_id"] for c in qa_chunks if c.get("qa_group_id") is not None}
        referenced_qa_ids = set()
        for plan in research_plan_data.get("category_plans", []):
            for qa_id in plan.get("relevant_qa_groups", []):
                referenced_qa_ids.add(qa_id)
        unreferenced_qa_ids = all_qa_ids - referenced_qa_ids
        if unreferenced_qa_ids:
            logger.warning(
                "etl.call_summary.qa_group_coverage_gap",
                execution_id=execution_id,
                total_qa_groups=len(all_qa_ids),
                unreferenced_count=len(unreferenced_qa_ids),
                unreferenced_qa_ids=sorted(unreferenced_qa_ids),
            )
        elif all_qa_ids:
            logger.info(
                "etl.call_summary.qa_group_coverage_complete",
                execution_id=execution_id,
                total_qa_groups=len(all_qa_ids),
                referenced_qa_groups=len(referenced_qa_ids),
            )

        marks.append(("research_plan", time.monotonic()))

        extraction_prompts = load_prompt_from_db(
            layer="call_summary_etl",
            name="category_extraction",
            compose_with_globals=False,
            available_databases=None,
            execution_id=execution_id,
        )

        dedup_prompts = load_prompt_from_db(
            layer="call_summary_etl",
            name="deduplication",
            compose_with_globals=False,
            available_databases=None,
            execution_id=execution_id,
        )

        # Cache sections to avoid redundant DB calls during category processing
        etl_context = {
            "retrieval_params": retrieval_params,
            "bank_info": bank_info,
            "quarter": quarter,
            "fiscal_year": fiscal_year,
            "context": context,
            "execution_id": execution_id,
            "bank_type": bank_info["bank_type"],
            "section_cache": {
                SECTIONS_KEY_ALL: chunks,
                SECTIONS_KEY_MD: [c for c in chunks if c["section_name"] == SECTION_MD],
                SECTIONS_KEY_QA: [c for c in chunks if c["section_name"] == SECTION_QA],
            },
        }

        category_results = await _process_categories(
            categories=categories,
            research_plan_data=research_plan_data,
            extraction_prompts=extraction_prompts,
            etl_context=etl_context,
        )

        marks.append(("extraction", time.monotonic()))

        category_results = await _deduplicate_results_llm(
            category_results, dedup_prompts, context, execution_id
        )

        marks.append(("deduplication", time.monotonic()))

        valid_categories = [c for c in category_results if not c.get("rejected", False)]

        if not valid_categories:
            raise CallSummaryUserError(
                "All categories were rejected - no content to generate document"
            )

        filepath, docx_filename = _generate_document(
            valid_categories=valid_categories, etl_context=etl_context
        )

        marks.append(("document", time.monotonic()))

        await _save_to_database(
            category_results=category_results,
            valid_categories=valid_categories,
            filepath=filepath,
            docx_filename=docx_filename,
            etl_context=etl_context,
        )

        marks.append(("save", time.monotonic()))

        cost_summary = _get_total_llm_cost(context)
        logger.info(
            "etl.call_summary.completed",
            execution_id=execution_id,
            num_categories=len(valid_categories),
            llm_calls=cost_summary["llm_calls"],
            total_tokens=cost_summary["total_tokens"],
            total_cost=cost_summary["total_cost"],
            **_timing_summary(marks),
        )

        return CallSummaryResult(
            filepath=filepath,
            total_categories=len(category_results),
            included_categories=len(valid_categories),
            total_cost=cost_summary["total_cost"],
            total_tokens=cost_summary["total_tokens"],
        )

    except CallSummaryError:
        raise
    except (ValueError, RuntimeError) as e:
        logger.error("etl.call_summary.error", execution_id=execution_id, error=str(e))
        raise CallSummaryUserError(str(e)) from e
    except Exception as e:
        error_msg = f"Error generating call summary: {str(e)}"
        logger.error(
            "etl.call_summary.error", execution_id=execution_id, error=error_msg, exc_info=True
        )
        raise CallSummarySystemError(error_msg) from e


def main():
    """Main entry point for command-line execution."""
    parser = argparse.ArgumentParser(
        description="Generate call summary reports",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )

    parser.add_argument("--bank", required=True, help="Bank ID, name, or symbol")
    parser.add_argument("--year", type=int, required=True, help="Fiscal year")
    parser.add_argument(
        "--quarter", required=True, choices=["Q1", "Q2", "Q3", "Q4"], help="Quarter"
    )

    args = parser.parse_args()

    postgresql_prompts()

    print(f"\n🔄 Generating report for {args.bank} {args.quarter} {args.year}...\n")

    try:
        result = asyncio.run(
            generate_call_summary(bank_name=args.bank, fiscal_year=args.year, quarter=args.quarter)
        )
        print(
            f"✅ Complete: {result.filepath}\n"
            f"   Categories: {result.included_categories}/{result.total_categories} included\n"
            f"   LLM cost: ${result.total_cost:.4f}, Tokens: {result.total_tokens:,}"
        )
    except CallSummaryUserError as e:
        print(f"⚠️ {e}", file=sys.stderr)
        sys.exit(1)
    except CallSummaryError as e:
        print(f"❌ {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
