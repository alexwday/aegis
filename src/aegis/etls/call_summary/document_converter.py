"""Document converter utilities for call summary ETL."""

import os
import re
from typing import Dict
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from aegis.utils.logging import get_logger

logger = get_logger()


def get_standard_report_metadata() -> Dict[str, str]:
    """
    Get standard metadata for call summary reports.

    Returns:
        Dictionary with report_name and report_description
    """
    return {
        "report_name": "Earnings Call Summary",
        "report_description": (
            "AI-generated comprehensive summary of quarterly earnings call transcripts, "
            "extracting key financial metrics, strategic insights, and management guidance. "
            "Includes categorized analysis of financial performance, credit quality, "
            "capital position, business segments, and forward-looking statements."
        ),
        "report_type": "call_summary",
    }


def add_page_numbers(doc):
    """Add page numbers to the footer of the document."""
    for section in doc.sections:
        footer = section.footer
        footer.paragraphs[0].clear()
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer_para.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        run._element.append(fld_char1)  # pylint: disable=protected-access

        instr_text = OxmlElement("w:instrText")
        instr_text.text = "PAGE"
        run._element.append(instr_text)  # pylint: disable=protected-access

        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._element.append(fld_char2)  # pylint: disable=protected-access


def setup_toc_styles(doc):
    """Setup custom TOC styles for formatting."""
    styles = doc.styles

    try:
        toc1_style = styles["TOC 1"]
    except KeyError:
        toc1_style = styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)

    toc1_style.font.size = Pt(8)
    toc1_style.font.bold = True
    toc1_style.paragraph_format.space_before = Pt(1)
    toc1_style.paragraph_format.space_after = Pt(1)
    toc1_style.paragraph_format.line_spacing = 0.9

    try:
        toc2_style = styles["TOC 2"]
    except KeyError:
        toc2_style = styles.add_style("TOC 2", WD_STYLE_TYPE.PARAGRAPH)

    toc2_style.font.size = Pt(7)
    toc2_style.font.bold = False
    toc2_style.paragraph_format.left_indent = Inches(0.2)
    toc2_style.paragraph_format.space_before = Pt(0)
    toc2_style.paragraph_format.space_after = Pt(0.5)
    toc2_style.paragraph_format.line_spacing = 0.9


def add_table_of_contents(doc):
    """Add a real table of contents field to the document."""
    setup_toc_styles(doc)

    toc_title = doc.add_paragraph()
    toc_title_run = toc_title.add_run("Contents")
    toc_title_run.font.size = Pt(10)
    toc_title_run.font.bold = True
    toc_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_title.paragraph_format.space_after = Pt(3)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run()

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    fld_char.set(qn("w:dirty"), "true")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-2" \\h \\z \\u'

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")

    fld_char3 = OxmlElement("w:t")
    fld_char3.text = "[Table of Contents will be generated here]"

    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")

    r_element = run._element  # pylint: disable=protected-access
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char3)
    r_element.append(fld_char4)

    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def mark_document_for_update(doc):
    """Mark the document settings to update fields on open."""
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def _apply_base_formatting(run, base_font_size, base_color, base_italic):
    """Apply base formatting to a text run."""
    if base_font_size:
        run.font.size = base_font_size
    if base_color:
        run.font.color.rgb = base_color
    run.italic = base_italic


def _find_next_format_match(text):
    """Find the next bold or underline match and return match and type."""
    bold_match = re.search(r"\*\*([^*]+)\*\*", text)
    underline_match = re.search(r"__([^_]+)__", text)

    if bold_match and underline_match:
        if bold_match.start() < underline_match.start():
            return bold_match, "bold"
        return underline_match, "underline"
    if bold_match:
        return bold_match, "bold"
    if underline_match:
        return underline_match, "underline"
    return None, None


def parse_and_format_text(
    paragraph, content: str, base_font_size=None, base_color=None, base_italic=False
) -> None:
    """
    Parse markdown-style formatting and add formatted runs to paragraph.
    Supports **bold** for emphasis and __underline__ for important phrases.

    Args:
        paragraph: Word paragraph object to add formatted text to
        content: Text containing markdown formatting
        base_font_size: Base font size for all runs (optional)
        base_color: Base RGB color for all runs (optional)
        base_italic: Whether base text should be italic (default False)
    """
    if "**" not in content and "__" not in content:
        run = paragraph.add_run(content)
        _apply_base_formatting(run, base_font_size, base_color, base_italic)
        return

    remaining_text = content
    while remaining_text:
        next_match, match_type = _find_next_format_match(remaining_text)

        if next_match:
            if next_match.start() > 0:
                run = paragraph.add_run(remaining_text[: next_match.start()])
                _apply_base_formatting(run, base_font_size, base_color, base_italic)

            run = paragraph.add_run(next_match.group(1))
            _apply_base_formatting(run, base_font_size, base_color, base_italic)

            if match_type == "bold":
                run.bold = True
            elif match_type == "underline":
                run.underline = True

            remaining_text = remaining_text[next_match.end() :]  # noqa: E203
        else:
            run = paragraph.add_run(remaining_text)
            _apply_base_formatting(run, base_font_size, base_color, base_italic)
            break


def _add_category_heading(doc, title: str, heading_level: int) -> None:
    """Add formatted category heading to document."""
    heading = doc.add_heading(title, level=heading_level)
    for run in heading.runs:
        run.font.size = Pt(11) if heading_level == 1 else Pt(10)
        run.font.bold = True
    heading.paragraph_format.space_before = Pt(6) if heading_level == 1 else Pt(4)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.page_break_before = False


def _add_evidence_paragraph(
    doc, evidence: dict, is_last_evidence: bool, is_last_statement: bool
) -> None:
    """Add formatted evidence paragraph to document."""
    evidence_para = doc.add_paragraph()
    evidence_para.paragraph_format.left_indent = Inches(0.75)
    evidence_para.paragraph_format.right_indent = Inches(0.5)
    evidence_para.paragraph_format.first_line_indent = Inches(0)
    evidence_para.paragraph_format.space_after = Pt(1)
    evidence_para.paragraph_format.line_spacing = 1.0

    if not is_last_evidence:
        evidence_para.paragraph_format.keep_with_next = True
    elif not is_last_statement:
        evidence_para.paragraph_format.space_after = Pt(4)

    evidence_para.paragraph_format.widow_control = True

    evidence_content = evidence["content"]

    if evidence["type"] == "quote":
        evidence_para.add_run('"').italic = True
        parse_and_format_text(
            evidence_para,
            evidence_content,
            base_font_size=Pt(8),
            base_color=RGBColor(64, 64, 64),
            base_italic=True,
        )
        closing_run = evidence_para.add_run('"')
        closing_run.italic = True
        closing_run.font.size = Pt(8)
        closing_run.font.color.rgb = RGBColor(64, 64, 64)
    else:
        parse_and_format_text(
            evidence_para,
            evidence_content,
            base_font_size=Pt(8),
            base_color=RGBColor(64, 64, 64),
            base_italic=True,
        )

    speaker = evidence.get("speaker", "Unknown")
    speaker_run = evidence_para.add_run(f" — {speaker}")
    speaker_run.italic = False
    speaker_run.font.size = Pt(7)
    speaker_run.font.color.rgb = RGBColor(96, 96, 96)


def setup_document_formatting(doc) -> None:
    """Configure document margins and add page numbers."""
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.5)
        section.gutter = Inches(0)
    add_page_numbers(doc)


def add_banner_image(doc, config_dir: str) -> None:
    """Add banner image to document if found in config directory."""
    banner_path = None
    for ext in ["jpg", "jpeg", "png"]:
        potential_banner = os.path.join(config_dir, f"banner.{ext}")
        if os.path.exists(potential_banner):
            banner_path = potential_banner
            break

    if banner_path:
        try:
            doc.add_picture(banner_path, width=Inches(7.4))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_paragraph.paragraph_format.space_after = Pt(3)
        except (FileNotFoundError, OSError, ValueError):
            pass


def add_document_title(doc, quarter: str, fiscal_year: int, bank_symbol: str) -> None:
    """Add formatted title to document."""
    title_text = f"{quarter}/{str(fiscal_year)[-2:]} Results and Call Summary - {bank_symbol}"
    title = doc.add_heading(title_text, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(14)
        run.font.bold = True
        try:
            run.font.name = "Arial"
        except (KeyError, ValueError, AttributeError):
            pass
    title.paragraph_format.space_after = Pt(4)


def add_section_heading(doc, section_name: str, is_first_section: bool = False) -> None:
    """Add formatted section heading to document."""
    section_heading = doc.add_heading(section_name, level=1)
    section_heading.paragraph_format.space_before = Pt(10)
    section_heading.paragraph_format.space_after = Pt(6)
    section_heading.paragraph_format.keep_with_next = True
    if not is_first_section:
        section_heading.paragraph_format.page_break_before = True


def add_structured_content_to_doc(doc, category_data: dict, heading_level: int = 2) -> None:
    """
    Add structured category data directly to Word document with proper formatting.

    Args:
        doc: Word document object
        category_data: Dictionary with title, summary_statements, evidence structure
        heading_level: Heading level for category title (default 2)
    """
    if category_data.get("rejected", False):
        return

    try:
        _add_category_heading(doc, category_data["title"], heading_level)

        statements = category_data.get("summary_statements", [])
        for idx, statement_data in enumerate(statements):
            statement_para = doc.add_paragraph(style="List Bullet")
            parse_and_format_text(statement_para, statement_data["statement"], base_font_size=Pt(9))
            statement_para.paragraph_format.space_after = Pt(2)
            statement_para.paragraph_format.line_spacing = 1.0
            statement_para.paragraph_format.keep_with_next = True

            evidence_list = statement_data.get("evidence", [])
            if evidence_list:
                for i, evidence in enumerate(evidence_list):
                    is_last_evidence = i == len(evidence_list) - 1
                    is_last_statement = idx == len(statements) - 1
                    _add_evidence_paragraph(doc, evidence, is_last_evidence, is_last_statement)

    except (KeyError, AttributeError, TypeError, ValueError) as e:
        logger.warning(f"Error formatting category content: {e}", exc_info=True)
