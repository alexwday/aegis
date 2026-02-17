"""
Document converter utilities for CM Readthrough ETL.

This module provides functions to create combined Word documents for capital markets analysis.
"""

import re
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml import parse_xml as docx_parse_xml
from docx.oxml.ns import nsdecls as docx_nsdecls
from aegis.utils.logging import get_logger

logger = get_logger()


def parse_xml(xml_string):
    """Parse XML string into OxmlElement using python-docx's parse_xml."""
    return docx_parse_xml(xml_string)


def nsdecls(*prefixes):
    """Generate namespace declarations for XML using python-docx's nsdecls."""
    return docx_nsdecls(*prefixes)


def _clean_text_for_word(text: str) -> str:
    """
    Clean text for Word document insertion by removing problematic characters.

    Args:
        text: Input text that may contain invalid characters

    Returns:
        Cleaned text safe for Word document insertion
    """
    if not text:
        return text
    text = text.replace("\x00", "")
    return text


def validate_document_content(doc) -> None:
    """
    Validate generated document has meaningful content before saving.

    Checks that:
    1. Document has at least one non-empty title paragraph
    2. Document has at least one table
    3. Document has body content beyond the title
    """
    if not doc.paragraphs:
        raise ValueError("Document has no paragraphs")

    has_title = any(
        "Read Through For Capital Markets:" in para.text
        for para in doc.paragraphs
        if para.text.strip()
    )
    has_table = len(doc.tables) > 0
    has_body_text = any(
        para.text.strip() and "Read Through For Capital Markets:" not in para.text
        for para in doc.paragraphs
    )

    missing = []
    if not has_title:
        missing.append("report title")
    if not has_table:
        missing.append("table content")
    if not has_body_text:
        missing.append("body text")

    if missing:
        raise ValueError(f"Document missing required content: {', '.join(missing)}")


# Patterns for financial metrics that should be emphasized
_METRIC_PATTERNS = [
    r"-?\$[\d,]+(?:\.\d+)?\s*(?:MM|BN|TN|K|M|B)\b",
    r"-?\$[\d,]+(?:\.\d+)?(?!\s*(?:MM|BN|TN|K|M|B))\b",
    r"\d+(?:\.\d+)?\s*bps\b",
    r"\d+(?:\.\d+)?%",
]
_METRIC_COMBINED_PATTERN = re.compile("|".join(f"(?:{p})" for p in _METRIC_PATTERNS))


def auto_bold_metrics(text: str) -> str:
    """
    Auto-bold financial metrics not already inside <strong><u>...</u></strong>.

    Args:
        text: Input text that may already include HTML emphasis
    """
    if not text:
        return text

    for pattern in _METRIC_PATTERNS:
        text = _bold_unbolded_matches(text, pattern)
    return text


def _bold_unbolded_matches(text: str, pattern: str) -> str:
    """Bold regex matches not already wrapped in strong or strong+underline tags."""

    def replacer(match):
        start = match.start()
        prefix = text[:start]
        # Check if inside <strong><u>...</u></strong>
        open_su = len(re.findall(r"<strong><u>", prefix, flags=re.IGNORECASE))
        close_su = len(re.findall(r"</u></strong>", prefix, flags=re.IGNORECASE))
        if open_su > close_su:
            return match.group(0)
        # Check if inside <strong>...</strong> (bold only)
        open_s = len(re.findall(r"<strong>(?!<u>)", prefix, flags=re.IGNORECASE))
        close_s = len(re.findall(r"(?<!</u>)</strong>", prefix, flags=re.IGNORECASE))
        if open_s > close_s:
            return match.group(0)
        return f"<strong>{match.group(0)}</strong>"

    return re.sub(pattern, replacer, text)


def _add_text_with_metric_bold(paragraph, text: str, font_size: int = 9) -> None:
    """Add plain text to a paragraph while bolding metric tokens."""
    clean = _clean_text_for_word(text)
    if not clean:
        return

    parts = re.split(f"({_METRIC_COMBINED_PATTERN.pattern})", clean)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        run.font.size = Pt(font_size)
        if _METRIC_COMBINED_PATTERN.fullmatch(part):
            run.font.bold = True


def _add_formatted_runs(paragraph, text: str, font_size: int = 9) -> None:
    """
    Add formatted runs to a paragraph processing HTML tags.

    Supports:
      - <strong><u>text</u></strong> for bold+underline (key insight phrases)
      - <strong>text</strong> for bold only (financial metrics)

    Args:
        paragraph: Document paragraph object
        text: Text with HTML formatting
        font_size: Font size in points
    """
    text = _clean_text_for_word(text)

    # Match bold+underline OR bold-only tags
    pattern = r"<strong><u>(.*?)</u></strong>|<strong>(.*?)</strong>"

    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end : match.start()])
            run.font.size = Pt(font_size)

        if match.group(1) is not None:
            # <strong><u>...</u></strong> — bold + underline
            run = paragraph.add_run(match.group(1))
            run.font.bold = True
            run.font.underline = True
            run.font.size = Pt(font_size)
        else:
            # <strong>...</strong> — bold only
            run = paragraph.add_run(match.group(2))
            run.font.bold = True
            run.font.size = Pt(font_size)

        last_end = match.end()

    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.size = Pt(font_size)


def generate_main_title(quarter: str, year: int) -> str:
    """Generate main title for the report."""
    return (
        f"Read Through For Capital Markets: {quarter}/{str(year)[2:]} Select U.S. & European Banks"
    )


def add_page_footer(section) -> None:
    """
    Add footer to page with horizontal line, left-aligned source, and right-aligned RBC.

    Args:
        section: Document section object
    """
    footer = section.footer
    footer.is_linked_to_previous = False

    line_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(3)

    pPr = line_para._element.get_or_add_pPr()
    pBdr = parse_xml(
        r"<w:pBdr {}>"
        r'<w:top w:val="single" w:sz="6" w:color="000000"/>'
        r"</w:pBdr>".format(nsdecls("w"))
    )
    pPr.append(pBdr)

    footer_table = footer.add_table(rows=1, cols=2, width=Inches(10))
    footer_table.autofit = False

    left_cell = footer_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.add_run("Source: Company Reports, Transcripts")
    left_run.font.size = Pt(8)
    left_run.font.italic = True

    right_cell = footer_table.rows[0].cells[1]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_run = right_para.add_run("RBC")
    right_run.font.size = Pt(8)
    right_run.font.italic = True

    tbl = footer_table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(r"<w:tblPr {}/> ".format(nsdecls("w")))
        tbl.insert(0, tbl_pr)

    tbl_borders = parse_xml(
        r"<w:tblBorders {}>"
        r'<w:top w:val="none"/>'
        r'<w:bottom w:val="none"/>'
        r'<w:left w:val="none"/>'
        r'<w:right w:val="none"/>'
        r'<w:insideH w:val="none"/>'
        r'<w:insideV w:val="none"/>'
        r"</w:tblBorders>".format(nsdecls("w"))
    )
    tbl_pr.append(tbl_borders)


def create_combined_document(results: Dict[str, Any], output_path: str) -> None:
    """
    Create a combined Word document for CM Readthrough analysis with 3 sections.

    Args:
        results: Structured results from processing with 3 sections
        output_path: Path to save the document
    """
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    metadata = results.get("metadata", {})

    add_section1_outlook(doc, results)

    doc.add_page_break()

    add_section2_qa(doc, results)

    doc.add_page_break()

    add_section3_qa(doc, results)

    add_page_footer(doc.sections[0])

    validate_document_content(doc)
    doc.save(output_path)
    logger.info(f"Document saved to {output_path}")


def add_static_contents(doc: Document, metadata: Dict[str, Any]) -> None:
    """Add deterministic contents page for fixed 3-section report layout."""
    heading = doc.add_paragraph()
    heading_run = heading.add_run("Contents")
    heading_run.font.bold = True
    heading_run.font.size = Pt(11)
    heading.paragraph_format.space_after = Pt(4)

    section_1 = metadata.get("subtitle_section1", "Outlook")
    section_2 = metadata.get("subtitle_section2", "Conference calls: Market dynamics")
    section_3 = metadata.get("subtitle_section3", "Conference calls: Pipeline dynamics")

    for line in [f"1. {section_1}", f"2. {section_2}", f"3. {section_3}"]:
        para = doc.add_paragraph(line)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(1)


def add_section1_outlook(doc: Document, results: Dict[str, Any]) -> None:
    """
    Add Section 1: Investment Banking & Trading Outlook.

    Args:
        doc: Document object
        results: Full results dictionary
    """
    metadata = results["metadata"]
    outlook_data = results.get("outlook", {})

    section = doc.sections[-1] if doc.sections else doc.sections[0]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    main_title = generate_main_title(metadata["quarter"], metadata["fiscal_year"])
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(main_title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 32, 96)  # Dark blue
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle_text = metadata.get("subtitle_section1", "Outlook: Capital markets activity")
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle_text)
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not outlook_data:
        doc.add_paragraph("No outlook statements available.")
        return

    banks_with_content = {
        bank: content for bank, content in outlook_data.items() if content.get("statements")
    }

    if not banks_with_content:
        return

    name_to_symbol = {
        bank_name: data.get("bank_symbol", bank_name[:4].upper())
        for bank_name, data in outlook_data.items()
    }

    table = doc.add_table(rows=len(banks_with_content) + 1, cols=2)
    table.autofit = False
    table.allow_autofit = False

    header_cells = table.rows[0].cells
    header_cells[0].text = "Banks/\nSegments"
    header_cells[1].text = "Investment Banking and Trading Outlook"

    header_cells[0].width = Inches(0.8)  # Ticker column
    header_cells[1].width = Inches(9.2)  # Content column

    for cell in header_cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="002060"/>'.format(nsdecls("w")))
        cell._tc.get_or_add_tcPr().append(shading_elm)

        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text

    row_idx = 1
    for bank_name, bank_data in banks_with_content.items():
        row = table.rows[row_idx]

        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(9.2)

        ticker = name_to_symbol.get(bank_name, bank_name[:4].upper())
        ticker_cell = row.cells[0]
        ticker_para = ticker_cell.paragraphs[0]
        ticker_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ticker_run = ticker_para.add_run(ticker)
        ticker_run.font.size = Pt(9)
        ticker_run.font.bold = True

        content_cell = row.cells[1]

        statements = bank_data.get("statements", [])

        # Build grouped structure: category_group -> category -> [statements]
        from collections import OrderedDict

        groups: Dict[str, Dict[str, list]] = OrderedDict()
        for statement in statements:
            group = statement.get("category_group", "")
            category = statement.get("category", "")
            if group not in groups:
                groups[group] = OrderedDict()
            if category not in groups[group]:
                groups[group][category] = []
            groups[group][category].append(statement)

        has_groups = any(g for g in groups if g)

        first_para = True
        for group_name, categories_in_group in groups.items():
            # Render group header if groups exist and group name is non-empty
            if has_groups and group_name:
                if first_para:
                    group_para = content_cell.paragraphs[0]
                    first_para = False
                else:
                    group_para = content_cell.add_paragraph()
                group_run = group_para.add_run(group_name)
                group_run.font.bold = True
                group_run.font.underline = True
                group_run.font.size = Pt(9)
                group_para.paragraph_format.space_before = Pt(2)
                group_para.paragraph_format.space_after = Pt(0)

            for category, category_statements in categories_in_group.items():
                # Sort statements by relevance_score descending within each category
                category_statements.sort(key=lambda s: s.get("relevance_score", 0), reverse=True)

                if first_para:
                    category_para = content_cell.paragraphs[0]
                    first_para = False
                else:
                    category_para = content_cell.add_paragraph()
                category_run = category_para.add_run(f"{category}:")
                category_run.font.bold = True
                category_run.font.size = Pt(9)
                category_para.paragraph_format.space_before = Pt(0)
                category_para.paragraph_format.space_after = Pt(0)

                for statement in category_statements:
                    content_text = statement.get("formatted_quote", statement.get("statement", ""))
                    content_text = auto_bold_metrics(content_text)

                    quote_para = content_cell.add_paragraph()
                    quote_para.paragraph_format.space_before = Pt(0)
                    quote_para.paragraph_format.space_after = Pt(0)

                    # Add relevance score prefix if available
                    score = statement.get("relevance_score", 0)
                    if score > 0:
                        score_run = quote_para.add_run(f"[{score}/10] ")
                        score_run.font.size = Pt(9)
                        score_run.font.bold = True
                        score_run.font.color.rgb = RGBColor(0, 32, 96)

                    opening_run = quote_para.add_run('"')
                    opening_run.font.size = Pt(9)

                    _add_formatted_runs(quote_para, content_text, font_size=9)

                    closing_run = quote_para.add_run('"')
                    closing_run.font.size = Pt(9)

        row_idx += 1

    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(r"<w:tblPr {}/> ".format(nsdecls("w")))
        tbl.insert(0, tbl_pr)

    tbl_borders = parse_xml(
        r"<w:tblBorders {}>"
        r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="4" w:color="000000"/>'
        r"</w:tblBorders>".format(nsdecls("w"))
    )
    tbl_pr.append(tbl_borders)


def add_section2_qa(doc: Document, results: Dict[str, Any]) -> None:
    """
    Add Section 2: Q&A for Global Markets, Risk Management, Corporate Banking, Regulatory Changes.
    Sorted by theme first, then by bank within each theme.

    Args:
        doc: Document object
        results: Full results dictionary
    """
    metadata = results["metadata"]
    questions_data = results.get("section2_questions", {})

    section = doc.sections[-1]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    main_title = generate_main_title(metadata["quarter"], metadata["fiscal_year"])
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(main_title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 32, 96)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle_text = metadata.get("subtitle_section2", "Conference calls: Market dynamics")
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle_text)
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not questions_data:
        doc.add_paragraph("No Section 2 questions available.")
        return

    category_bank_questions = {}  # {category: {bank: [questions]}}
    for bank_name, bank_data in questions_data.items():
        ticker = bank_data.get("bank_symbol", "")
        for question in bank_data.get("questions", []):
            category = question.get("category", "Uncategorized")
            if category not in category_bank_questions:
                category_bank_questions[category] = {}
            if bank_name not in category_bank_questions[category]:
                category_bank_questions[category][bank_name] = {"ticker": ticker, "questions": []}
            category_bank_questions[category][bank_name]["questions"].append(question)

    if not category_bank_questions:
        doc.add_paragraph("No Section 2 questions available.")
        return

    total_rows = sum(len(banks) for banks in category_bank_questions.values())

    table = doc.add_table(rows=total_rows + 1, cols=3)
    table.autofit = False
    table.allow_autofit = False

    header_cells = table.rows[0].cells
    header_cells[0].text = "Themes"
    header_cells[1].text = "Banks"
    header_cells[2].text = "Relevant Questions"

    header_cells[0].width = Inches(0.8)  # Themes
    header_cells[1].width = Inches(0.8)  # Banks
    header_cells[2].width = Inches(8.4)  # Questions

    for cell in header_cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="002060"/>'.format(nsdecls("w")))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)

    row_idx = 1
    for category in sorted(category_bank_questions.keys()):
        banks_in_category = category_bank_questions[category]
        category_start_row = row_idx
        category_end_row = category_start_row + len(banks_in_category) - 1

        for bank_name in sorted(banks_in_category.keys()):
            bank_data = banks_in_category[bank_name]
            row = table.rows[row_idx]
            is_first_in_category = row_idx == category_start_row
            is_last_in_category = row_idx == category_end_row

            row.cells[0].width = Inches(0.8)
            row.cells[1].width = Inches(0.8)
            row.cells[2].width = Inches(8.4)

            if is_first_in_category:
                row.cells[0].text = category
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.bold = True
            else:
                row.cells[0].text = ""  # Empty for subsequent rows in same category

            row.cells[1].text = bank_data["ticker"]
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.bold = True

            first_question = True
            for i, question in enumerate(bank_data["questions"]):
                question_text = question.get("verbatim_question", "")
                question_text = _clean_text_for_word(question_text)

                if first_question:
                    q_para = row.cells[2].paragraphs[0]
                    q_para.style = "List Bullet"
                    first_question = False
                else:
                    q_para = row.cells[2].add_paragraph(style="List Bullet")
                _add_text_with_metric_bold(q_para, question_text, font_size=9)
                q_para.paragraph_format.space_before = Pt(0)
                q_para.paragraph_format.space_after = Pt(0)

            if is_last_in_category:
                for col_idx in [0, 1, 2]:
                    cell = row.cells[col_idx]
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = parse_xml(
                        r"<w:tcBorders {}>"
                        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
                        r"</w:tcBorders>".format(nsdecls("w"))
                    )
                    tcPr.append(tcBorders)
            else:
                for col_idx in [1, 2]:
                    cell = row.cells[col_idx]
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = parse_xml(
                        r"<w:tcBorders {}>"
                        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
                        r"</w:tcBorders>".format(nsdecls("w"))
                    )
                    tcPr.append(tcBorders)
                cell = row.cells[0]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    r"<w:tcBorders {}>"
                    r'<w:bottom w:val="none"/>'
                    r"</w:tcBorders>".format(nsdecls("w"))
                )
                tcPr.append(tcBorders)

            row_idx += 1

    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(r"<w:tblPr {}/> ".format(nsdecls("w")))
        tbl.insert(0, tbl_pr)

    tbl_borders = parse_xml(
        r"<w:tblBorders {}>"
        r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r"</w:tblBorders>".format(nsdecls("w"))
    )
    tbl_pr.append(tbl_borders)


def add_section3_qa(doc: Document, results: Dict[str, Any]) -> None:
    """
    Add Section 3: Q&A for Investment Banking/M&A and Transaction Banking.
    Sorted by theme first, then by bank within each theme.

    Args:
        doc: Document object
        results: Full results dictionary
    """
    metadata = results["metadata"]
    questions_data = results.get("section3_questions", {})

    section = doc.sections[-1]
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    main_title = generate_main_title(metadata["quarter"], metadata["fiscal_year"])
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(main_title)
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 32, 96)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle_text = metadata.get("subtitle_section3", "Conference calls: Pipeline dynamics")
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle_text)
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if not questions_data:
        doc.add_paragraph("No Section 3 questions available.")
        return

    category_bank_questions = {}  # {category: {bank: [questions]}}
    for bank_name, bank_data in questions_data.items():
        ticker = bank_data.get("bank_symbol", "")
        for question in bank_data.get("questions", []):
            category = question.get("category", "Uncategorized")
            if category not in category_bank_questions:
                category_bank_questions[category] = {}
            if bank_name not in category_bank_questions[category]:
                category_bank_questions[category][bank_name] = {"ticker": ticker, "questions": []}
            category_bank_questions[category][bank_name]["questions"].append(question)

    if not category_bank_questions:
        doc.add_paragraph("No Section 3 questions available.")
        return

    total_rows = sum(len(banks) for banks in category_bank_questions.values())

    table = doc.add_table(rows=total_rows + 1, cols=3)
    table.autofit = False
    table.allow_autofit = False

    header_cells = table.rows[0].cells
    header_cells[0].text = "Themes"
    header_cells[1].text = "Banks"
    header_cells[2].text = "Relevant Questions"

    header_cells[0].width = Inches(0.8)  # Themes
    header_cells[1].width = Inches(0.8)  # Banks
    header_cells[2].width = Inches(8.4)  # Questions

    for cell in header_cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="002060"/>'.format(nsdecls("w")))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)

    row_idx = 1
    for category in sorted(category_bank_questions.keys()):
        banks_in_category = category_bank_questions[category]
        category_start_row = row_idx
        category_end_row = category_start_row + len(banks_in_category) - 1

        for bank_name in sorted(banks_in_category.keys()):
            bank_data = banks_in_category[bank_name]
            row = table.rows[row_idx]
            is_first_in_category = row_idx == category_start_row
            is_last_in_category = row_idx == category_end_row

            row.cells[0].width = Inches(0.8)
            row.cells[1].width = Inches(0.8)
            row.cells[2].width = Inches(8.4)

            if is_first_in_category:
                row.cells[0].text = category
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.bold = True
            else:
                row.cells[0].text = ""  # Empty for subsequent rows in same category

            row.cells[1].text = bank_data["ticker"]
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.bold = True

            first_question = True
            for i, question in enumerate(bank_data["questions"]):
                question_text = question.get("verbatim_question", "")
                question_text = _clean_text_for_word(question_text)

                if first_question:
                    q_para = row.cells[2].paragraphs[0]
                    q_para.style = "List Bullet"
                    first_question = False
                else:
                    q_para = row.cells[2].add_paragraph(style="List Bullet")
                _add_text_with_metric_bold(q_para, question_text, font_size=9)
                q_para.paragraph_format.space_before = Pt(0)
                q_para.paragraph_format.space_after = Pt(0)

            if is_last_in_category:
                for col_idx in [0, 1, 2]:
                    cell = row.cells[col_idx]
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = parse_xml(
                        r"<w:tcBorders {}>"
                        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
                        r"</w:tcBorders>".format(nsdecls("w"))
                    )
                    tcPr.append(tcBorders)
            else:
                for col_idx in [1, 2]:
                    cell = row.cells[col_idx]
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = parse_xml(
                        r"<w:tcBorders {}>"
                        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
                        r"</w:tcBorders>".format(nsdecls("w"))
                    )
                    tcPr.append(tcBorders)
                cell = row.cells[0]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = parse_xml(
                    r"<w:tcBorders {}>"
                    r'<w:bottom w:val="none"/>'
                    r"</w:tcBorders>".format(nsdecls("w"))
                )
                tcPr.append(tcBorders)

            row_idx += 1

    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = parse_xml(r"<w:tblPr {}/> ".format(nsdecls("w")))
        tbl.insert(0, tbl_pr)

    tbl_borders = parse_xml(
        r"<w:tblBorders {}>"
        r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r"</w:tblBorders>".format(nsdecls("w"))
    )
    tbl_pr.append(tbl_borders)
