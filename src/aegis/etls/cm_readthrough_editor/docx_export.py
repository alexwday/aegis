"""DOCX export helpers for the CM readthrough editor ETL."""

from __future__ import annotations

import re
from collections import OrderedDict
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


_METRIC_PATTERN = re.compile(
    r"(-?\$[\d,]+(?:\.\d+)?\s*(?:MM|BN|TN|K|M|B)?\b|\d+(?:\.\d+)?\s*bps\b|\d+(?:\.\d+)?%)",
    re.IGNORECASE,
)
_OUTLOOK_COLUMN_WIDTHS = (Inches(0.75), Inches(9.45))
_QA_COLUMN_WIDTHS = (Inches(1.25), Inches(0.7), Inches(8.25))


def _clean_text(value: Any) -> str:
    """Normalize a text-ish value for Word output."""
    return " ".join(str(value or "").replace("\x00", "").split())


def _bank_state(report_state: Dict[str, Any], bank_id: str) -> Dict[str, Any]:
    """Return the mutable editor state for a bank."""
    return report_state.get("bank_states", {}).get(bank_id, {}) or {}


def _report_bank_ids(report_state: Dict[str, Any]) -> List[str]:
    """Return bank IDs in the same order used by the HTML report preview."""
    ordered = report_state.get("report_bank_order") or []
    if ordered:
        return [str(bank_id) for bank_id in ordered]
    return [str(bank_id) for bank_id in report_state.get("banks", {})]


def _bank_ticker_label(bank_id: str, bank_data: Dict[str, Any]) -> str:
    """Return the bank ticker label shown in the report preview."""
    raw = (
        bank_data.get("symbol")
        or bank_data.get("ticker")
        or bank_data.get("full_ticker")
        or bank_id
    )
    return re.sub(r"-[A-Z]{2,4}$", "", _clean_text(raw), flags=re.IGNORECASE)


def _bucket_display_title(report_state: Dict[str, Any], bucket: Dict[str, Any]) -> str:
    """Return the category title shown in the report preview."""
    bucket_id = str(bucket.get("id") or "")
    user_titles = report_state.get("bucket_user_titles", {}) or {}
    return _clean_text(user_titles.get(bucket_id) or bucket.get("name")) or "Uncategorized"


def _sentence_status(sentence: Dict[str, Any], bank_state: Dict[str, Any]) -> str:
    """Return sentence review status after editor overrides."""
    sid = str(sentence.get("sid") or "")
    overrides = bank_state.get("sentence_status_overrides", {}) or {}
    return str(overrides.get(sid) or sentence.get("status") or "").strip()


def _effective_bucket_id(sentence: Dict[str, Any], bank_state: Dict[str, Any]) -> str:
    """Return sentence bucket after editor reassignment overrides."""
    sid = str(sentence.get("sid") or "")
    user_primary = bank_state.get("sentence_user_primary", {}) or {}
    return str(
        user_primary.get(sid)
        or sentence.get("selected_bucket_id")
        or sentence.get("primary")
        or ""
    ).strip()


def _subquote_sentences(subquote: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Return the sentence list used as visible text for a report finding."""
    if subquote.get("type") == "md":
        return subquote.get("sentences", []) or []
    if subquote.get("render_mode") == "question":
        return subquote.get("question_sentences", []) or []
    return subquote.get("answer_sentences", []) or []


def _subquote_text(subquote: Dict[str, Any], bank_state: Dict[str, Any]) -> str:
    """Return plain text matching the CM report preview's visible finding text."""
    overrides = bank_state.get("subquote_text_overrides", {}) or {}
    override = overrides.get(subquote.get("id"))
    if isinstance(override, str) and override.strip():
        return _clean_text(override)
    return _clean_text(
        " ".join(_clean_text(sentence.get("text")) for sentence in _subquote_sentences(subquote))
    )


def _subquote_importance(subquote: Dict[str, Any]) -> float:
    """Return the preview sorting importance for a subquote."""
    values = [
        float(sentence.get("importance_score") or 0)
        for sentence in _subquote_sentences(subquote)
    ]
    return max(values) if values else 0.0


def _subquote_bucket_score(subquote: Dict[str, Any]) -> float:
    """Return the preview sorting bucket score for a subquote."""
    bucket_id = str(subquote.get("effective_bucket") or "")
    scores = []
    for sentence in _subquote_sentences(subquote):
        sentence_scores = sentence.get("scores", {}) or {}
        scores.append(float(sentence_scores.get(bucket_id) or 0))
    return max(scores) if scores else 0.0


def _cluster_key(subquote: Dict[str, Any]) -> str:
    """Return the same coarse cluster key used by the preview ordering."""
    if subquote.get("type") == "qa":
        return f"qa::{subquote.get('conv_id') or subquote.get('id')}"
    return "md::{}||{}||{}".format(
        _clean_text(subquote.get("speaker")),
        _clean_text(subquote.get("speaker_title")),
        _clean_text(subquote.get("speaker_affiliation")),
    )


def _cluster_subquotes(subquotes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Cluster sorted subquotes the same way the report preview does."""
    order: List[str] = []
    grouped: Dict[str, List[Dict[str, Any]]] = OrderedDict()
    for subquote in subquotes:
        key = _cluster_key(subquote)
        if key not in grouped:
            grouped[key] = []
            order.append(key)
        grouped[key].append(subquote)

    clustered: List[Dict[str, Any]] = []
    for key in order:
        values = grouped[key]
        if key.startswith("qa::") and len(values) > 1:
            clustered.extend(sorted(values, key=lambda item: str(item.get("id") or "")))
        else:
            clustered.extend(values)
    return clustered


def _order_subquotes(
    subquotes: List[Dict[str, Any]],
    bucket_id: str,
    bank_state: Dict[str, Any],
) -> List[Dict[str, Any]]:
    """Return report subquotes in preview order for one bucket."""
    order = (bank_state.get("bucket_subquote_order", {}) or {}).get(bucket_id) or []
    if order:
        by_id = {str(subquote.get("id")): subquote for subquote in subquotes}
        ordered = [by_id[subquote_id] for subquote_id in order if subquote_id in by_id]
        ordered_ids = {str(subquote.get("id")) for subquote in ordered}
        ordered.extend(
            subquote
            for subquote in subquotes
            if str(subquote.get("id")) not in ordered_ids
        )
        return ordered

    ranked = sorted(
        subquotes,
        key=lambda item: (
            -_subquote_importance(item),
            -_subquote_bucket_score(item),
            str(item.get("id") or ""),
        ),
    )
    return _cluster_subquotes(ranked)


def _qa_answer_speaker_key(conversation: Dict[str, Any], sentence: Dict[str, Any]) -> str:
    """Return the QA answer speaker grouping key used by the preview."""
    speaker = sentence.get("speaker") or conversation.get("executive_name") or "Executive"
    title = sentence.get("speaker_title") or conversation.get("executive_title") or ""
    affiliation = (
        sentence.get("speaker_affiliation")
        or conversation.get("executive_affiliation")
        or ""
    )
    return "||".join([_clean_text(speaker), _clean_text(title), _clean_text(affiliation)])


def _qa_question_speaker_key(conversation: Dict[str, Any]) -> str:
    """Return the QA question speaker grouping key used by the preview."""
    return "||".join(
        [
            _clean_text(conversation.get("analyst_name") or "Analyst"),
            "",
            _clean_text(conversation.get("analyst_affiliation")),
        ]
    )


def _append_sentence_to_subquote(
    subquote: Dict[str, Any],
    sentence: Dict[str, Any],
) -> None:
    """Append a sentence to the correct subquote sentence collection."""
    if subquote.get("type") == "md":
        subquote.setdefault("sentences", []).append(sentence)
    elif subquote.get("render_mode") == "question":
        subquote.setdefault("question_sentences", []).append(sentence)
    else:
        subquote.setdefault("answer_sentences", []).append(sentence)


def _build_md_subquotes(
    block: Dict[str, Any],
    bank_state: Dict[str, Any],
) -> List[Dict[str, Any]]:
    """Build MD subquotes using the report-preview grouping rules."""
    subquotes: List[Dict[str, Any]] = []
    current: Dict[str, Any] | None = None

    for sentence in block.get("sentences", []) or []:
        if _sentence_status(sentence, bank_state) != "selected":
            if current:
                subquotes.append(current)
                current = None
            continue

        bucket_id = _effective_bucket_id(sentence, bank_state)
        if not bucket_id:
            if current:
                subquotes.append(current)
                current = None
            continue

        tentative_id = str(current.get("id")) if current else f"SQ_{sentence.get('sid')}"
        bucket_id = (
            (bank_state.get("subquote_bucket_overrides", {}) or {}).get(tentative_id)
            or bucket_id
        )
        if not current or current.get("effective_bucket") != bucket_id:
            if current:
                subquotes.append(current)
            current = {
                "id": f"SQ_{sentence.get('sid')}",
                "type": "md",
                "block_id": block.get("id"),
                "speaker": block.get("speaker", ""),
                "speaker_title": block.get("speaker_title", ""),
                "speaker_affiliation": block.get("speaker_affiliation", ""),
                "effective_bucket": bucket_id,
                "sentences": [sentence],
            }
        else:
            _append_sentence_to_subquote(current, sentence)

    if current:
        subquotes.append(current)
    return subquotes


def _build_qa_subquotes_for_mode(
    conversation: Dict[str, Any],
    sentences: List[Dict[str, Any]],
    bank_state: Dict[str, Any],
    mode: str,
) -> List[Dict[str, Any]]:
    """Build QA subquotes using the report-preview grouping rules."""
    subquotes: List[Dict[str, Any]] = []
    current: Dict[str, Any] | None = None
    question_mode = mode == "question"

    for sentence in sentences:
        if _sentence_status(sentence, bank_state) != "selected":
            if current:
                subquotes.append(current)
                current = None
            continue

        bucket_id = _effective_bucket_id(sentence, bank_state)
        if not bucket_id:
            if current:
                subquotes.append(current)
                current = None
            continue

        speaker_key = (
            _qa_question_speaker_key(conversation)
            if question_mode
            else _qa_answer_speaker_key(conversation, sentence)
        )
        if (
            not current
            or current.get("effective_bucket") != bucket_id
            or current.get("_speaker_key") != speaker_key
        ):
            if current:
                current.pop("_speaker_key", None)
                subquotes.append(current)
            current = {
                "id": f"SQ_{sentence.get('sid')}",
                "type": "qa",
                "render_mode": "question" if question_mode else "answer",
                "conv_id": conversation.get("id"),
                "analyst_name": conversation.get("analyst_name"),
                "analyst_affiliation": conversation.get("analyst_affiliation"),
                "executive_name": (
                    conversation.get("executive_name") if question_mode else sentence.get("speaker")
                ),
                "executive_title": (
                    conversation.get("executive_title")
                    if question_mode
                    else sentence.get("speaker_title")
                ),
                "executive_affiliation": (
                    conversation.get("executive_affiliation")
                    if question_mode
                    else sentence.get("speaker_affiliation")
                ),
                "analyst_question_summary": conversation.get("analyst_question_summary", ""),
                "effective_bucket": bucket_id,
                "question_sentences": [sentence] if question_mode else [],
                "answer_sentences": [] if question_mode else [sentence],
                "_speaker_key": speaker_key,
            }
        else:
            _append_sentence_to_subquote(current, sentence)

    if current:
        current.pop("_speaker_key", None)
        subquotes.append(current)
    return subquotes


def _build_qa_subquotes(
    conversation: Dict[str, Any],
    bank_state: Dict[str, Any],
) -> List[Dict[str, Any]]:
    """Build all QA subquotes for one conversation."""
    return [
        *_build_qa_subquotes_for_mode(
            conversation,
            conversation.get("question_sentences", []) or [],
            bank_state,
            "question",
        ),
        *_build_qa_subquotes_for_mode(
            conversation,
            conversation.get("answer_sentences", []) or [],
            bank_state,
            "answer",
        ),
    ]


def _all_subquotes(bank_data: Dict[str, Any], bank_state: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Build all report-preview subquotes for one bank."""
    subquotes: List[Dict[str, Any]] = []
    for block in bank_data.get("md_blocks", []) or []:
        subquotes.extend(_build_md_subquotes(block, bank_state))
    for conversation in bank_data.get("qa_conversations", []) or []:
        subquotes.extend(_build_qa_subquotes(conversation, bank_state))
    return subquotes


def _add_metric_runs(paragraph, text: str, *, size: int = 9) -> None:
    """Add text to a paragraph, bolding financial metrics."""
    cursor = 0
    clean = _clean_text(text)
    for match in _METRIC_PATTERN.finditer(clean):
        if match.start() > cursor:
            end = match.start()
            run = paragraph.add_run(clean[cursor:end])
            run.font.size = Pt(size)
        run = paragraph.add_run(match.group(0))
        run.font.size = Pt(size)
        run.font.bold = True
        cursor = match.end()
    if cursor < len(clean):
        run = paragraph.add_run(clean[cursor:])
        run.font.size = Pt(size)


def _shade_cell(cell, fill: str = "002060") -> None:
    """Apply solid fill shading to one table cell."""
    shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), fill))
    cell._tc.get_or_add_tcPr().append(shading)  # pylint: disable=protected-access


def _replace_child(parent, child) -> None:
    """Replace an OOXML child element with the same tag."""
    existing = parent.find(child.tag)
    if existing is not None:
        parent.remove(existing)
    parent.append(child)


def _border_element(name: str, *, size: int, color: str) -> OxmlElement:
    """Build one Word border element."""
    element = OxmlElement(f"w:{name}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(size))
    element.set(qn("w:space"), "0")
    element.set(qn("w:color"), color)
    return element


def _set_cell_borders(cell, *, size: int = 4, color: str = "D9D0BC") -> None:
    """Apply visible borders to one table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()  # pylint: disable=protected-access
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        borders.append(_border_element(edge, size=size, color=color))
    _replace_child(tc_pr, borders)


def _set_cell_width(cell, width: Inches) -> None:
    """Set an exact Word cell width."""
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()  # pylint: disable=protected-access
    tc_width = tc_pr.tcW
    tc_width.type = "dxa"
    tc_width.w = round(width.inches * 1440)


def _set_table_width(table, widths: Tuple[Inches, ...]) -> None:
    """Apply fixed table layout and exact column widths to all rows."""
    table.autofit = False
    table.allow_autofit = False
    total_width = sum(width.inches for width in widths)
    tbl_pr = table._tbl.tblPr  # pylint: disable=protected-access

    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(round(total_width * 1440)))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            _set_cell_width(cell, width)
            _set_cell_borders(cell)


def _set_table_borders(table) -> None:
    """Apply visible outer and inner table borders."""
    tbl_pr = table._tbl.tblPr  # pylint: disable=protected-access
    borders = OxmlElement("w:tblBorders")
    for edge, size, color in (
        ("top", 6, "B8AD94"),
        ("left", 6, "B8AD94"),
        ("bottom", 6, "B8AD94"),
        ("right", 6, "B8AD94"),
        ("insideH", 4, "D9D0BC"),
        ("insideV", 4, "D9D0BC"),
    ):
        borders.append(_border_element(edge, size=size, color=color))
    _replace_child(tbl_pr, borders)


def _repeat_header_row(row) -> None:
    """Mark a Word table row to repeat at the top of each page."""
    tr_pr = row._tr.get_or_add_trPr()  # pylint: disable=protected-access
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _set_header_cell(cell, text: str, width: Inches) -> None:
    """Format a table header cell."""
    cell.text = text
    _set_cell_width(cell, width)
    _set_cell_borders(cell, size=6, color="B8AD94")
    _shade_cell(cell)
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)


def _main_title(report_state: Dict[str, Any]) -> str:
    """Return the report title used by the editor state."""
    meta = report_state.get("meta", {})
    return _clean_text(meta.get("cm_main_title")) or (
        f"Read Through For Capital Markets: {meta.get('fiscal_quarter')}/"
        f"{str(meta.get('fiscal_year'))[-2:]} Select Banks"
    )


def _report_period_label(report_state: Dict[str, Any]) -> str:
    """Return the report fiscal period label."""
    meta = report_state.get("meta", {})
    return " ".join(
        _clean_text(value)
        for value in (meta.get("fiscal_quarter"), meta.get("fiscal_year"))
        if _clean_text(value)
    )


def _cover_meta(report_state: Dict[str, Any]) -> str:
    """Return the cover-page metadata line used by the editor export."""
    meta = report_state.get("meta", {})
    cover_entity = _clean_text(meta.get("cover_entity"))
    period = _report_period_label(report_state)
    return " - ".join(part for part in (cover_entity, period) if part)


def _cover_title(report_state: Dict[str, Any]) -> str:
    """Return the cover-page report title."""
    meta = report_state.get("meta", {})
    return _clean_text(meta.get("report_title")) or "Capital Markets Readthrough"


def _section_subtitle(report_state: Dict[str, Any], key: str, fallback: str) -> str:
    """Return a section subtitle from editor state metadata."""
    subtitles = report_state.get("meta", {}).get("section_subtitles", {})
    return _clean_text(subtitles.get(key)) or fallback


def _set_portrait_section(section) -> None:
    """Apply portrait title/contents page layout."""
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def _set_landscape_section(section) -> None:
    """Apply landscape CM report formatting."""
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)


def _add_paragraph_bottom_border(paragraph, *, size: int, color: str) -> None:
    """Add a bottom border to a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()  # pylint: disable=protected-access
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    _replace_child(borders, _border_element("bottom", size=size, color=color))


def _add_cover_and_contents(doc: Document, report_state: Dict[str, Any]) -> None:
    """Add title page and contents before the landscape report tables."""
    cover_meta = _cover_meta(report_state) or "Capital Markets Readthrough"
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.paragraph_format.space_before = Pt(24)
    meta_para.paragraph_format.space_after = Pt(42)
    _add_paragraph_bottom_border(meta_para, size=4, color="D9D0BC")
    meta_run = meta_para.add_run(cover_meta.upper())
    meta_run.font.bold = True
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = RGBColor(122, 115, 100)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(72)
    title_para.paragraph_format.space_after = Pt(24)
    title_run = title_para.add_run(_cover_title(report_state))
    title_run.font.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(15, 27, 45)

    accent_para = doc.add_paragraph()
    accent_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    accent_para.paragraph_format.left_indent = Inches(3.0)
    accent_para.paragraph_format.right_indent = Inches(3.0)
    accent_para.paragraph_format.space_after = Pt(96)
    _add_paragraph_bottom_border(accent_para, size=16, color="B25A1E")

    contents_title = doc.add_paragraph()
    contents_title.paragraph_format.space_after = Pt(12)
    contents_run = contents_title.add_run("Contents")
    contents_run.font.bold = True
    contents_run.font.size = Pt(18)
    contents_run.font.color.rgb = RGBColor(15, 27, 45)
    _add_paragraph_bottom_border(contents_title, size=4, color="D9D0BC")

    contents = [
        ("I.", "Outlook", _section_subtitle(report_state, "outlook", "Outlook")),
        ("II.", "Q&A", _section_subtitle(report_state, "qa", "Q&A")),
    ]
    for ordinal, section_name, subtitle in contents:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        ordinal_run = paragraph.add_run(f"{ordinal}   ")
        ordinal_run.font.bold = True
        ordinal_run.font.color.rgb = RGBColor(178, 90, 30)
        name_run = paragraph.add_run(section_name)
        name_run.font.bold = True
        name_run.font.size = Pt(12)
        if subtitle and subtitle != section_name:
            subtitle_para = doc.add_paragraph()
            subtitle_para.paragraph_format.left_indent = Inches(0.35)
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.italic = True
            subtitle_run.font.size = Pt(10)
            subtitle_run.font.color.rgb = RGBColor(74, 82, 104)


def _add_title(doc: Document, report_state: Dict[str, Any], subtitle: str) -> None:
    """Add the common CM report title and subtitle."""
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(_main_title(report_state))
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 32, 96)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.italic = True
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _setup_document(doc: Document) -> None:
    """Apply initial title/contents document formatting."""
    _set_portrait_section(doc.sections[0])


def _start_landscape_body(doc: Document) -> None:
    """Start the landscape report body after the title/contents section."""
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    _set_landscape_section(section)


def _bucket_lookup(report_state: Dict[str, Any]) -> Dict[str, Dict[str, Any]]:
    """Return buckets keyed by id."""
    return {
        str(bucket.get("id")): bucket
        for bucket in report_state.get("buckets", [])
        if bucket.get("id")
    }


def _iter_outlook_findings(
    subquotes: List[Dict[str, Any]],
    bucket_id: str,
) -> Iterable[Dict[str, Any]]:
    """Yield selected Outlook subquotes from MD and management answers."""
    for subquote in subquotes:
        if subquote.get("effective_bucket") != bucket_id:
            continue
        if subquote.get("type") == "md" or subquote.get("render_mode") == "answer":
            yield subquote


def _iter_question_findings(
    subquotes: List[Dict[str, Any]],
    bucket_id: str,
) -> Iterable[Dict[str, Any]]:
    """Yield selected Q&A question subquotes."""
    for subquote in subquotes:
        if (
            subquote.get("type") == "qa"
            and subquote.get("render_mode") == "question"
            and subquote.get("effective_bucket") == bucket_id
        ):
            yield subquote


def _collect_outlook(
    report_state: Dict[str, Any],
) -> "OrderedDict[str, Dict[str, List[Tuple[Dict[str, Any], Dict[str, Any]]]]]":
    """Collect Outlook findings as bank -> bucket -> findings."""
    buckets = _bucket_lookup(report_state)
    grouped: "OrderedDict[str, Dict[str, List[Tuple[Dict[str, Any], Dict[str, Any]]]]]" = (
        OrderedDict()
    )
    for bank_key in _report_bank_ids(report_state):
        bank_state = _bank_state(report_state, bank_key)
        bank_data = report_state.get("banks", {}).get(bank_key, {})
        bank_label = _bank_ticker_label(bank_key, bank_data)
        subquotes = _all_subquotes(bank_data, bank_state)
        for bucket_id, bucket in buckets.items():
            if bucket.get("report_section") != "Outlook":
                continue
            findings = _order_subquotes(
                list(_iter_outlook_findings(subquotes, bucket_id)),
                bucket_id,
                bank_state,
            )
            if findings:
                annotated_findings = []
                for finding in findings:
                    annotated = dict(finding)
                    annotated["_bank_id"] = bank_key
                    annotated_findings.append(annotated)
                grouped.setdefault(bank_label, OrderedDict())[bucket_id] = [
                    (bucket, finding) for finding in annotated_findings
                ]
    return grouped


def _collect_questions(
    report_state: Dict[str, Any],
) -> "OrderedDict[str, Dict[str, Any]]":
    """Collect Q&A findings as bucket/category -> bank -> questions."""
    buckets = _bucket_lookup(report_state)
    grouped: "OrderedDict[str, Dict[str, Any]]" = OrderedDict()
    for bucket_id, bucket in buckets.items():
        if bucket.get("report_section") != "Q&A":
            continue
        for bank_key in _report_bank_ids(report_state):
            bank_state = _bank_state(report_state, bank_key)
            bank_data = report_state.get("banks", {}).get(bank_key, {})
            bank_label = _bank_ticker_label(bank_key, bank_data)
            questions = _order_subquotes(
                list(_iter_question_findings(_all_subquotes(bank_data, bank_state), bucket_id)),
                bucket_id,
                bank_state,
            )
            if questions:
                annotated_questions = []
                for question in questions:
                    annotated = dict(question)
                    annotated["_bank_id"] = bank_key
                    annotated_questions.append(annotated)
                if bucket_id not in grouped:
                    grouped[bucket_id] = {"bucket": bucket, "bank_groups": OrderedDict()}
                grouped[bucket_id]["bank_groups"].setdefault(bank_label, []).extend(
                    annotated_questions
                )
    return grouped


def _add_outlook_section(doc: Document, report_state: Dict[str, Any]) -> None:
    """Add the Outlook table."""
    _add_title(
        doc,
        report_state,
        _section_subtitle(report_state, "outlook", "Outlook: Capital markets activity"),
    )
    outlook = _collect_outlook(report_state)
    table = doc.add_table(rows=max(len(outlook), 1) + 1, cols=2)
    _set_table_width(table, _OUTLOOK_COLUMN_WIDTHS)
    _set_table_borders(table)
    _repeat_header_row(table.rows[0])
    _set_header_cell(table.rows[0].cells[0], "Banks", _OUTLOOK_COLUMN_WIDTHS[0])
    _set_header_cell(
        table.rows[0].cells[1],
        "Investment Banking and Trading Outlook",
        _OUTLOOK_COLUMN_WIDTHS[1],
    )

    if not outlook:
        table.rows[1].cells[0].text = "-"
        table.rows[1].cells[1].text = "No selected Outlook findings available."
        return

    for row_idx, (bank_label, bucket_groups) in enumerate(outlook.items(), start=1):
        row = table.rows[row_idx]
        row.cells[0].text = bank_label
        for paragraph in row.cells[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

        content_cell = row.cells[1]
        first_para = True
        for bucket_findings in bucket_groups.values():
            bucket = bucket_findings[0][0]
            category_para = (
                content_cell.paragraphs[0] if first_para else content_cell.add_paragraph()
            )
            first_para = False
            category_run = category_para.add_run(f"{_bucket_display_title(report_state, bucket)}:")
            category_run.font.bold = True
            category_run.font.size = Pt(9)
            headline = _clean_text(bucket.get("generated_headline"))
            if headline:
                headline_para = content_cell.add_paragraph()
                headline_run = headline_para.add_run(headline)
                headline_run.font.italic = True
                headline_run.font.size = Pt(9)
                headline_run.font.color.rgb = RGBColor(74, 82, 104)
            for _, finding in bucket_findings:
                para = content_cell.add_paragraph(style="List Bullet")
                _add_metric_runs(
                    para,
                    _subquote_text(finding, _bank_state(report_state, finding.get("_bank_id", ""))),
                    size=9,
                )
                para.paragraph_format.space_after = Pt(0)


def _add_qa_section(doc: Document, report_state: Dict[str, Any]) -> None:
    """Add the Q&A table."""
    doc.add_page_break()
    _add_title(
        doc,
        report_state,
        _section_subtitle(report_state, "qa", "Conference calls: Capital markets questions"),
    )
    grouped = _collect_questions(report_state)
    row_count = sum(
        len(group["bank_groups"]) for group in grouped.values()
    ) or 1
    table = doc.add_table(rows=row_count + 1, cols=3)
    _set_table_width(table, _QA_COLUMN_WIDTHS)
    _set_table_borders(table)
    _repeat_header_row(table.rows[0])
    _set_header_cell(table.rows[0].cells[0], "Themes", _QA_COLUMN_WIDTHS[0])
    _set_header_cell(table.rows[0].cells[1], "Banks", _QA_COLUMN_WIDTHS[1])
    _set_header_cell(table.rows[0].cells[2], "Relevant Questions", _QA_COLUMN_WIDTHS[2])

    if not grouped:
        table.rows[1].cells[0].text = "-"
        table.rows[1].cells[1].text = "-"
        table.rows[1].cells[2].text = "No selected Q&A findings available."
        return

    row_idx = 1
    for group in grouped.values():
        bucket = group["bucket"]
        category = _bucket_display_title(report_state, bucket)
        bank_groups = group["bank_groups"]
        first_category_row = True
        for bank_label, questions in bank_groups.items():
            row = table.rows[row_idx]
            row.cells[0].text = category if first_category_row else ""
            row.cells[1].text = bank_label
            for paragraph in row.cells[0].paragraphs + row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

            first_question = True
            for question in questions:
                para = (
                    row.cells[2].paragraphs[0] if first_question else row.cells[2].add_paragraph()
                )
                para.style = "List Bullet"
                _add_metric_runs(
                    para,
                    _subquote_text(
                        question,
                        _bank_state(report_state, question.get("_bank_id", "")),
                    ),
                    size=9,
                )
                first_question = False
            first_category_row = False
            row_idx += 1


def _validate_document(doc: Document) -> None:
    """Validate the generated DOCX has meaningful report content."""
    has_title = any("Read Through For Capital Markets:" in p.text for p in doc.paragraphs)
    has_table = bool(doc.tables)
    has_body = any(
        p.text.strip() and "Read Through For Capital Markets:" not in p.text for p in doc.paragraphs
    ) or any(cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells)
    missing = []
    if not has_title:
        missing.append("report title")
    if not has_table:
        missing.append("table content")
    if not has_body:
        missing.append("body text")
    if missing:
        raise ValueError(f"Document missing required content: {', '.join(missing)}")


def create_cm_readthrough_docx_from_state(
    *,
    report_state: Dict[str, Any],
    output_path: str,
) -> None:
    """Create a DOCX CM readthrough report from the editor's initial report state."""
    doc = Document()
    _setup_document(doc)
    _add_cover_and_contents(doc, report_state)
    _start_landscape_body(doc)
    _add_outlook_section(doc, report_state)
    _add_qa_section(doc, report_state)
    _validate_document(doc)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
