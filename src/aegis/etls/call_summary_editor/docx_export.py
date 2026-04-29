"""DOCX export helpers for the call summary editor ETL."""

from __future__ import annotations

import re
from collections import OrderedDict
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


_METRIC_PATTERN = re.compile(
    r"(-?\$[\d,]+(?:\.\d+)?\s*(?:MM|BN|TN|K|M|B)?\b|\d+(?:\.\d+)?\s*bps\b|\d+(?:\.\d+)?%)",
    re.IGNORECASE,
)


def _clean_text(value: Any) -> str:
    """Normalize a text-ish value for Word output."""
    return " ".join(str(value or "").replace("\x00", "").split())


def _bucket_id(sentence: Dict[str, Any]) -> str:
    """Return the active bucket id for a selected sentence/finding."""
    return str(sentence.get("selected_bucket_id") or sentence.get("primary") or "").strip()


def _add_page_numbers(doc: Document) -> None:
    """Add page numbers to document footers."""
    for section in doc.sections:
        footer_para = section.footer.paragraphs[0]
        footer_para.clear()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer_para.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._element.append(fld_begin)  # pylint: disable=protected-access
        run._element.append(instr)  # pylint: disable=protected-access
        run._element.append(fld_end)  # pylint: disable=protected-access


def _setup_document(doc: Document) -> None:
    """Apply compact report formatting."""
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.5)
    _add_page_numbers(doc)


def _add_metric_runs(paragraph, text: str, *, size: Pt, italic: bool = False) -> None:
    """Add text to a paragraph, bolding financial metrics."""
    cursor = 0
    for match in _METRIC_PATTERN.finditer(text):
        if match.start() > cursor:
            end = match.start()
            run = paragraph.add_run(text[cursor:end])
            run.font.size = size
            run.italic = italic
        run = paragraph.add_run(match.group(0))
        run.font.size = size
        run.font.bold = True
        run.italic = italic
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        run.font.size = size
        run.italic = italic


def _sentence_speaker(
    sentence: Dict[str, Any],
    *,
    block: Optional[Dict[str, Any]] = None,
    conversation: Optional[Dict[str, Any]] = None,
) -> str:
    """Resolve a useful speaker attribution for a finding."""
    for value in (
        sentence.get("speaker"),
        (block or {}).get("speaker"),
        (conversation or {}).get("executive_name"),
        (conversation or {}).get("analyst_name"),
    ):
        cleaned = _clean_text(value)
        if cleaned:
            return cleaned
    return "Unknown"


def _iter_selected_findings(
    bank_data: Dict[str, Any],
    bucket_id: str,
) -> Iterable[Dict[str, str]]:
    """Yield selected initial-report findings for one bucket."""
    for block in bank_data.get("md_blocks", []):
        for sentence in block.get("sentences", []):
            if sentence.get("status") != "selected" or _bucket_id(sentence) != bucket_id:
                continue
            yield {
                "statement": _clean_text(sentence.get("condensed") or sentence.get("text")),
                "quote": _clean_text(sentence.get("verbatim_text") or sentence.get("text")),
                "speaker": _sentence_speaker(sentence, block=block),
            }

    for conversation in bank_data.get("qa_conversations", []):
        for sentence in conversation.get("answer_sentences", []):
            if sentence.get("status") != "selected" or _bucket_id(sentence) != bucket_id:
                continue
            yield {
                "statement": _clean_text(sentence.get("condensed") or sentence.get("text")),
                "quote": _clean_text(sentence.get("verbatim_text") or sentence.get("text")),
                "speaker": _sentence_speaker(sentence, conversation=conversation),
            }


def _collect_sections(report_state: Dict[str, Any]) -> "OrderedDict[str, List[Dict[str, Any]]]":
    """Collect selected findings by report section and bucket."""
    banks = report_state.get("banks", {})
    sections: "OrderedDict[str, List[Dict[str, Any]]]" = OrderedDict()

    for bucket in report_state.get("buckets", []):
        bucket_id = str(bucket.get("id", "")).strip()
        if not bucket_id:
            continue
        findings: List[Dict[str, str]] = []
        for bank_data in banks.values():
            findings.extend(_iter_selected_findings(bank_data, bucket_id))
        if not findings:
            continue
        section_name = _clean_text(bucket.get("report_section")) or "Results Summary"
        sections.setdefault(section_name, []).append(
            {
                "title": _clean_text(bucket.get("name")) or "Untitled",
                "findings": findings,
            }
        )

    return sections


def _add_title(doc: Document, *, quarter: str, fiscal_year: int, bank_symbol: str) -> None:
    """Add the report title."""
    title = doc.add_heading(
        f"{quarter}/{str(fiscal_year)[-2:]} Results and Call Summary - {bank_symbol}", 0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(14)
        run.font.bold = True
    title.paragraph_format.space_after = Pt(4)


def _add_section(doc: Document, section_name: str, *, is_first: bool) -> None:
    """Add a level-one report section heading."""
    heading = doc.add_heading(section_name, level=1)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.page_break_before = not is_first


def _add_category(doc: Document, category: Dict[str, Any]) -> None:
    """Add one category and its selected findings."""
    heading = doc.add_heading(category["title"], level=2)
    for run in heading.runs:
        run.font.size = Pt(10)
        run.font.bold = True
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.keep_with_next = True

    for finding in category["findings"]:
        statement = finding["statement"]
        quote = finding["quote"]
        speaker = finding["speaker"]
        if not statement and not quote:
            continue

        bullet = doc.add_paragraph(style="List Bullet")
        _add_metric_runs(bullet, statement or quote, size=Pt(9))
        bullet.paragraph_format.space_after = Pt(2)
        bullet.paragraph_format.keep_with_next = True

        if quote:
            evidence = doc.add_paragraph()
            evidence.paragraph_format.left_indent = Inches(0.75)
            evidence.paragraph_format.right_indent = Inches(0.5)
            evidence.paragraph_format.space_after = Pt(1)
            evidence.add_run('"').italic = True
            _add_metric_runs(evidence, quote, size=Pt(8), italic=True)
            evidence.add_run('"').italic = True
            speaker_run = evidence.add_run(f" - {speaker}")
            speaker_run.font.size = Pt(7)
            speaker_run.font.color.rgb = RGBColor(96, 96, 96)


def _validate_document(doc: Document) -> None:
    """Validate the generated DOCX has report content."""
    has_section = any(p.style and p.style.name == "Heading 1" for p in doc.paragraphs)
    has_category = any(p.style and p.style.name == "Heading 2" for p in doc.paragraphs)
    has_body = any(
        p.text.strip() and (not p.style or p.style.name not in {"Title", "Heading 1", "Heading 2"})
        for p in doc.paragraphs
    )
    missing = []
    if not has_section:
        missing.append("section heading")
    if not has_category:
        missing.append("category heading")
    if not has_body:
        missing.append("body text")
    if missing:
        raise ValueError(f"Document missing required content: {', '.join(missing)}")


def create_call_summary_docx_from_state(
    *,
    report_state: Dict[str, Any],
    output_path: str,
    bank_symbol: str,
) -> None:
    """Create a DOCX call-summary report from the editor's initial report state."""
    meta = report_state.get("meta", {})
    fiscal_year = int(meta.get("fiscal_year"))
    quarter = str(meta.get("fiscal_quarter"))
    sections = _collect_sections(report_state)
    if not sections:
        raise ValueError("No selected report findings available for DOCX output")

    doc = Document()
    _setup_document(doc)
    _add_title(doc, quarter=quarter, fiscal_year=fiscal_year, bank_symbol=bank_symbol)

    for idx, (section_name, categories) in enumerate(sections.items()):
        _add_section(doc, section_name, is_first=idx == 0)
        for category in categories:
            _add_category(doc, category)

    _validate_document(doc)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
